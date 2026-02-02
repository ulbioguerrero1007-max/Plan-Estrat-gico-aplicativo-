import streamlit as st
import re
import google.generativeai as genai
import pandas as pd
import sqlite3
import io
from io import StringIO, BytesIO  # <-- Asegúrate de tener esta línea
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4, letter
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, Table, TableStyle, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.colors import navy, grey, red, green, black
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT, TA_RIGHT
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import matplotlib.pyplot as plt
import plotly.express as px
import plotly.graph_objects as go
from plotly.subplots import make_subplots
import numpy as np
import unicodedata
import time
from supabase import create_client, Client
from reportlab.lib.colors import HexColor, Color, white

class ColorPalette:
    """Paleta de colores profesional para documentos ejecutivos"""
    PRIMARY = HexColor('#1e3a5f')      # Azul marino corporativo
    PRIMARY_LIGHT = HexColor('#2c5282') # Azul medio
    PRIMARY_DARK = HexColor('#0f1f33')  # Azul oscuro
    SECONDARY = HexColor('#c9a227')    # Dorado/Diesel (acento)
    SUCCESS = HexColor('#059669')      # Verde esmeralda
    WARNING = HexColor('#d97706')      # Ámbar
    DANGER = HexColor('#dc2626')       # Rojo rubí
    INFO = HexColor('#2563eb')         # Azul eléctrico
    TEXT_PRIMARY = HexColor('#1f2937')     # Gris muy oscuro
    TEXT_SECONDARY = HexColor('#4b5563')   # Gris medio
    TEXT_LIGHT = HexColor('#6b7280')       # Gris claro
    BACKGROUND = HexColor('#ffffff')       # Blanco puro
    BACKGROUND_ALT = HexColor('#f8fafc')   # Gris muy claro
    BORDER = HexColor('#e2e8f0')       # Gris borde

# ============================================================================
# FUNCIONES DE SANITIZACIÓN PARA PDF (AGREGAR ESTAS)
# ============================================================================

def sanitizar_texto_para_pdf(texto):
    """
    Limpia el texto para que ReportLab lo renderice correctamente.
    Elimina emojis, normaliza unicode y quita caracteres problemáticos.
    """
    if not texto:
        return ""
    
    # Convertir a string
    texto = str(texto)
    
    # 1. Normalizar unicode (NFD separa tildes de letras)
    texto = unicodedata.normalize('NFKD', texto)
    
    # 2. Eliminar emojis y caracteres de control (excepto \n, \t)
    texto = ''.join(c for c in texto if unicodedata.category(c)[0] != 'C' or c in '\n\t')
    
    # 3. Eliminar caracteres de formato de la IA (markdown)
    texto = re.sub(r'[\*#_`~\[\]\(\)\{\}]', '', texto)  # Markdown básico
    texto = re.sub(r'!\[.*?\]\(.*?\)', '', texto)  # Imágenes markdown
    texto = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', texto)  # Links -> solo texto
    
    # 4. Limpiar espacios múltiples y saltos de línea excesivos
    texto = re.sub(r' +', ' ', texto)
    texto = re.sub(r'\n{3,}', '\n\n', texto)
    texto = re.sub(r'\t+', ' ', texto)
    
    # 5. Eliminar caracteres no imprimibles
    texto = ''.join(c for c in texto if ord(c) >= 32 or c in '\n\t')
    
    # 6. Codificar a latin-1 (ISO-8859-1) que soporta ReportLab con Helvetica
    # Esto elimina caracteres que no se pueden representar
    try:
        texto = texto.encode('latin-1', 'ignore').decode('latin-1')
    except:
        texto = texto.encode('ascii', 'ignore').decode('ascii')
    
    return texto.strip()


def limpiar_para_paragraph(texto, max_length=None):
    """
    Versión segura para usar con Paragraph de ReportLab.
    Opcionalmente trunca a max_length caracteres.
    """
    texto = sanitizar_texto_para_pdf(texto)
    if max_length and len(texto) > max_length:
        texto = texto[:max_length-3] + "..."
    return texto


def crear_parrafo_seguro(texto, estilo, max_length=None):
    """
    Crea un Paragraph de ReportLab con texto sanitizado.
    Maneja errores silenciosamente.
    """
    try:
        texto_limpio = limpiar_para_paragraph(texto, max_length)
        return Paragraph(texto_limpio, estilo)
    except Exception as e:
        # Si falla, intentar con texto ultra-limpio
        try:
            texto_ultra = re.sub(r'[^\w\s\.\,\;\:\-\(\)\n]', '', str(texto))
            return Paragraph(texto_ultra[:500], estilo)
        except:
            # Último recurso: texto de error
            return Paragraph("[Error al renderizar texto]", estilo)


def safe_append_story(story, elemento, fallback=None):
    """
    Agrega elemento a story con manejo de errores.
    """
    try:
        story.append(elemento)
        return True
    except Exception as e:
        print(f"Error agregando elemento a story: {e}")
        if fallback:
            try:
                story.append(fallback)
            except:
                pass
        return False

class Typography:
    """Configuración tipográfica profesional con soporte UTF-8 mejorado"""
    # CAMBIO CLAVE: Helvetica tiene mejor soporte UTF-8 que Times-Roman
    FONT_MAIN = 'Helvetica'
    FONT_BOLD = 'Helvetica-Bold'
    FONT_ITALIC = 'Helvetica-Oblique'
    FONT_BOLD_ITALIC = 'Helvetica-BoldOblique'
    SIZE_TITLE = 16
    SIZE_H1 = 14
    SIZE_H2 = 12
    SIZE_H3 = 11
    SIZE_H4 = 10
    SIZE_BODY = 11
    SIZE_SMALL = 9
    SIZE_FOOTER = 8
    LEADING_TITLE = 20
    LEADING_H1 = 18
    LEADING_H2 = 16
    LEADING_H3 = 14
    LEADING_BODY = 14

def formatear_contenido_plan(contenido, styles):
    """
    Convierte el contenido estructurado de un plan en elementos PDF bien formateados.
    Detecta numerales, viñetas y jerarquía automáticamente.
    """
    import re
    from reportlab.platypus import Paragraph, Spacer
    from reportlab.lib.units import inch
    
    story = []
    lineas = contenido.split('\n')
    i = 0
    
    while i < len(lineas):
        linea = lineas[i].rstrip()
        
        # Saltar líneas vacías al inicio de secciones
        if not linea.strip():
            story.append(Spacer(1, 0.05*inch))
            i += 1
            continue
        
        # Detectar títulos de sección principal (1., 2., 3., etc.)
        if re.match(r'^\d+\.\s+[A-ZÁÉÍÓÚÑ]', linea):
            story.append(Paragraph(linea, styles['Heading2Enhanced']))
            story.append(Spacer(1, 0.08*inch))
            
        # Detectar subtítulos (1.1, 1.2, etc.)
        elif re.match(r'^\d+\.\d+\s+', linea):
            story.append(Paragraph(linea, styles['Heading3Enhanced']))
            story.append(Spacer(1, 0.05*inch))
            
        # Detectar letras de estrategia (A., B., C., D.)
        elif re.match(r'^[A-D]\.\s+', linea.strip()):
            story.append(Paragraph(linea.strip(), styles['Heading3Enhanced']))
            
        # Detectar viñetas (•, -, *)
        elif linea.strip().startswith('•') or linea.strip().startswith('-') or linea.strip().startswith('*'):
            viñeta_texto = linea.strip()[1:].strip()
            
            # Buscar continuación de la viñeta (líneas siguientes indentadas)
            j = i + 1
            while j < len(lineas):
                siguiente = lineas[j]
                # Si la siguiente línea está indentada o es sub-item, es continuación
                if siguiente.startswith('       -') or siguiente.startswith('        -'):
                    viñeta_texto += "<br/>" + "&nbsp;" * 6 + "• " + siguiente.replace('       -', '').replace('        -', '').strip()
                    j += 1
                elif siguiente.startswith('    ') and len(siguiente.strip()) > 0:
                    viñeta_texto += " " + siguiente.strip()
                    j += 1
                else:
                    break
            
            story.append(Paragraph("• " + viñeta_texto, styles['APA_List']))
            i = j - 1  # Saltar líneas procesadas
            
        # Detectar sub-items indentados con guión
        elif linea.startswith('       -') or linea.startswith('        -'):
            texto = linea.replace('       -', '').replace('        -', '').strip()
            story.append(Paragraph("&nbsp;" * 6 + "• " + texto, styles['APA_List']))
            
        # Texto normal (párrafo)
        elif len(linea.strip()) > 5:
            story.append(Paragraph(linea.strip(), styles['BodyTextEnhanced']))
            
        i += 1
    
    return story

def get_ia_client():
    api_key = st.secrets.get("GEMINI_API_KEY")
    if not api_key:
        st.error("❌ No se encontró la API Key de Gemini en st.secrets")
        st.stop()
    genai.configure(api_key=api_key)
    return True

def generar_analisis_ia(tipo_matriz, datos_contexto):
    if not get_ia_client():
        return "Error: No se encontro la API Key de Gemini en st.secrets."
    prompt = f"Actua como un consultor senior de estrategia. Analiza la siguiente matriz {tipo_matriz} y proporciona conclusiones estrategicas clave, riesgos y recomendaciones. Datos: {datos_contexto}"
    return generar_analisis(prompt)


def generar_analisis(prompt, client=None):
    """
    Genera analisis con IA y sanitiza el texto para PDF seguro.
    """
    errores = []
    
    # Prompt mejorado que EXIGE estructura jerarquica clara Y texto plano
    prompt_estructurado = prompt + """

ESTRUCTURA OBLIGATORIA DEL ANALISIS:

Usa el siguiente formato EXACTO para organizar el contenido:

1. TITULO PRINCIPAL
   [Escribe aqui el titulo general del analisis]

2. RESUMEN EJECUTIVO
   [1-2 parrafos con los hallazgos mas importantes]

3. ANALISIS DETALLADO
   3.1 [Subtitulo especifico 1]
       - Punto clave 1: [Descripcion detallada]
       - Punto clave 2: [Descripcion detallada]
       
   3.2 [Subtitulo especifico 2]
       - Punto clave 1: [Descripcion detallada]
       - Punto clave 2: [Descripcion detallada]

4. CONCLUSIONES
   - Conclusion 1
   - Conclusion 2
   - Conclusion 3

5. RECOMENDACIONES
   5.1 Recomendacion prioritaria: [Descripcion]
   5.2 Recomendacion secundaria: [Descripcion]

REGLAS DE FORMATO ESTRICTAS:
- Usa NUMERALES (1., 2., 3.) para titulos principales
- Usa NUMERALES CON PUNTO (3.1, 3.2) para subtitulos
- Usa GUIONES MEDIOS (-) para listas de puntos clave. PROHIBIDO usar viñetas (•)
- Deja LINEAS EN BLANCO entre secciones
- NO uses Markdown (*, #, **, __)
- NO uses emojis ni caracteres especiales
- NO escribas todo seguido; respeta los saltos de linea entre parrafos
- Usa SOLO letras, numeros, puntos, comas, dos puntos, punto y coma, parentesis y guiones
"""
    
    try:
        modelos_disponibles = [m.name for m in genai.list_models() if 'generateContent' in m.supported_generation_methods]
        modelos_a_probar = []
        modelos_a_probar.extend([m for m in modelos_disponibles if 'flash' in m.lower()])
        modelos_a_probar.extend([m for m in modelos_disponibles if 'pro' in m.lower() and m not in modelos_a_probar])
        
        if not modelos_a_probar:
            return "No se encontraron modelos de Gemini disponibles."
            
        for nombre_modelo in modelos_a_probar:
            try:
                model = genai.GenerativeModel(
                    model_name=nombre_modelo,
                    system_instruction="""Eres un consultor senior de estrategia empresarial. 
                    
Tu especialidad es redactar informes ejecutivos con ESTRUCTURA JERARQUICA CLARA:
- Titulos numerados (1., 2., 3.)
- Subtitulos (3.1, 3.2)
- Listas con GUIONES MEDIOS (-), NUNCA viñetas (•)
- Saltos de linea entre secciones
- TEXTO PLANO: sin emojis, sin markdown, sin caracteres especiales

NUNCA escribas todo el texto seguido. Siempre organiza el contenido en secciones bien diferenciadas."""
                )
                
                response = model.generate_content(prompt_estructurado)
                texto = response.text
                
                # =========================================================================
                # SANITIZACION AGRESIVA PARA PDF SEGURO
                # =========================================================================
                
                # 1. Eliminar emojis y caracteres de control
                texto = ''.join(c for c in texto if unicodedata.category(c)[0] not in ['C', 'So'] or c in '\n\t')
                
                # 2. Eliminar markdown
                texto = re.sub(r'\*\*', '', texto)      # Negritas
                texto = re.sub(r'__', '', texto)        # Negritas alt
                texto = re.sub(r'\*', '', texto)        # Asteriscos sueltos
                texto = re.sub(r'_{1,2}', '', texto)    # Guiones bajos
                texto = re.sub(r'#{1,6}\s*', '', texto) # Almohadillas
                
                # 3. REEMPLAZAR viñetas problemáticas por guiones seguros
                texto = texto.replace('•', '-')         # Viñeta bullet → guion
                texto = texto.replace('●', '-')         # Circulo negro → guion
                texto = texto.replace('○', '-')         # Circulo blanco → guion
                texto = texto.replace('▪', '-')         # Cuadrado → guion
                texto = texto.replace('■', '-')         # Cuadrado negro → guion
                texto = texto.replace('‣', '-')         # Triangulo → guion
                texto = texto.replace('⁃', '-')         # Guion raro → guion normal
                texto = texto.replace('◦', '-')         # Circulo pequeño → guion
                
                # 4. Normalizar espacios y saltos de linea
                texto = re.sub(r' +', ' ', texto)       # Espacios multiples → uno
                texto = re.sub(r'\n{3,}', '\n\n', texto) # Saltos excesivos → doble
                texto = re.sub(r'\t+', '    ', texto)   # Tabs → 4 espacios
                
                # 5. Limpiar caracteres no latin-1 (para Helvetica/ReportLab)
                texto = texto.encode('latin-1', 'ignore').decode('latin-1')
                
                # 6. Asegurar saltos de linea despues de numerales para mejor legibilidad
                texto = re.sub(r'(\d+\.\s+[^\n]{50,})\s+(?=\d+\.)', r'\1\n\n', texto)
                
                return texto.strip()
                
            except Exception as e:
                errores.append(f"{nombre_modelo}: {str(e)}")
                continue
                
    except Exception as e:
        return f"Error de conexion: {str(e)}"
        
    return f"Error en analisis. Intentados: {', '.join(errores)}"

st.set_page_config(page_title="Estratega Pro UG-UCE", page_icon="♟️", layout="wide", initial_sidebar_state="expanded")
st.markdown("""
    <style>
    .main { background-color: #f5f5f5; }
    .stButton>button { width: 100%; border-radius: 5px; height: 3em; }
    h1, h2, h3 { color: #1f77b4; }
    </style>
    """, unsafe_allow_html=True)

def init_supabase():
    try:
        if "supabase_url" in st.secrets and "supabase_key" in st.secrets:
            url = st.secrets["supabase_url"]
            key = st.secrets["supabase_key"]
            return create_client(url, key)
        return None
    except Exception:
        return None

def get_datos_empresa(empresa_id):
    if supabase and empresa_id:
        try:
            response = supabase.table('empresas').select('*').eq('id', empresa_id).single().execute()
            return response.data
        except Exception as e:
            st.error(f"Error al cargar datos de la empresa: {e}")
    return {}

def get_datos_tabla(tabla, empresa_id, tipo_matriz_filter=None):
    if supabase and empresa_id:
        try:
            query = supabase.table(tabla).select('*').eq('empresa_id', empresa_id)
            if tipo_matriz_filter:
                query = query.eq('tipo_matriz', tipo_matriz_filter)
            response = query.execute()
            return pd.DataFrame(response.data) if response.data else pd.DataFrame()
        except Exception as e:
            st.error(f"Error al cargar datos de la tabla {tabla}: {e}")
    return pd.DataFrame()

def guardar_analisis_db(empresa_id, tipo_analisis, contenido):
    if supabase and empresa_id:
        try:
            supabase.table('empresas').update({f'analisis_{tipo_analisis}': contenido}).eq('id', empresa_id).execute()
            st.success(f"✅ Análisis de {tipo_analisis.upper()} guardado en la nube.")
            st.rerun()
        except Exception as e:
            st.error(f"Error al guardar análisis: {e}")

def get_empresas():
    """
    Obtiene las empresas del usuario actual:
    - Empresas creadas por el usuario
    - Empresas compartidas con el usuario (como lector o editor)
    """
    if supabase and st.session_state.get("user"):
        try:
            user_id = st.session_state.user.id

            # 1. Obtener empresas creadas por el usuario
            response_propias = supabase.table('empresas').select('id, nombre, propietario_id').eq('propietario_id', user_id).execute()
            empresas_propias = response_propias.data if response_propias.data else []

            # 2. Obtener empresas compartidas con el usuario
            response_compartidas = supabase.table('empresas_compartidas').select('empresa_id, permiso').eq('usuario_compartido_id', user_id).execute()
            compartidas = response_compartidas.data if response_compartidas.data else []

            # 3. Obtener detalles de las empresas compartidas
            empresas_compartidas_list = []
            for comp in compartidas:
                try:
                    resp_emp = supabase.table('empresas').select('id, nombre, propietario_id').eq('id', comp['empresa_id']).single().execute()
                    if resp_emp.data:
                        emp_data = resp_emp.data
                        emp_data['permiso_compartido'] = comp['permiso']  # 'lector' o 'editor'
                        empresas_compartidas_list.append(emp_data)
                except:
                    continue

            # 4. Combinar y marcar propiedad
            todas_empresas = []

            # Agregar propias
            for emp in empresas_propias:
                emp['es_propietario'] = True
                emp['permiso'] = 'propietario'
                todas_empresas.append(emp)

            # Agregar compartidas (evitar duplicados si es propietario y compartida)
            ids_existentes = {e['id'] for e in todas_empresas}
            for emp in empresas_compartidas_list:
                if emp['id'] not in ids_existentes:
                    emp['es_propietario'] = False
                    emp['permiso'] = emp.get('permiso_compartido', 'lector')
                    todas_empresas.append(emp)

            return pd.DataFrame(todas_empresas) if todas_empresas else pd.DataFrame(columns=['id', 'nombre', 'es_propietario', 'permiso'])

        except Exception as e:
            st.error(f"Error al cargar empresas: {e}")
    return pd.DataFrame(columns=['id', 'nombre', 'es_propietario', 'permiso'])

def compartir_empresa(empresa_id, email_usuario, permiso='lector'):
    """
    Comparte una empresa con otro usuario.
    """
    if supabase and empresa_id and email_usuario:
        try:
            # Buscar el usuario por email
            resp_usuario = supabase.rpc('get_user_id_by_email', {'email_input': email_usuario}).execute()
            
            if not resp_usuario.data:
                return False, "Usuario no encontrado. Asegúrate de que el email esté registrado."
            
            usuario_compartir_id = resp_usuario.data
            
            # Verificar que no sea el propietario
            resp_empresa = supabase.table('empresas').select('propietario_id').eq('id', empresa_id).single().execute()
            if resp_empresa.data and resp_empresa.data['propietario_id'] == usuario_compartir_id:
                return False, "No puedes compartir con el propietario"

            # Verificar si ya está compartida
            resp_existente = supabase.table('empresas_compartidas').select('*').eq('empresa_id', empresa_id).eq('usuario_compartido_id', usuario_compartir_id).execute()
            if resp_existente.data:
                # Actualizar permiso y guardar email
                supabase.table('empresas_compartidas').update({
                    'permiso': permiso,
                    'email_compartido': email_usuario  # <-- Guardar email
                }).eq('empresa_id', empresa_id).eq('usuario_compartido_id', usuario_compartir_id).execute()
                return True, "Permiso actualizado"

            # Crear nuevo registro con email
            supabase.table('empresas_compartidas').insert({
                'empresa_id': empresa_id,
                'usuario_compartido_id': usuario_compartir_id,
                'email_compartido': email_usuario,  # <-- Guardar email aquí
                'permiso': permiso
            }).execute()

            return True, f"Empresa compartida con {email_usuario} como {permiso}"

        except Exception as e:
            return False, f"Error al compartir: {e}"
    return False, "Datos incompletos"


def eliminar_compartir(empresa_id, usuario_compartido_id):
    """Elimina el acceso compartido a un usuario."""
    if supabase and empresa_id and usuario_compartido_id:
        try:
            supabase.table('empresas_compartidas').delete().eq('empresa_id', empresa_id).eq('usuario_compartido_id', usuario_compartido_id).execute()
            return True, "Acceso eliminado"
        except Exception as e:
            return False, f"Error: {e}"
    return False, "Datos incompletos"


def get_usuarios_compartidos(empresa_id):
    """Obtiene la lista de usuarios con quienes se compartió la empresa."""
    if supabase and empresa_id:
        try:
            resp = supabase.table('empresas_compartidas').select('*').eq('empresa_id', empresa_id).execute()
            
            if resp.data:
                usuarios = []
                for comp in resp.data:
                    try:
                        usuario_id = comp.get('usuario_compartido_id', 'unknown')
                        permiso = comp.get('permiso', 'lector')
                        
                        # PRIORIDAD 1: Usar email_compartido si existe
                        # PRIORIDAD 2: Usar email si existe  
                        # PRIORIDAD 3: Mostrar ID truncado como fallback
                        email = comp.get('email_compartido') or comp.get('email') or f"ID: {str(usuario_id)[:8]}..."
                        
                        usuarios.append({
                            'usuario_id': usuario_id,
                            'email': email,
                            'permiso': permiso
                        })
                    except Exception as e:
                        print(f"Error procesando registro: {e}")
                        continue
                return usuarios
        except Exception as e:
            st.error(f"Error al cargar usuarios compartidos: {e}")
    return []
    
def save_image(uploaded_file):
    """
    Guarda una imagen subida y retorna los bytes listos para guardar en BD.
    """
    if uploaded_file is not None:
        try:
            # Leer los bytes directamente
            bytes_data = uploaded_file.getvalue()
            return bytes_data  # Retornar bytes directamente
        except Exception as e:
            st.error(f"Error al procesar imagen: {e}")
            return None
    return None

def mostrar_imagen_bd(imagen_data, caption="Imagen", width=None):
    """Muestra imagen guardada en base64"""
    if not imagen_data:
        return False
    try:
        import base64
        if isinstance(imagen_data, str):
            img_bytes = base64.b64decode(imagen_data)
            st.image(img_bytes, caption=caption, width=width)
            return True
    except Exception as e:
        pass
    return False
    
def analizar_foda(df_foda):
    if df_foda.empty: 
        return None, None, None, pd.Series(dtype='float64')
    estrategias = {'FO': 'Ofensiva (F+O)', 'FA': 'Defensiva (F+A)', 'DO': 'Adaptativa (D+O)', 'DA': 'Supervivencia (D+A)'}
    puntajes = df_foda.groupby('cuadrante')['impacto'].sum().reindex(estrategias.keys(), fill_value=0)
    puntajes_ordenados = puntajes.sort_values(ascending=False)
    analisis_df = pd.DataFrame({
        'Estrategia': [estrategias[c] for c in puntajes_ordenados.index],
        'Puntaje Total': puntajes_ordenados.values
    }).reset_index(drop=True)
    estrategia_principal = analisis_df.iloc[0]['Estrategia']
    resumen = f"La estrategia principal recomendada es **{analisis_df.iloc[0]['Estrategia']}** ({analisis_df.iloc[0]['Puntaje Total']} puntos), seguida por **{analisis_df.iloc[1]['Estrategia']}** ({analisis_df.iloc[1]['Puntaje Total']} puntos)."
    return analisis_df, resumen, estrategia_principal, puntajes_ordenados

def generar_planes_por_plantilla(estrategia_foda, pest_total, empresa_id=None):
    """
    Genera los 7 planes funcionales usando exclusivamente IA.
    Si falla, devuelve mensaje de error en lugar de plantillas.
    """
    
    # Obtener contexto adicional si hay empresa_id
    contexto_empresa = ""
    if empresa_id:
        empresa = get_datos_empresa(empresa_id)
        if empresa:
            contexto_empresa = f"""
Nombre de la empresa: {empresa.get('nombre', 'La empresa')}
Giro: {empresa.get('giro', 'No especificado')}
Misión: {empresa.get('mision', 'No disponible')[:200]}
"""
    
    # Determinar contexto estratégico
    es_ofensiva = "Ofensiva" in str(estrategia_foda)
    es_adaptativa = "Adaptativa" in str(estrategia_foda)
    contexto_postura = "crecimiento agresivo" if (es_ofensiva or es_adaptativa) else "consolidación defensiva"
    
    # Prompt maestro para IA
    prompt_maestro = f"""Actúa como un consultor senior de estrategia empresarial con 20 años de experiencia.

CONTEXTO ESTRATÉGICO:
{contexto_empresa}
Estrategia FODA principal: {estrategia_foda}
Postura estratégica: {contexto_postura}
Entorno PEST score: {pest_total:.2f} ({'favorable' if pest_total > 2.5 else 'desafiante'})

GENERA LOS 7 PLANES FUNCIONALES ESTRATÉGICOS SIGUIENDO ESTA ESTRUCTURA EXACTA:

=== 1. PLAN ADMINISTRATIVO ===

1.1 FUNDAMENTO ESTRATÉGICO
    [2-3 párrafos explicando por qué este plan es crítico]

1.2 OBJETIVO GENERAL DEL PLAN
    [Objetivo SMART específico]

1.3 OBJETIVOS ESPECÍFICOS
    • Objetivo 1: [Descripción]
    • Objetivo 2: [Descripción]
    • Objetivo 3: [Descripción]

1.4 ESTRATEGIAS DE IMPLEMENTACIÓN
    A. [Nombre estrategia]
       - Descripción detallada
       - Recursos necesarios
       
    B. [Nombre estrategia]
       - Descripción detallada
       - Recursos necesarios

1.5 KPIs Y METAS
    • KPI 1: [Nombre] | Meta: [X] | Frecuencia: [mensual]

1.6 RECURSOS REQUERIDOS
    • Humanos: [Detalle]
    • Financieros: [Presupuesto]

1.7 RESPONSABLES Y GOBIERNO
    • Responsable: [Rol]

1.8 RIESGOS Y MITIGACIÓN
    • Riesgo: [Descripción] | Mitigación: [Acción]

1.9 ALINEACIÓN CON ESTRATEGIA FODA
    [Explicación]

=== 2. PLAN OPERATIVO ===
[Repetir estructura 2.1 a 2.9]

=== 3. PLAN TECNOLÓGICO ===
[Repetir estructura 3.1 a 3.9]

=== 4. PLAN FINANCIERO ===
[Repetir estructura 4.1 a 4.9]

=== 5. PLAN DE MONITOREO Y CONTROL ===
[Repetir estructura 5.1 a 5.9]

=== 6. PLAN DE MEJORA ===
[Repetir estructura 6.1 a 6.9]

=== 7. PLAN DE CONTINGENCIA ===
[Repetir estructura 7.1 a 7.9]

REGLAS:
- USA numeración X.Y para subtítulos
- USA viñetas (•) para listas
- DEJA líneas en blanco entre secciones
- Adapta contenido a: {contexto_postura} y PEST {pest_total:.2f}

Genera los 7 planes completos ahora:"""

    # Llamar a IA
    planes_generados = generar_analisis(prompt_maestro)
    
    return planes_generados

def generar_cuadro_de_mando_ia(estrategias_df):
    if estrategias_df.empty:
        return pd.DataFrame(columns=['Estrategia', 'Perspectiva', 'KPIs', 'Formulas', 'Frecuencia', 'LI', 'LC', 'LS'])
    
    contexto_estrategias = ""
    for _, row in estrategias_df.iterrows():
        contexto_estrategias += f"- Estrategia: {row['estrategia']} (Plan: {row['plan_asignado']})\n"
    
    prompt = f"""Actúa como un experto en Balanced Scorecard. Basado en las siguientes estrategias:
{contexto_estrategias}
Genera una tabla de Cuadro de Mando Integral (CMI) con las siguientes columnas exactas:
1. (estrategia): La estrategia proporcionada.
2. (perspectiva): A qué perspectiva corresponde (Financiera, Cliente, Procesos, Aprendizaje y Control).
3. (KPIs): El indicador clave de desempeño que medirá la estrategia.
4. (Formulas): La fórmula de cálculo del KPI.
5. (Frecuencia): Cada cuánto se mide (Mensual, Trimestral, etc.).
6. (LI): Límite Inferior (Rojo/Crítico).
7. (LC): Límite de Control (Amarillo/Preventivo).
8. (LS): Límite Superior (Verde/Satisfactorio).
Formato de salida: ESTRATEGIA|PERSPECTIVA|KPI|FORMULA|FRECUENCIA|LI|LC|LS
No incluyas encabezados ni texto adicional, solo las líneas de datos separadas por pipe (|)."""
    
    resultado_ia = generar_analisis(prompt)
    cmi_rows = []
    
    for line in resultado_ia.strip().split("\n"):
        partes = line.split("|")
        if len(partes) >= 8:
            cmi_rows.append({
                "Estrategia": partes[0].strip(),
                "Perspectiva": partes[1].strip(),
                "KPIs": partes[2].strip(),
                "Formulas": partes[3].strip(),
                "Frecuencia": partes[4].strip(),
                "LI": partes[5].strip(),
                "LC": partes[6].strip(),
                "LS": partes[7].strip()
            })
    
    df_cmi = pd.DataFrame(cmi_rows)
    
    # VERIFICAR QUE LAS COLUMNAS NECESARIAS EXISTAN
    columnas_requeridas = ['Estrategia', 'Perspectiva', 'KPIs', 'Formulas', 'Frecuencia', 'LI', 'LC', 'LS']
    
    # Si el DataFrame está vacío o faltan columnas, retornar DataFrame vacío con columnas correctas
    if df_cmi.empty:
        return pd.DataFrame(columns=columnas_requeridas)
    
    # Asegurar que todas las columnas existan
    for col in columnas_requeridas:
        if col not in df_cmi.columns:
            df_cmi[col] = ''
    
    # Ordenar columnas
    df_cmi = df_cmi[columnas_requeridas]
    
    # Ordenar por perspectiva solo si la columna existe y tiene datos válidos
    perspectiva_orden = ['Financiera', 'Cliente', 'Procesos', 'Aprendizaje y Control']
    
    # Filtrar solo perspectivas válidas antes de ordenar
    df_cmi['Perspectiva'] = df_cmi['Perspectiva'].apply(
        lambda x: x if x in perspectiva_orden else 'Procesos'
    )
    
    try:
        df_cmi['Perspectiva'] = pd.Categorical(df_cmi['Perspectiva'], categories=perspectiva_orden, ordered=True)
        df_cmi = df_cmi.sort_values(by='Perspectiva').reset_index(drop=True)
    except Exception as e:
        # Si falla el ordenamiento, simplemente retornar el DataFrame sin ordenar
        pass
    
    return df_cmi
# ============================================================================
# FUNCIONES DE GRÁFICOS MEJORADAS (Agregar estas funciones a tu app.py)
# ============================================================================

import matplotlib.pyplot as plt
import numpy as np

def generar_grafico_foda_radar_mejorado(puntajes):
    """Genera gráfico de radar FODA con diseño profesional mejorado"""
    if puntajes is None or puntajes.empty:
        return None
    
    labels = ['Ofensiva\n(FO)', 'Defensiva\n(FA)', 
              'Adaptativa\n(DO)', 'Supervivencia\n(DA)']
    stats = puntajes.reindex(['FO', 'FA', 'DO', 'DA']).fillna(0).values
    
    # Configurar estilo
    plt.style.use('default')
    
    fig, ax = plt.subplots(figsize=(7, 7), subplot_kw=dict(polar=True))
    
    # Ángulos
    angles = np.linspace(0, 2 * np.pi, len(labels), endpoint=False).tolist()
    stats = np.concatenate((stats, [stats[0]]))
    angles += angles[:1]
    
    # Colores profesionales convertidos a RGB
    color_fill = (30/255, 58/255, 95/255)  # PRIMARY #1e3a5f
    color_line = (201/255, 162/255, 39/255)  # SECONDARY #c9a227
    
    # Dibujar área
    ax.fill(angles, stats, color=color_fill, alpha=0.25)
    ax.plot(angles, stats, color=color_line, linewidth=3, marker='o', 
            markersize=8, markerfacecolor='white', markeredgewidth=2)
    
    # Configurar ejes
    ax.set_yticklabels([])
    ax.set_xticks(angles[:-1])
    ax.set_xticklabels(labels, fontsize=11, fontweight='bold', color='#1f2937')
    
    # Título
    ax.set_title("Posicionamiento Estratégico FODA", 
                fontsize=14, fontweight='bold', color='#1e3a5f', pad=20)
    
    # Grid mejorado
    ax.grid(True, linestyle='--', alpha=0.5, color='gray')
    ax.spines['polar'].set_color('#e2e8f0')
    
    plt.tight_layout()
    buf = BytesIO()
    plt.savefig(buf, format='PNG', dpi=200, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    buf.seek(0)
    return buf


def generar_grafico_barras_pest_mejorado(df_pest):
    """Genera gráfico de barras PEST con diseño profesional"""
    if df_pest.empty:
        return None
    
    pest_scores = df_pest.groupby('categoria')['valor_ponderado'].sum().sort_values(ascending=True)
    
    # Colores por categoría
    colores_categoria = {
        'Político': '#dc2626',
        'Económico': '#059669',
        'Social': '#2563eb',
        'Tecnológico': '#c9a227'
    }
    
    colores = [colores_categoria.get(cat, '#1e3a5f') for cat in pest_scores.index]
    
    fig, ax = plt.subplots(figsize=(8, 5))
    
    bars = ax.barh(pest_scores.index, pest_scores.values, 
                   color=colores, edgecolor='white', linewidth=2)
    
    # Valores en las barras
    for i, (bar, val) in enumerate(zip(bars, pest_scores.values)):
        ax.text(val + 0.5, i, f'{val:.1f}', va='center', 
               fontsize=10, fontweight='bold', color='#1f2937')
    
    # Título y etiquetas
    ax.set_title('Análisis PEST - Impacto por Categoría', 
                fontsize=14, fontweight='bold', color='#1e3a5f', pad=15)
    ax.set_xlabel('Puntuación Ponderada', fontsize=11, color='#4b5563')
    
    # Estilo limpio
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.spines['left'].set_color('#e2e8f0')
    ax.spines['bottom'].set_color('#e2e8f0')
    ax.tick_params(colors='#4b5563')
    
    # Grid
    ax.grid(axis='x', alpha=0.3, linestyle='--', color='gray')
    ax.set_axisbelow(True)
    
    plt.tight_layout()
    buf = BytesIO()
    plt.savefig(buf, format='PNG', dpi=200, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    buf.seek(0)
    return buf


def generar_grafico_proyeccion_mejorado(df_proy):
    """Genera gráfico de proyección financiera con diseño profesional"""
    if df_proy.empty:
        return None
    
    fig, ax = plt.subplots(figsize=(10, 6))
    
    # Colores profesionales
    color_ingresos = '#059669'  # Verde éxito
    color_costos = '#dc2626'    # Rojo peligro
    color_utilidad = '#1e3a5f'  # Azul primario
    
    ax.plot(df_proy['anio'], df_proy['ingresos_proyectados'], 
            marker='o', linewidth=2.5, label='Ingresos Proyectados', 
            color=color_ingresos, markersize=8)
    ax.plot(df_proy['anio'], df_proy['costos_proyectados'], 
            marker='s', linewidth=2.5, label='Costos Proyectados', 
            color=color_costos, markersize=8)
    ax.plot(df_proy['anio'], df_proy['utilidad_neta_proyectada'], 
            marker='^', linewidth=3, label='Utilidad Neta', 
            color=color_utilidad, markersize=10, markerfacecolor='white',
            markeredgewidth=2)
    
    ax.set_title('Proyección Financiera a 5 Años', 
                fontsize=14, fontweight='bold', color='#1e3a5f', pad=15)
    ax.set_xlabel('Año', fontsize=11, color='#4b5563')
    ax.set_ylabel('Monto ($)', fontsize=11, color='#4b5563')
    ax.legend(loc='best', framealpha=0.9, fontsize=10)
    ax.grid(True, alpha=0.3, linestyle='--', color='gray')
    ax.set_axisbelow(True)
    
    # Formato de ejes
    ax.yaxis.set_major_formatter(plt.FuncFormatter(
        lambda x, p: f'${x/1e6:.1f}M' if x >= 1e6 else f'${x/1e3:.0f}K'
    ))
    
    # Estilo limpio
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.spines['left'].set_color('#e2e8f0')
    ax.spines['bottom'].set_color('#e2e8f0')
    ax.tick_params(colors='#4b5563')
    
    plt.tight_layout()
    buf = BytesIO()
    plt.savefig(buf, format='PNG', dpi=200, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    buf.seek(0)
    return buf
    
def get_enhanced_styles():
    """Genera estilos de párrafo profesionales mejorados"""
    styles = getSampleStyleSheet()
    
    # =========================================================================
    # Función auxiliar para agregar estilo solo si no existe
    # =========================================================================
    def add_style_safe(name, parent=None, **kwargs):
        if name in styles:
            # Si ya existe, lo eliminamos primero
            del styles[name]
        
        if parent is not None:
            if isinstance(parent, str):
                parent = styles[parent]
            styles.add(ParagraphStyle(name=name, parent=parent, **kwargs))
        else:
            styles.add(ParagraphStyle(name=name, **kwargs))
    
    # =========================================================================
    # 1. Estilo base
    # =========================================================================
    add_style_safe(
        'EnhancedBase',
        fontName=Typography.FONT_MAIN,
        fontSize=Typography.SIZE_BODY,
        leading=Typography.LEADING_BODY,
        textColor=ColorPalette.TEXT_PRIMARY,
        alignment=TA_JUSTIFY,
        spaceAfter=6,
    )
    
    # =========================================================================
    # 2. Títulos (dependen de EnhancedBase)
    # =========================================================================
    add_style_safe(
        'Heading1Enhanced',
        parent='EnhancedBase',
        fontName=Typography.FONT_BOLD,
        fontSize=Typography.SIZE_H1,
        leading=Typography.LEADING_H1,
        textColor=ColorPalette.PRIMARY,
        alignment=TA_LEFT,
        spaceBefore=24,
        spaceAfter=12,
    )
    
    add_style_safe(
        'Heading2Enhanced',
        parent='EnhancedBase',
        fontName=Typography.FONT_BOLD,
        fontSize=Typography.SIZE_H2,
        leading=Typography.LEADING_H2,
        textColor=ColorPalette.PRIMARY_DARK,
        alignment=TA_LEFT,
        spaceBefore=18,
        spaceAfter=8,
    )
    
    add_style_safe(
        'Heading3Enhanced',
        parent='EnhancedBase',
        fontName=Typography.FONT_BOLD_ITALIC,
        fontSize=Typography.SIZE_H3,
        leading=Typography.LEADING_H3,
        textColor=ColorPalette.TEXT_PRIMARY,
        alignment=TA_LEFT,
        spaceBefore=12,
        spaceAfter=6,
        leftIndent=0.25*inch,
    )
    
    # =========================================================================
    # 3. Texto body (depende de EnhancedBase)
    # =========================================================================
    add_style_safe(
        'BodyTextEnhanced',
        parent='EnhancedBase',
        fontName=Typography.FONT_MAIN,
        fontSize=Typography.SIZE_BODY,
        leading=Typography.LEADING_BODY,
        alignment=TA_JUSTIFY,
        spaceAfter=8,
        firstLineIndent=0.5*inch,
    )
    
    # =========================================================================
    # 4. Variantes de body (dependen de BodyTextEnhanced)
    # =========================================================================
    add_style_safe(
        'BodyTextNoIndent',
        parent='BodyTextEnhanced',
        firstLineIndent=0,
    )
    
    add_style_safe(
        'BodyTextHighlight',
        parent='BodyTextEnhanced',
        backColor=ColorPalette.BACKGROUND_ALT,
        borderPadding=8,
        borderWidth=1,
        borderColor=ColorPalette.BORDER,
        leftIndent=0.25*inch,
        rightIndent=0.25*inch,
        firstLineIndent=0,
    )
    
    # =========================================================================
    # 5. Otros estilos
    # =========================================================================
    add_style_safe(
        'CaptionEnhanced',
        parent='EnhancedBase',
        fontName=Typography.FONT_ITALIC,
        fontSize=Typography.SIZE_SMALL,
        leading=11,
        textColor=ColorPalette.TEXT_SECONDARY,
        alignment=TA_CENTER,
        spaceBefore=6,
        spaceAfter=12,
    )
    
    add_style_safe(
        'TableText',
        parent='EnhancedBase',
        fontSize=9,
        leading=11,
        alignment=TA_LEFT,
        spaceAfter=0,
        firstLineIndent=0,
    )
    
    add_style_safe(
        'TableHeader',
        parent='TableText',
        fontName=Typography.FONT_BOLD,
        textColor=white,
        alignment=TA_CENTER,
    )
    
    # Lista con viñetas (compatible con código antiguo)
    add_style_safe(
        'APA_List',
        parent='BodyTextEnhanced',
        leftIndent=0.5*inch,
        firstLineIndent=-0.25*inch,
        spaceAfter=4,
    )
    
    # =========================================================================
    # 6. ALIAS para compatibilidad con código antiguo (opcional pero recomendado)
    # =========================================================================
    alias_map = {
        'APA_H1': 'Heading1Enhanced',
        'APA_H2': 'Heading2Enhanced',
        'APA_H3': 'Heading3Enhanced',
        'APA_Title': 'Heading1Enhanced',
        'APA_Body': 'BodyTextEnhanced',
        'APA_Body_No_Indent': 'BodyTextNoIndent',
        'APA_Footer': 'CaptionEnhanced',
    }
    
    for old_name, new_name in alias_map.items():
        add_style_safe(old_name, parent=new_name)
    
    return styles



def create_header_footer(canvas, doc, logo_bytes=None, empresa_nombre="", 
                         version="", elaborado="", revisado="", aprobado="", fecha=""):
    """Dibuja encabezado y pie de página profesional"""
    canvas.saveState()
    width, height = A4
    margin = inch
    
    # Header
    header_bottom = height - 1.2*inch
    
    # Barra de color
    canvas.setFillColor(ColorPalette.PRIMARY)
    canvas.rect(0, header_bottom, width, 0.7*inch, fill=1, stroke=0)
    
    # Línea dorada
    canvas.setStrokeColor(ColorPalette.SECONDARY)
    canvas.setLineWidth(2)
    canvas.line(margin, header_bottom, width - margin, header_bottom)
    
    # Logo
    if logo_bytes:
        try:
            logo = Image(logo_bytes, width=0.6*inch, height=0.6*inch)
            logo.drawOn(canvas, margin, header_bottom + 0.05*inch)
        except:
            pass
    
    # Textos header
    canvas.setFillColor(white)
    canvas.setFont(Typography.FONT_BOLD, 12)
    canvas.drawCentredString(width/2, header_bottom + 0.35*inch, 
                            empresa_nombre[:50].upper())
    
    canvas.setFont(Typography.FONT_MAIN, 9)
    canvas.drawRightString(width - margin, header_bottom + 0.35*inch, 
                          f"Versión {version}")
    
    canvas.setFont(Typography.FONT_MAIN, 8)
    canvas.drawRightString(width - margin, header_bottom + 0.15*inch, fecha)
    
    # Footer
    footer_top = 0.8*inch
    
    canvas.setStrokeColor(ColorPalette.BORDER)
    canvas.setLineWidth(0.5)
    canvas.line(margin, footer_top, width - margin, footer_top)
    
    canvas.setFillColor(ColorPalette.TEXT_SECONDARY)
    canvas.setFont(Typography.FONT_MAIN, 8)
    
    canvas.drawString(margin, footer_top - 0.15*inch, f"Elaborado: {elaborado}")
    canvas.drawCentredString(width/2, footer_top - 0.15*inch, f"Revisado: {revisado}")
    canvas.drawRightString(width - margin, footer_top - 0.15*inch, f"Aprobado: {aprobado}")
    
    # Número de página
    canvas.setFillColor(ColorPalette.PRIMARY)
    canvas.setFont(Typography.FONT_BOLD, 10)
    canvas.drawCentredString(width/2, 0.35*inch, f"Página {doc.page}")
    
    # Línea decorativa
    canvas.setStrokeColor(ColorPalette.SECONDARY)
    canvas.setLineWidth(1.5)
    canvas.line(width/2 - 0.5*inch, 0.2*inch, width/2 + 0.5*inch, 0.2*inch)
    
    canvas.restoreState()

def generar_grafico_foda_radar_mejorado(puntajes):
    """Genera gráfico de radar FODA para el PDF."""
    if puntajes is None or puntajes.empty: 
        return None
    
    labels = np.array(['Ofensiva\n(FO)', 'Defensiva\n(FA)', 'Adaptativa\n(DO)', 'Supervivencia\n(DA)'])
    stats = puntajes.reindex(['FO', 'FA', 'DO', 'DA']).fillna(0).values
    angles = np.linspace(0, 2 * np.pi, len(labels), endpoint=False).tolist()
    stats = np.concatenate((stats, [stats[0]]))
    angles += angles[:1]
    
    fig, ax = plt.subplots(figsize=(6, 6), subplot_kw=dict(polar=True))
    ax.fill(angles, stats, color='#1f77b4', alpha=0.25)
    ax.plot(angles, stats, color='#1f77b4', linewidth=2, marker='o')
    ax.set_yticklabels([])
    ax.set_xticks(angles[:-1])
    ax.set_xticklabels(labels, fontsize=10)
    ax.set_title("Posicionamiento Estratégico FODA", size=12, color='black', y=1.1, fontweight='bold')
    ax.grid(True, linestyle='--', alpha=0.7)
    
    buf = BytesIO()
    plt.tight_layout()
    plt.savefig(buf, format='PNG', dpi=150, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    buf.seek(0)
    return buf

def generar_grafico_foda_radar_mejorado(puntajes):
    """Genera gráfico de radar FODA con diseño profesional mejorado"""
    if puntajes is None or puntajes.empty:
        return None
    
    labels = ['Ofensiva\n(FO)', 'Defensiva\n(FA)', 
              'Adaptativa\n(DO)', 'Supervivencia\n(DA)']
    stats = puntajes.reindex(['FO', 'FA', 'DO', 'DA']).fillna(0).values
    
    # Configurar estilo
    plt.style.use('seaborn-v0_8-whitegrid')
    
    fig, ax = plt.subplots(figsize=(7, 7), subplot_kw=dict(polar=True))
    
    # Ángulos
    angles = np.linspace(0, 2 * np.pi, len(labels), endpoint=False).tolist()
    stats = np.concatenate((stats, [stats[0]]))
    angles += angles[:1]
    
    # Colores profesionales convertidos a RGB
    color_fill = (30/255, 58/255, 95/255)  # PRIMARY
    color_line = (201/255, 162/255, 39/255)  # SECONDARY
    
    # Dibujar área
    ax.fill(angles, stats, color=color_fill, alpha=0.25)
    ax.plot(angles, stats, color=color_line, linewidth=3, marker='o', 
            markersize=8, markerfacecolor='white', markeredgewidth=2)
    
    # Configurar ejes
    ax.set_yticklabels([])
    ax.set_xticks(angles[:-1])
    ax.set_xticklabels(labels, fontsize=11, fontweight='bold', color='#1f2937')
    
    # Título
    ax.set_title("Posicionamiento Estratégico FODA", 
                fontsize=14, fontweight='bold', color='#1e3a5f', pad=20)
    
    # Grid mejorado
    ax.grid(True, linestyle='--', alpha=0.5, color='gray')
    ax.spines['polar'].set_color('#e2e8f0')
    
    plt.tight_layout()
    buf = BytesIO()
    plt.savefig(buf, format='PNG', dpi=200, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    buf.seek(0)
    return buf


def generar_grafico_barras_pest_mejorado(df_pest):
    """Genera gráfico de barras PEST con diseño profesional"""
    if df_pest.empty:
        return None
    
    pest_scores = df_pest.groupby('categoria')['valor_ponderado'].sum().sort_values(ascending=True)
    
    # Colores por categoría
    colores_categoria = {
        'Político': '#dc2626',
        'Económico': '#059669',
        'Social': '#2563eb',
        'Tecnológico': '#c9a227'
    }
    
    colores = [colores_categoria.get(cat, '#1e3a5f') for cat in pest_scores.index]
    
    fig, ax = plt.subplots(figsize=(8, 5))
    
    bars = ax.barh(pest_scores.index, pest_scores.values, 
                   color=colores, edgecolor='white', linewidth=2)
    
    # Valores en las barras
    for i, (bar, val) in enumerate(zip(bars, pest_scores.values)):
        ax.text(val + 0.5, i, f'{val:.1f}', va='center', 
               fontsize=10, fontweight='bold', color='#1f2937')
    
    # Título y etiquetas
    ax.set_title('Análisis PEST - Impacto por Categoría', 
                fontsize=14, fontweight='bold', color='#1e3a5f', pad=15)
    ax.set_xlabel('Puntuación Ponderada', fontsize=11, color='#4b5563')
    
    # Estilo limpio
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.spines['left'].set_color('#e2e8f0')
    ax.spines['bottom'].set_color('#e2e8f0')
    ax.tick_params(colors='#4b5563')
    
    # Grid
    ax.grid(axis='x', alpha=0.3, linestyle='--', color='gray')
    ax.set_axisbelow(True)
    
    plt.tight_layout()
    buf = BytesIO()
    plt.savefig(buf, format='PNG', dpi=200, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    buf.seek(0)
    return buf


def procesar_imagen_mejorada(imagen_data):
    """Procesa datos de imagen de la base de datos"""
    if not imagen_data:
        return None
    try:
        if isinstance(imagen_data, bytes):
            return BytesIO(imagen_data)
        elif isinstance(imagen_data, str):
            hex_clean = imagen_data.replace('\\x', '').replace('0x', '').replace("'", "").strip()
            try:
                image_bytes = bytes.fromhex(hex_clean)
            except ValueError:
                import base64
                image_bytes = base64.b64decode(imagen_data)
            
            from PIL import Image as PILImage
            test_img = PILImage.open(BytesIO(image_bytes))
            test_img.verify()
            return BytesIO(image_bytes)
    except Exception as e:
        print(f"Error procesando imagen: {e}")
        return None
    return None

def create_professional_table(data, col_widths=None, has_header=True):
    """Crea tabla con diseño profesional moderno"""
    if col_widths is None:
        num_cols = len(data[0]) if data else 1
        col_widths = [inch * 1.5] * num_cols
    
    table = Table(data, colWidths=col_widths, repeatRows=1 if has_header else 0)
    
    style_commands = [
        ('FONTNAME', (0, 0), (-1, -1), Typography.FONT_MAIN),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('TEXTCOLOR', (0, 0), (-1, -1), ColorPalette.TEXT_PRIMARY),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('LEFTPADDING', (0, 0), (-1, -1), 8),
        ('RIGHTPADDING', (0, 0), (-1, -1), 8),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('LINEBELOW', (0, 0), (-1, -2), 0.5, ColorPalette.BORDER),
        ('LINEABOVE', (0, 0), (-1, 0), 1, ColorPalette.PRIMARY),
        ('LINEBELOW', (0, -1), (-1, -1), 1, ColorPalette.PRIMARY),
        ('GRID', (0, 0), (-1, -1), 0.25, ColorPalette.BORDER),
    ]
    
    if has_header:
        style_commands.extend([
            ('BACKGROUND', (0, 0), (-1, 0), ColorPalette.PRIMARY),
            ('TEXTCOLOR', (0, 0), (-1, 0), white),
            ('FONTNAME', (0, 0), (-1, 0), Typography.FONT_BOLD),
            ('FONTSIZE', (0, 0), (-1, 0), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
            ('TOPPADDING', (0, 0), (-1, 0), 8),
        ])
    
    # Zebra striping
    for i in range(1 if has_header else 0, len(data)):
        if i % 2 == 0:
            style_commands.append(
                ('BACKGROUND', (0, i), (-1, i), ColorPalette.BACKGROUND_ALT)
            )
    
    table.setStyle(TableStyle(style_commands))
    return table

# Alias para compatibilidad con código existente
create_table_pdf = create_professional_table

def parsear_texto_a_story(texto, styles):
    """
    Convierte texto estructurado de la IA en elementos de ReportLab.
    Respeta titulos, subtitulos, listas y parrafos.
    """
    import re
    
    if not texto or not str(texto).strip():
        return []
    
    elementos = []
    lineas = str(texto).split('\n')
    i = 0
    
    while i < len(lineas):
        linea = lineas[i].rstrip()
        
        # Saltar lineas vacias
        if not linea.strip():
            i += 1
            continue
        
        linea_stripped = linea.strip()
        
        # DETECTAR TITULO PRINCIPAL (1., 2., 3., etc. sin subpunto)
        # Ejemplo: "1. TITULO PRINCIPAL" o "2. RESUMEN EJECUTIVO"
        if re.match(r'^\d+\.\s+[A-ZÁÉÍÓÚÑ]', linea_stripped) and not re.match(r'^\d+\.\d+', linea_stripped):
            titulo = limpiar_para_paragraph(linea_stripped)
            elementos.append(Paragraph(titulo, styles['Heading2Enhanced']))
            elementos.append(Spacer(1, 0.15*inch))
            i += 1
            continue
        
        # DETECTAR SUBTITULO (3.1, 3.2, 5.1, etc.)
        # Ejemplo: "3.1 Oportunidades Clave"
        if re.match(r'^\d+\.\d+\s+', linea_stripped):
            subtitulo = limpiar_para_paragraph(linea_stripped)
            elementos.append(Paragraph(subtitulo, styles['Heading3Enhanced']))
            elementos.append(Spacer(1, 0.1*inch))
            i += 1
            continue
        
        # DETECTAR ELEMENTO DE LISTA (- al inicio con espacio)
        # Ejemplo: "- Punto clave 1: descripcion"
        if re.match(r'^-\s+', linea_stripped):
            # Recolectar toda la lista consecutiva
            items_lista = []
            while i < len(lineas):
                linea_actual = lineas[i].rstrip().strip()
                if not linea_actual or not re.match(r'^-\s+', linea_actual):
                    break
                # Quitar el guion inicial y limpiar
                item = re.sub(r'^-\s+', '', linea_actual)
                items_lista.append(item)
                i += 1
            
            # Crear elementos de lista con indentacion
            for item in items_lista:
                item_limpio = limpiar_para_paragraph(item, max_length=300)
                if item_limpio:
                    # Usar estilo de lista con sangria
                    elementos.append(Paragraph(f"        - {item_limpio}", styles['APA_List']))
            
            if items_lista:
                elementos.append(Spacer(1, 0.05*inch))
            continue
        
        # DETECTAR PARRAFO NORMAL
        # Recolectar lineas consecutivas hasta encontrar separador
        parrafo_lineas = [linea_stripped]
        i += 1
        
        while i < len(lineas):
            siguiente = lineas[i].rstrip()
            siguiente_stripped = siguiente.strip()
            
            # Condiciones para detener: vacio, titulo, subtitulo, o lista
            if (not siguiente_stripped or 
                re.match(r'^\d+\.\s+[A-ZÁÉÍÓÚÑ]', siguiente_stripped) or
                re.match(r'^\d+\.\d+', siguiente_stripped) or
                re.match(r'^-\s+', siguiente_stripped)):
                break
            
            parrafo_lineas.append(siguiente_stripped)
            i += 1
        
        # Unir lineas del parrafo (wrapping natural del texto)
        parrafo_completo = ' '.join(parrafo_lineas)
        parrafo_limpio = limpiar_para_paragraph(parrafo_completo, max_length=800)
        
        if parrafo_limpio and len(parrafo_limpio) > 5:
            elementos.append(Paragraph(parrafo_limpio, styles['BodyTextEnhanced']))
            elementos.append(Spacer(1, 0.08*inch))
    
    return elementos


def agregar_analisis_estructurado(story, titulo_seccion, contenido_analisis, styles):
    """
    Funcion auxiliar para agregar un analisis completo al story.
    Maneja errores y asegura que siempre haya contenido.
    """
    if not contenido_analisis or not str(contenido_analisis).strip():
        # Si no hay analisis, agregar mensaje placeholder
        story.append(Paragraph(f"{titulo_seccion}: No se ha generado analisis.", styles['BodyTextEnhanced']))
        story.append(Spacer(1, 0.1*inch))
        return
    
    # Agregar titulo de la subseccion
    story.append(Paragraph(titulo_seccion, styles['Heading3Enhanced']))
    story.append(Spacer(1, 0.05*inch))
    
    # Parsear y agregar contenido estructurado
    try:
        elementos = parsear_texto_a_story(contenido_analisis, styles)
        
        if elementos:
            for elem in elementos:
                story.append(elem)
        else:
            # Fallback: si el parseo devuelve vacio, agregar como parrafo simple
            texto_limpio = limpiar_para_paragraph(contenido_analisis, max_length=1000)
            story.append(Paragraph(texto_limpio, styles['BodyTextEnhanced']))
            story.append(Spacer(1, 0.1*inch))
            
    except Exception as e:
        # Si falla el parseo, agregar texto plano
        print(f"Error parseando analisis '{titulo_seccion}': {e}")
        texto_fallback = limpiar_para_paragraph(str(contenido_analisis)[:500])
        story.append(Paragraph(texto_fallback, styles['BodyTextEnhanced']))
        story.append(Spacer(1, 0.1*inch))
    
def generar_pdf_completo_mejorado(empresa_id, version, elaborado, revisado, aprobado):
    """
    Genera el documento PDF completo con formato APA y estructura solicitada:
    1. Resumen Ejecutivo (máx 5 hojas)
    2. Plan Estratégico (máx 30 hojas) con estructura específica
    3. Anexos (ilimitado)
    """
    from io import BytesIO
    from datetime import datetime
    
    empresa = get_datos_empresa(empresa_id)
    if not empresa:
        st.error("No se pueden generar el PDF, no se encontraron datos de la empresa.")
        return None
    
    # Obtener todos los datos necesarios
    df_pest = get_datos_tabla('matrices', empresa_id, tipo_matriz_filter='PEST')
    df_foda = get_datos_tabla('foda_cruzado', empresa_id)
    df_estrategias = get_datos_tabla('estrategias_generadas', empresa_id)
    df_oper = get_datos_tabla('operativizacion', empresa_id)
    df_pg = get_datos_tabla('perdida_ganancia', empresa_id)
    df_proy = get_datos_tabla('proyeccion_financiera', empresa_id)
    df_cb = get_datos_tabla('analisis_costo_beneficio', empresa_id)
    df_made = get_datos_tabla('matriz_marketing', empresa_id, tipo_matriz_filter='MADE')
    df_madi = get_datos_tabla('matriz_marketing', empresa_id, tipo_matriz_filter='MADI')
    
    # Calcular análisis FODA
    analisis_foda_df, resumen_foda, estrategia_principal, puntajes_foda = analizar_foda(df_foda)
    pest_total = df_pest['valor_ponderado'].sum() if not df_pest.empty else 0
    total_costo = pd.to_numeric(df_oper['costo'], errors='coerce').fillna(0).sum() if not df_oper.empty else 0
    
    # Fecha actual
    fecha_actual = datetime.now().strftime('%d/%m/%Y')
    
    # Preparar buffer y documento
    pdf_buffer = BytesIO()
    
    # Configuración de márgenes APA (1 pulgada = 2.54 cm en todos los lados)
    doc = SimpleDocTemplate(
        pdf_buffer, 
        pagesize=A4,
        leftMargin=1*inch,
        rightMargin=1*inch,
        topMargin=1.5*inch,  # Espacio para encabezado
        bottomMargin=1.2*inch,  # Espacio para pie de página
    )

    # Crear Frame para limitar el área de contenido (evita solapamiento con header/footer)
    from reportlab.platypus import Frame
    content_frame = Frame(
        doc.leftMargin,  # x
        doc.bottomMargin,  # y
        doc.width,  # width
        letter[1] - doc.topMargin - doc.bottomMargin,  # height (página - márgenes)
        id='content_frame',
        showBoundary=0  # 0 = no mostrar borde, 1 = debug
    )
    
    styles = get_enhanced_styles()
    story = []
    
    # Preparar logo
    logo_bytes_data = empresa.get('logo')
    logo_bytes = None
    if logo_bytes_data:
        try:
            # Intentar diferentes formatos de decodificación
            image_bytes = None
            
            if isinstance(logo_bytes_data, bytes):
                image_bytes = logo_bytes_data
            elif isinstance(logo_bytes_data, str):
                # Limpiar el string de prefijos comunes
                hex_clean = logo_bytes_data.replace('\\x', '').replace('0x', '').replace("'", "").strip()
                try:
                    image_bytes = bytes.fromhex(hex_clean)
                except ValueError:
                    # Si no es hex válido, intentar como base64
                    try:
                        import base64
                        image_bytes = base64.b64decode(logo_bytes_data)
                    except:
                        pass
            
            # Verificar que sea una imagen válida antes de usarla
            if image_bytes:
                from PIL import Image as PILImage
                from io import BytesIO
                
                try:
                    # Intentar abrir con PIL para validar
                    test_img = PILImage.open(BytesIO(image_bytes))
                    test_img.verify()  # Verificar que no esté corrupta
                    logo_bytes = BytesIO(image_bytes)
                except Exception as e:
                    print(f"Logo no es una imagen válida: {e}")
                    logo_bytes = None
                    
        except Exception as e:
            print(f"Error al procesar logo: {e}")
            logo_bytes = None

    # Preparar organigrama
    organigrama_bytes = None
    org_bytes_data = empresa.get('organigrama')
    if org_bytes_data:
        try:
            image_bytes = None
            
            if isinstance(org_bytes_data, bytes):
                image_bytes = org_bytes_data
            elif isinstance(org_bytes_data, str):
                hex_clean = org_bytes_data.replace('\\x', '').replace('0x', '').replace("'", "").strip()
                try:
                    image_bytes = bytes.fromhex(hex_clean)
                except ValueError:
                    try:
                        import base64
                        image_bytes = base64.b64decode(org_bytes_data)
                    except:
                        pass
            
            if image_bytes:
                from PIL import Image as PILImage
                from io import BytesIO
                
                try:
                    test_img = PILImage.open(BytesIO(image_bytes))
                    test_img.verify()
                    organigrama_bytes = BytesIO(image_bytes)
                except Exception as e:
                    print(f"Organigrama no es una imagen válida: {e}")
                    organigrama_bytes = None
                    
        except Exception as e:
            print(f"Error al procesar organigrama: {e}")
            organigrama_bytes = None

    # Función para encabezado/pie en cada página
    def header_footer(canvas, doc):
        create_header_footer(
            canvas, doc, logo_bytes, 
            empresa.get('nombre', 'Empresa'), 
            version, elaborado, revisado, aprobado, fecha_actual
        )
    
    # ============================================
    # PARTE 1: RESUMEN EJECUTIVO (máx 5 hojas)
    # ============================================
    
    story.append(Paragraph("RESUMEN EJECUTIVO", styles['Heading1Enhanced']))
    story.append(Spacer(1, 0.2*inch))
    
    story.append(Paragraph(
        f"El presente documento constituye el Plan Estratégico de <b>{empresa.get('nombre', 'la empresa')}</b>, "
        f"elaborado con fecha {fecha_actual}. Este resumen ejecutivo presenta los hallazgos más relevantes "
        f"del diagnóstico estratégico y las recomendaciones prioritarias para la alta dirección.",
        styles['BodyTextEnhanced']
    ))
    story.append(Spacer(1, 0.2*inch))
    
    # Datos clave del diagnóstico
    story.append(Paragraph("Diagnóstico Estratégico Clave", styles['Heading2Enhanced']))
    
    if estrategia_principal:
        story.append(Paragraph(
            f"<b>Estrategia Principal Recomendada:</b> {estrategia_principal}. "
            f"Esta postura estratégica se determina a partir del análisis FODA cruzado y representa "
            f"la orientación prioritaria para el período de planificación.",
            styles['BodyTextEnhanced']
        ))
    
    # Gráfico FODA radar (compacto)
    if puntajes_foda is not None and not puntajes_foda.empty:
        story.append(Spacer(1, 0.1*inch))
        grafico_foda = generar_grafico_foda_radar_mejorado(puntajes_foda)
        if grafico_foda:
            story.append(Image(grafico_foda, width=3.5*inch, height=3.5*inch))
            story.append(Paragraph(
                "<i>Figura 1. Posicionamiento estratégico según análisis FODA cruzado.</i>",
                ParagraphStyle(name='Caption', parent=styles['BodyTextEnhanced'], alignment=TA_CENTER, fontSize=10, firstLineIndent=0)
            ))
    
    story.append(Spacer(1, 0.2*inch))
    
    # Análisis PEST resumido
    if not df_pest.empty:
        story.append(Paragraph("Factores Críticos del Entorno", styles['Heading2Enhanced']))
        
        pest_criticos = df_pest.nlargest(3, 'valor_ponderado')
        factores_texto = []
        for _, row in pest_criticos.iterrows():
            factores_texto.append(f"<b>{row['categoria']}</b>: {row['factor'][:100]}...")
        
        for ft in factores_texto:
            story.append(Paragraph(f"• {ft}", styles['APA_List']))
        
        # Gráfico PEST compacto
        grafico_pest = generar_grafico_barras_pest_mejorado(df_pest)
        if grafico_pest:
            story.append(Spacer(1, 0.1*inch))
            story.append(Image(grafico_pest, width=4*inch, height=2.5*inch))
            story.append(Paragraph(
                "<i>Figura 2. Distribución de factores PEST por impacto.</i>",
                ParagraphStyle(name='Caption', parent=styles['BodyTextEnhanced'], alignment=TA_CENTER, fontSize=10, firstLineIndent=0)
            ))
    
    story.append(Spacer(1, 0.2*inch))
    
    # Estrategias priorizadas
    if not df_estrategias.empty:
        story.append(Paragraph("Estrategias Prioritarias", styles['Heading2Enhanced']))
        
        estrategias_alta = df_estrategias[df_estrategias['importancia'].isin(['Alta', 'Media Alta'])].head(5)
        
        for idx, row in estrategias_alta.iterrows():
            story.append(Paragraph(
                f"<b>{row['cuadrante']} - {row['plan_asignado']}</b>: {row['estrategia'][:150]}...",
                styles['APA_List']
            ))
    
    story.append(Spacer(1, 0.2*inch))
    
    # Indicadores financieros clave
    if not df_cb.empty:
        story.append(Paragraph("Viabilidad Financiera", styles['Heading2Enhanced']))
        
        datos_cb = df_cb.iloc[0]
        
        # Tabla resumen financiero
        datos_fin = [
            ['Indicador', 'Valor', 'Interpretación'],
            ['Relación Costo-Beneficio', f"{float(datos_cb['relacion_costo_beneficio_dolares']):.2f}", 
             'Rentable' if datos_cb['relacion_costo_beneficio_dolares'] >= 1 else 'No rentable'],
            ['Periodo de Recuperación', 
             f"{float(datos_cb['payback_periodo_anios']):.1f} años" if datos_cb['payback_periodo_anios'] else 'N/A',
             'Aceptable' if datos_cb['payback_periodo_anios'] and datos_cb['payback_periodo_anios'] <= 5 else 'Revisar'],
            ['Inversión Total', f"${float(datos_cb['inversion_total']):,.2f}", 'Requerimiento de capital'],
        ]
        
        tabla_fin = create_professional_table(datos_fin, col_widths=[2*inch, 1.5*inch, 2*inch])
        story.append(tabla_fin)
        
        # Gráfico de proyección
        if not df_proy.empty:
            story.append(Spacer(1, 0.1*inch))
            grafico_proy = generar_grafico_proyeccion_mejorado(df_proy)
            if grafico_proy:
                story.append(Image(grafico_proy, width=5*inch, height=3*inch))
                story.append(Paragraph(
                    "<i>Figura 3. Proyección financiera del plan estratégico.</i>",
                    ParagraphStyle(name='Caption', parent=styles['BodyTextEnhanced'], alignment=TA_CENTER, fontSize=10, firstLineIndent=0)
                ))
    
    story.append(Spacer(1, 0.2*inch))
    
    # Recomendación final
    story.append(Paragraph("Recomendación Ejecutiva", styles['Heading2Enhanced']))
    
    if not df_cb.empty:
        puntos_positivos = sum([
            datos_cb['relacion_costo_beneficio_dolares'] >= 1,
            datos_cb['payback_periodo_anios'] is not None and datos_cb['payback_periodo_anios'] <= 5,
            datos_cb['relacion_cb_unidades'] >= 1
        ])
        
        if puntos_positivos >= 2:
            recomendacion = (
                "<b>APROBAR EL PROYECTO.</b> El análisis integral indica que el Plan Estratégico es "
                "viable financieramente y alineado con la posición competitiva de la empresa. "
                "Se recomienda iniciar la implementación priorizando las estrategias de alta importancia."
            )
        elif puntos_positivos == 1:
            recomendacion = (
                "<b>EVALUAR CON PRECAUCIÓN.</b> El proyecto presenta riesgos moderados. "
                "Se sugiere revisar los supuestos críticos y desarrollar planes de contingencia "
                "antes de aprobar la inversión total."
            )
        else:
            recomendacion = (
                "<b>RECHAZAR O REPLANTEAR.</b> El análisis revela riesgos significativos. "
                "Se recomienda revisar la estrategia, reducir el alcance inicial o buscar "
                "alternativas de menor inversión."
            )
        
        story.append(Paragraph(recomendacion, styles['BodyTextEnhanced']))
    
    story.append(PageBreak())
    
    # ============================================
    # PARTE 2: PLAN ESTRATÉGICO (máx 30 hojas)
    # ============================================
    
    # PORTADA INTERNA DEL PLAN
    story.append(Spacer(1, 1*inch))
    story.append(Paragraph("PLAN ESTRATÉGICO", styles['APA_Title']))
    story.append(Spacer(1, 0.3*inch))
    story.append(Paragraph(empresa.get('nombre', 'EMPRESA').upper(), styles['APA_Title']))
    story.append(Spacer(1, 0.5*inch))
    
    if logo_bytes:
        story.append(Image(logo_bytes, width=2*inch, height=2*inch))
        story.append(Spacer(1, 0.5*inch))
    
    story.append(Paragraph(f"<b>Versión:</b> {version}", styles['BodyTextNoIndent']))
    story.append(Paragraph(f"<b>Fecha:</b> {fecha_actual}", styles['BodyTextNoIndent']))
    story.append(Spacer(1, 1*inch))
    story.append(Paragraph(f"<b>Elaborado por:</b> {elaborado}", styles['BodyTextNoIndent']))
    story.append(Paragraph(f"<b>Revisado por:</b> {revisado}", styles['BodyTextNoIndent']))
    story.append(Paragraph(f"<b>Aprobado por:</b> {aprobado}", styles['BodyTextNoIndent']))
    
    story.append(PageBreak())
    
    # 1. INTRODUCCIÓN Y FUNDAMENTOS (Datos Generales + Cultura Organizacional)
    story.append(Paragraph("1. INTRODUCCIÓN Y FUNDAMENTOS", styles['Heading1Enhanced']))
    
    # 1.1 Datos Generales
    story.append(Paragraph("1.1 Datos Generales de la Empresa", styles['Heading2Enhanced']))
    
    datos_gen = [
        ['Campo', 'Información'],
        ['Nombre de la Empresa', empresa.get('nombre', 'N/A')],
        ['Giro del Negocio', empresa.get('giro', 'N/A')],
        ['Fecha del Plan', fecha_actual],
        ['Versión del Documento', version],
    ]
    
    tabla_datos_gen = create_table_pdf(datos_gen, col_widths=[2*inch, 4*inch])
    story.append(tabla_datos_gen)
    story.append(Spacer(1, 0.2*inch))
    
    # Logo de la empresa
    if logo_bytes:
        story.append(Paragraph("Logo de la Empresa", styles['Heading3Enhanced']))
        story.append(Image(logo_bytes, width=2*inch, height=2*inch))
        story.append(Spacer(1, 0.2*inch))
    
    # 1.2 Cultura Organizacional (Misión, Visión, Valores)
    story.append(Paragraph("1.2 Elementos Orientadores de la Cultura Organizacional", styles['Heading2Enhanced']))
    
    if empresa.get('mision'):
        story.append(Paragraph("Misión", styles['Heading3Enhanced']))
        story.append(Paragraph(empresa['mision'], styles['BodyTextEnhanced']))
        story.append(Spacer(1, 0.1*inch))
    
    if empresa.get('vision'):
        story.append(Paragraph("Visión", styles['Heading3Enhanced']))
        story.append(Paragraph(empresa['vision'], styles['BodyTextEnhanced']))
        story.append(Spacer(1, 0.1*inch))
    
    if empresa.get('valores'):
        story.append(Paragraph("Valores y Principios", styles['Heading3Enhanced']))
        valores = empresa['valores']
        if len(valores) > 200:
            lineas_valores = [v.strip() for v in valores.replace('\n', ',').split(',') if v.strip()]
            for val in lineas_valores[:5]:
                story.append(Paragraph(f"• {val}", styles['APA_List']))
        else:
            story.append(Paragraph(valores, styles['BodyTextEnhanced']))
    
    if empresa.get('objetivo_plan'):
        story.append(Paragraph("Objetivo del Plan Estratégico", styles['Heading3Enhanced']))
        story.append(Paragraph(empresa['objetivo_plan'], styles['BodyTextEnhanced']))
    
    # Organigrama
    if organigrama_bytes:
        try:
            story.append(Paragraph("1.3 Organigrama de la Empresa", styles['Heading2Enhanced']))
            story.append(Image(organigrama_bytes, width=6*inch, height=4*inch))
            story.append(Paragraph(
                "<i>Figura 1.1. Estructura organizacional de la empresa.</i>",
                ParagraphStyle(name='Caption', parent=styles['BodyTextEnhanced'], alignment=TA_CENTER, fontSize=10, firstLineIndent=0)
            ))
        except Exception as e:
            story.append(Paragraph("1.3 Organigrama de la Empresa", styles['Heading2Enhanced']))
            story.append(Paragraph(
                "[El organigrama no pudo ser cargado - formato de imagen no válido]",
                ParagraphStyle(name='Error', parent=styles['BodyTextEnhanced'], textColor='red', alignment=TA_CENTER)
            ))
    
    story.append(PageBreak())
    
    # =========================================================================
    # SECCION 2: ANALISIS SITUACIONAL (VERSION CORREGIDA)
    # =========================================================================
    
    story.append(Paragraph("2. ANALISIS SITUACIONAL", styles['Heading1Enhanced']))
    story.append(Paragraph(
        "El analisis situacional examina tanto los factores internos como externos que afectan "
        "a la organizacion, permitiendo identificar fortalezas, debilidades, oportunidades y amenazas.",
        styles['BodyTextEnhanced']
    ))
    story.append(Spacer(1, 0.2*inch))
    
    # 2.1 Diagnostico Interno
    story.append(Paragraph("2.1 Diagnostico Interno", styles['Heading2Enhanced']))
    
    # Analisis MADE (Marketing Interno)
    if not df_made.empty:
        story.append(Paragraph("2.1.1 Analisis de Marketing Interno (MADE)", styles['Heading3Enhanced']))
        story.append(Paragraph(
            "La matriz MADE evalua las variables internas de marketing: Producto, Precio, Plaza y Promocion.",
            styles['BodyTextEnhanced']
        ))
        story.append(Spacer(1, 0.1*inch))
        
        # Tabla resumen MADE
        datos_made_resumen = [['Variable', 'Factor', 'Rating', 'Ponderacion']]
        for _, row in df_made.head(10).iterrows():
            datos_made_resumen.append([
                limpiar_para_paragraph(str(row.get('variable', '')))[:20],
                limpiar_para_paragraph(str(row.get('factor', '')))[:30],
                str(row.get('rating', '')),
                f"{row.get('weight_percent', '')}%"
            ])
        
        tabla_made = create_professional_table(datos_made_resumen, col_widths=[1.2*inch, 2.5*inch, 0.8*inch, 1*inch])
        story.append(tabla_made)
        story.append(Spacer(1, 0.15*inch))
        
        # Analisis de IA con formato estructurado
        analisis_made = empresa.get('analisis_made', '')
        agregar_analisis_estructurado(
            story, 
            "Analisis de Marketing Interno", 
            analisis_made, 
            styles
        )
    
    story.append(Spacer(1, 0.2*inch))
    
    # 2.2 Diagnostico Externo
    story.append(Paragraph("2.2 Diagnostico Externo", styles['Heading2Enhanced']))
    
    # Analisis MADI (Marketing Externo)
    if not df_madi.empty:
        story.append(Paragraph("2.2.1 Analisis de Marketing Externo (MADI)", styles['Heading3Enhanced']))
        story.append(Paragraph(
            "La matriz MADI evalua las variables externas de marketing que impactan en la posicion competitiva.",
            styles['BodyTextEnhanced']
        ))
        story.append(Spacer(1, 0.1*inch))
        
        datos_madi_resumen = [['Variable', 'Factor', 'Rating', 'Ponderacion']]
        for _, row in df_madi.head(10).iterrows():
            datos_madi_resumen.append([
                limpiar_para_paragraph(str(row.get('variable', '')))[:20],
                limpiar_para_paragraph(str(row.get('factor', '')))[:30],
                str(row.get('rating', '')),
                f"{row.get('weight_percent', '')}%"
            ])
        
        tabla_madi = create_professional_table(datos_madi_resumen, col_widths=[1.2*inch, 2.5*inch, 0.8*inch, 1*inch])
        story.append(tabla_madi)
        story.append(Spacer(1, 0.15*inch))
        
        # Analisis de IA con formato estructurado
        analisis_madi = empresa.get('analisis_madi', '')
        agregar_analisis_estructurado(
            story, 
            "Analisis de Marketing Externo", 
            analisis_madi, 
            styles
        )
    
    story.append(Spacer(1, 0.2*inch))
    
    # Analisis PEST
    if not df_pest.empty:
        story.append(Paragraph("2.2.2 Analisis del Entorno PEST", styles['Heading3Enhanced']))
        story.append(Paragraph(
            "El analisis PEST examina los factores Politicos, Economicos, Sociales y Tecnologicos.",
            styles['BodyTextEnhanced']
        ))
        story.append(Spacer(1, 0.1*inch))
        
        datos_pest_resumen = [['Categoria', 'Factor', 'Tipo FODA', 'Puntaje', 'Ponderado']]
        for _, row in df_pest.head(15).iterrows():
            datos_pest_resumen.append([
                limpiar_para_paragraph(str(row.get('categoria', ''))),
                limpiar_para_paragraph(str(row.get('factor', '')))[:35] + '...' if len(str(row.get('factor', ''))) > 35 else limpiar_para_paragraph(str(row.get('factor', ''))),
                limpiar_para_paragraph(str(row.get('tipo_foda', ''))),
                str(row.get('puntaje', '')),
                f"{row.get('valor_ponderado', 0):.2f}"
            ])
        
        tabla_pest = create_professional_table(datos_pest_resumen, col_widths=[1*inch, 2.5*inch, 1*inch, 0.8*inch, 1*inch])
        story.append(tabla_pest)
        story.append(Spacer(1, 0.15*inch))
        
        # Grafico PEST
        grafico_pest = generar_grafico_barras_pest_mejorado(df_pest)
        if grafico_pest:
            story.append(Image(grafico_pest, width=5*inch, height=3*inch))
            story.append(Paragraph(
                "<i>Figura 2.1. Analisis PEST - Distribucion por categoria.</i>",
                ParagraphStyle(name='Caption', parent=styles['BodyTextEnhanced'], alignment=TA_CENTER, fontSize=10, firstLineIndent=0)
            ))
            story.append(Spacer(1, 0.1*inch))
        
        # Analisis de IA con formato estructurado
        analisis_pest = empresa.get('analisis_pest', '')
        agregar_analisis_estructurado(
            story, 
            "Interpretacion del Analisis PEST", 
            analisis_pest, 
            styles
        )
    
    story.append(Spacer(1, 0.2*inch))
    
    # Matriz FODA Cruzado
    if not df_foda.empty:
        story.append(Paragraph("2.3 Matriz FODA Cruzado", styles['Heading2Enhanced']))
        story.append(Paragraph(
            "El analisis FODA cruzado identifica estrategias a partir de la combinacion de "
            "fortalezas, debilidades, oportunidades y amenazas.",
            styles['BodyTextEnhanced']
        ))
        story.append(Spacer(1, 0.1*inch))
        
        datos_foda_resumen = [['Cuadrante', 'Factor Fila', 'Factor Columna', 'Impacto']]
        for _, row in df_foda.head(20).iterrows():
            datos_foda_resumen.append([
                limpiar_para_paragraph(str(row.get('cuadrante', ''))),
                limpiar_para_paragraph(str(row.get('factor_fila', '')))[:30] + '...' if len(str(row.get('factor_fila', ''))) > 30 else limpiar_para_paragraph(str(row.get('factor_fila', ''))),
                limpiar_para_paragraph(str(row.get('factor_columna', '')))[:30] + '...' if len(str(row.get('factor_columna', ''))) > 30 else limpiar_para_paragraph(str(row.get('factor_columna', ''))),
                str(row.get('impacto', ''))
            ])
        
        tabla_foda = create_professional_table(datos_foda_resumen, col_widths=[1*inch, 2.5*inch, 2.5*inch, 0.8*inch])
        story.append(tabla_foda)
        story.append(Spacer(1, 0.15*inch))
        
        # Grafico FODA
        if puntajes_foda is not None and not puntajes_foda.empty:
            grafico_foda = generar_grafico_foda_radar_mejorado(puntajes_foda)
            if grafico_foda:
                story.append(Image(grafico_foda, width=4*inch, height=4*inch))
                story.append(Paragraph(
                    "<i>Figura 2.2. Posicionamiento estrategico segun FODA cruzado.</i>",
                    ParagraphStyle(name='Caption', parent=styles['BodyTextEnhanced'], alignment=TA_CENTER, fontSize=10, firstLineIndent=0)
                ))
                story.append(Spacer(1, 0.1*inch))
        
        # Postura estrategica
        if analisis_foda_df is not None:
            story.append(Paragraph("Postura Estrategica Recomendada", styles['Heading3Enhanced']))
            story.append(Paragraph(
                f"Basado en el analisis cruzado, la estrategia principal recomendada es "
                f"<b>{estrategia_principal}</b>. La distribucion de puntajes por estrategia es:",
                styles['BodyTextEnhanced']
            ))
            story.append(Spacer(1, 0.1*inch))
            
            datos_postura = [['Estrategia', 'Puntaje Total']]
            for _, row in analisis_foda_df.iterrows():
                datos_postura.append([
                    limpiar_para_paragraph(str(row['Estrategia'])),
                    str(row['Puntaje Total'])
                ])
            
            tabla_postura = create_professional_table(datos_postura, col_widths=[3*inch, 2*inch])
            story.append(tabla_postura)
            story.append(Spacer(1, 0.15*inch))
        
        # Analisis de IA con formato estructurado
        analisis_foda_texto = empresa.get('analisis_foda', '')
        agregar_analisis_estructurado(
            story, 
            "Interpretacion del Analisis FODA", 
            analisis_foda_texto, 
            styles
        )
    
    story.append(PageBreak())
    
    
    # 3. ESTRATEGIAS
    story.append(Paragraph("3. ESTRATEGIAS", styles['Heading1Enhanced']))
    story.append(Paragraph(
        "Las estrategias representan las acciones específicas diseñadas para alcanzar los objetivos "
        "organizacionales, derivadas del análisis situacional.",
        styles['BodyTextEnhanced']
    ))
    story.append(Spacer(1, 0.2*inch))
    
    if not df_estrategias.empty:
        story.append(Paragraph(
            f"Se han formulado {len(df_estrategias)} estrategias distribuidas en los cuatro cuadrantes "
            f"del análisis FODA cruzado. La estrategia principal recomendada es <b>{estrategia_principal}</b>.",
            styles['BodyTextEnhanced']
        ))
        story.append(Spacer(1, 0.2*inch))
        
        story.append(Paragraph("3.1 Resumen de Estrategias por Cuadrante", styles['Heading2Enhanced']))
        
        resumen_est = df_estrategias.groupby('cuadrante').agg({
            'estrategia': 'count',
            'plan_asignado': lambda x: ', '.join(x.unique())
        }).reset_index()
        resumen_est.columns = ['Cuadrante', 'N° Estrategias', 'Planes Asignados']
        
        datos_est_tabla = [['Cuadrante', 'N° Estrategias', 'Planes Asignados']]
        for _, row in resumen_est.iterrows():
            datos_est_tabla.append([row['Cuadrante'], str(row['N° Estrategias']), row['Planes Asignados']])
        
        tabla_est = create_table_pdf(datos_est_tabla, col_widths=[1.5*inch, 1.5*inch, 3*inch])
        story.append(tabla_est)
        story.append(Spacer(1, 0.2*inch))
        
        story.append(Paragraph("3.2 Estrategias Detalladas", styles['Heading2Enhanced']))
        
        for idx, row in df_estrategias.iterrows():
            story.append(Paragraph(f"Estrategia {idx + 1}: [{row['cuadrante']}] {row['estrategia']}", styles['Heading3Enhanced']))
            story.append(Paragraph(f"<b>Plan asignado:</b> {row['plan_asignado']}", styles['BodyTextNoIndent']))
            story.append(Paragraph(f"<b>Importancia:</b> {row['importancia']}", styles['BodyTextNoIndent']))
            story.append(Paragraph(f"<b>Actividades clave:</b> {row['actividades']}", styles['BodyTextEnhanced']))
            story.append(Spacer(1, 0.15*inch))
    else:
        story.append(Paragraph("No se han generado estrategias en el sistema.", styles['BodyTextEnhanced']))
    
    story.append(PageBreak())
    
    # ============================================
    # 4. PLAN DE ACCIÓN - PLANES FUNCIONALES
    # ============================================
    
    story.append(Paragraph("4. PLAN DE ACCIÓN", styles['Heading1Enhanced']))
    story.append(Paragraph(
        "El plan de acción detalla las actividades específicas, responsables, tiempos y recursos "
        "necesarios para implementar las estrategias formuladas.",
        styles['BodyTextEnhanced']
    ))
    story.append(Spacer(1, 0.2*inch))
    
    
    # 4.2 Planes Funcionales Detallados (si existe análisis guardado)
    # Planes Maestros con formato estructurado
    planes_maestros = empresa.get('analisis_operativo', '')
    if planes_maestros and str(planes_maestros).strip():
        story.append(Paragraph("4.1 Planes Funcionales Detallados", styles['Heading2Enhanced']))
        story.append(Spacer(1, 0.1*inch))
        
        # Los planes maestros son muy largos, dividir en secciones si es posible
        try:
            # Intentar dividir por los 7 planes si estan marcados
            secciones_planes = re.split(r'\n(?=\d+\.\s+PLAN\s+)', str(planes_maestros))
            
            if len(secciones_planes) > 1:
                # Hay secciones claramente divididas
                for seccion in secciones_planes:
                    if seccion.strip():
                        elementos_plan = parsear_texto_a_story(seccion, styles)
                        for elem in elementos_plan:
                            story.append(elem)
                        story.append(Spacer(1, 0.1*inch))
            else:
                # No hay division clara, parsear todo junto
                elementos_planes = parsear_texto_a_story(planes_maestros, styles)
                for elem in elementos_planes:
                    story.append(elem)
                    
        except Exception as e:
            # Fallback: texto plano
            print(f"Error parseando planes maestros: {e}")
            texto_planes = limpiar_para_paragraph(planes_maestros, max_length=2000)
            # Dividir en chunks si es muy largo
            chunks = [texto_planes[i:i+500] for i in range(0, len(texto_planes), 500)]
            for chunk in chunks:
                story.append(Paragraph(chunk, styles['BodyTextEnhanced']))
                story.append(Spacer(1, 0.05*inch))
        story.append(PageBreak())
    
    # 4.3 Operativización (Cuadro de Operativización)
    story.append(Paragraph("4.2 Operativización y Presupuesto", styles['Heading2Enhanced']))
    
    if not df_oper.empty:
        story.append(Paragraph(
            f"La operativización detalla {len(df_oper)} actividades derivadas de las estrategias formuladas, "
            f"con una inversión total estimada de <b>${total_costo:,.2f}</b>.",
            styles['BodyTextEnhanced']
        ))
        story.append(Spacer(1, 0.2*inch))
        
        story.append(Paragraph("4.2.1 Resumen de Inversión por Plan", styles['Heading3Enhanced']))
        
        presupuesto_plan = df_oper.groupby('plan_asignado').agg({
            'costo': 'sum',
            'descripcion_actividad': 'count'
        }).reset_index()
        presupuesto_plan.columns = ['Plan', 'Costo Total', 'N° Actividades']
        
        datos_pres = [['Plan Estratégico', 'Costo Total', 'N° Actividades']]
        for _, row in presupuesto_plan.iterrows():
            datos_pres.append([
                row['Plan'],
                f"${row['Costo Total']:,.2f}",
                str(int(row['N° Actividades']))
            ])
        
        tabla_pres = create_table_pdf(datos_pres, col_widths=[3*inch, 1.5*inch, 1.5*inch])
        story.append(tabla_pres)
        story.append(Spacer(1, 0.2*inch))
        
        story.append(Paragraph("4.2.2 Cuadro de Operativización Detallado", styles['Heading3Enhanced']))
        
        datos_oper = [['Estrategia', 'Actividad', 'Plazo', 'Responsable', 'Costo']]
        for _, row in df_oper.head(30).iterrows():  # Mostrar primeras 30 para no sobrecargar
            datos_oper.append([
                row['estrategia_nombre'][:25] + '...' if len(row['estrategia_nombre']) > 25 else row['estrategia_nombre'],
                row['descripcion_actividad'][:35] + '...' if len(row['descripcion_actividad']) > 35 else row['descripcion_actividad'],
                row['plazo'] or 'Pendiente',
                row['responsable'] or 'Sin asignar',
                f"${row['costo']:,.2f}"
            ])
        
        tabla_oper = create_table_pdf(datos_oper, col_widths=[1.8*inch, 2.2*inch, 1*inch, 1.2*inch, 1*inch])
        story.append(tabla_oper)
        story.append(Spacer(1, 0.1*inch))
        
        if len(df_oper) > 30:
            story.append(Paragraph(
                f"<i>Nota: Se muestran 30 de {len(df_oper)} actividades. El detalle completo está en los anexos.</i>",
                styles['BodyTextEnhanced']
            ))
        
        # Actividades de mayor inversión
        story.append(Paragraph("4.2.3 Actividades de Mayor Inversión", styles['Heading3Enhanced']))
        
        actividades_criticas = df_oper.nlargest(5, 'costo')
        for idx, row in actividades_criticas.iterrows():
            story.append(Paragraph(
                f"${row['costo']:,.2f} - {row['descripcion_actividad'][:80]}... "
                f"({row['responsable'] or 'Sin asignar'})",
                styles['APA_List']
            ))
    else:
        story.append(Paragraph("No se ha completado la operativización de estrategias.", styles['BodyTextEnhanced']))
    
    story.append(PageBreak())
    
    # 5. EVALUACIÓN Y CONTROL (CMI + Semaforización)
    story.append(Paragraph("5. EVALUACIÓN Y CONTROL", styles['Heading1Enhanced']))
    story.append(Paragraph(
        "La evaluación y control permiten monitorear el desempeño de las estrategias mediante "
        "indicadores clave de desempeño (KPIs) y sistemas de alerta temprana.",
        styles['BodyTextEnhanced']
    ))
    story.append(Spacer(1, 0.2*inch))
    
    # 5.1 Cuadro de Mando Integral (CMI)
    story.append(Paragraph("5.1 Cuadro de Mando Integral (CMI)", styles['Heading2Enhanced']))
    
    if not df_estrategias.empty:
        story.append(Paragraph(
            "El Cuadro de Mando Integral traduce las estrategias en indicadores medibles "
            "desde cuatro perspectivas: Financiera, Cliente, Procesos Internos, y Aprendizaje y Crecimiento.",
            styles['BodyTextEnhanced']
        ))
        story.append(Spacer(1, 0.2*inch))
        
        # Intentar obtener CMI guardado
        cmi_guardado = empresa.get('analisis_cmi', '')
        df_cmi = pd.DataFrame()
        
        if cmi_guardado and '|' in str(cmi_guardado):
            try:
                df_cmi = pd.read_csv(io.StringIO(cmi_guardado), sep="|")
            except:
                pass
        
        if df_cmi.empty:
            try:
                df_cmi = generar_cuadro_de_mando_ia(df_estrategias)
            except:
                pass
        
        if not df_cmi.empty:
            story.append(Paragraph("5.1.1 Indicadores Clave por Perspectiva", styles['Heading3Enhanced']))
            
            datos_cmi = [['Estrategia', 'Perspectiva', 'KPI', 'Frecuencia', 'Límites']]
            for _, row in df_cmi.head(15).iterrows():
                limites = f"LI:{row.get('LI','')} LC:{row.get('LC','')} LS:{row.get('LS','')}"
                datos_cmi.append([
                    row['Estrategia'][:35] + '...' if len(row['Estrategia']) > 35 else row['Estrategia'],
                    row['Perspectiva'],
                    row['KPIs'][:30] + '...' if len(row['KPIs']) > 30 else row['KPIs'],
                    row['Frecuencia'],
                    limites
                ])
            
            tabla_cmi = create_table_pdf(datos_cmi, col_widths=[2*inch, 1.2*inch, 1.8*inch, 0.8*inch, 1.5*inch])
            story.append(tabla_cmi)
            
            if len(df_cmi) > 15:
                story.append(Paragraph(
                    f"<i>Nota: Se muestran 15 de {len(df_cmi)} indicadores. El detalle completo está en los anexos.</i>",
                    styles['BodyTextEnhanced']
                ))
        else:
            story.append(Paragraph("No se pudo generar el CMI automáticamente.", styles['BodyTextEnhanced']))
    else:
        story.append(Paragraph("No hay estrategias disponibles para construir el CMI.", styles['BodyTextEnhanced']))
    
    story.append(Spacer(1, 0.2*inch))
    
    # 5.2 Semaforización Estratégica
    story.append(Paragraph("5.2 Semaforización Estratégica", styles['Heading2Enhanced']))
    
    analisis_semaforo = empresa.get('analisis_semaforo_resumen', '')
    if analisis_semaforo:
        story.append(Paragraph(analisis_semaforo, styles['BodyTextEnhanced']))
    else:
        story.append(Paragraph(
            "El sistema de semaforización evalúa el estado de cada estrategia considerando alineación "
            "con objetivos, recursos asignados y contexto externo. Las estrategias se clasifican en: "
            "<b>🟢 Óptimas</b> (implementar según plan), <b>🟡 Atención</b> (requieren ajustes), "
            "y <b>🔴 Críticas</b> (necesitan revisión inmediata).",
            styles['BodyTextEnhanced']
        ))
    
    story.append(PageBreak())
    
    # ============================================
    # PARTE 3: ANEXOS (ilimitado)
    # ============================================
    
    story.append(Paragraph("ANEXOS", styles['Heading1Enhanced']))
    story.append(Spacer(1, 0.2*inch))
    
    # Anexo A: Análisis Detallados de Matrices
    story.append(Paragraph("Anexo A. Análisis Detallados de Matrices", styles['Heading2Enhanced']))
    
    # A.1 MADE Completo
    if not df_made.empty:
        story.append(Paragraph("A.1 Matriz MADE (Marketing Interno) - Datos Completos", styles['Heading3Enhanced']))
        
        datos_made_full = [['Variable', 'Factor', 'Producto', 'Precio', 'Plaza', 'Promoción', 'Rating', 'Peso %', 'Valor']]
        for _, row in df_made.iterrows():
            datos_made_full.append([
                str(row.get('variable', ''))[:15],
                str(row.get('factor', ''))[:25],
                str(row.get('producto', ''))[:10],
                str(row.get('precio', ''))[:10],
                str(row.get('plaza', ''))[:10],
                str(row.get('promocion', ''))[:10],
                str(row.get('rating', '')),
                f"{row.get('weight_percent', '')}%",
                str(row.get('valor', ''))
            ])
        
        tabla_made_full = create_table_pdf(datos_made_full, col_widths=[0.9*inch, 1.5*inch, 0.7*inch, 0.7*inch, 0.7*inch, 0.7*inch, 0.6*inch, 0.7*inch, 0.6*inch])
        story.append(tabla_made_full)
        story.append(PageBreak())
    
    # A.2 MADI Completo
    if not df_madi.empty:
        story.append(Paragraph("A.2 Matriz MADI (Marketing Externo) - Datos Completos", styles['Heading3Enhanced']))
        
        datos_madi_full = [['Variable', 'Factor', 'Producto', 'Precio', 'Plaza', 'Promoción', 'Rating', 'Peso %', 'Valor']]
        for _, row in df_madi.iterrows():
            datos_madi_full.append([
                str(row.get('variable', ''))[:15],
                str(row.get('factor', ''))[:25],
                str(row.get('producto', ''))[:10],
                str(row.get('precio', ''))[:10],
                str(row.get('plaza', ''))[:10],
                str(row.get('promocion', ''))[:10],
                str(row.get('rating', '')),
                f"{row.get('weight_percent', '')}%",
                str(row.get('valor', ''))
            ])
        
        tabla_madi_full = create_table_pdf(datos_madi_full, col_widths=[0.9*inch, 1.5*inch, 0.7*inch, 0.7*inch, 0.7*inch, 0.7*inch, 0.6*inch, 0.7*inch, 0.6*inch])
        story.append(tabla_madi_full)
        story.append(PageBreak())
    
    # A.3 PEST Completo
    if not df_pest.empty:
        story.append(Paragraph("A.3 Análisis PEST Completo", styles['Heading3Enhanced']))
        
        datos_pest_full = [['Categoría', 'Factor', 'Tipo FODA', 'Puntaje', 'Importancia', 'Ponderado']]
        for _, row in df_pest.iterrows():
            datos_pest_full.append([
                row['categoria'],
                row['factor'][:40] + '...' if len(row['factor']) > 40 else row['factor'],
                row['tipo_foda'],
                str(row['puntaje']),
                f"{row['importancia']}%",
                f"{row['valor_ponderado']:.2f}"
            ])
        
        tabla_pest_full = create_table_pdf(datos_pest_full, col_widths=[1*inch, 3*inch, 1*inch, 0.8*inch, 1*inch, 1*inch])
        story.append(tabla_pest_full)
        story.append(Spacer(1, 0.1*inch))
        
        # Gráfico PEST detallado
        grafico_pest = generar_grafico_barras_pest_mejorado(df_pest)
        if grafico_pest:
            story.append(Image(grafico_pest, width=6*inch, height=4*inch))
            story.append(Paragraph(
                "Figura A.1. Análisis PEST - Distribución de factores por categoría.",
                ParagraphStyle(name='Caption', parent=styles['BodyTextEnhanced'], alignment=TA_CENTER, fontSize=10, firstLineIndent=0)
            ))
        story.append(PageBreak())
    
    # A.4 FODA Cruzado Completo
    if not df_foda.empty:
        story.append(Paragraph("A.4 Matriz FODA Cruzado Completa", styles['Heading3Enhanced']))
        
        datos_foda_full = [['Cuadrante', 'Factor Fila', 'Factor Columna', 'Impacto']]
        for _, row in df_foda.iterrows():
            datos_foda_full.append([
                row['cuadrante'],
                str(row['factor_fila']),
                str(row['factor_columna']),
                str(row['impacto'])
            ])
        
        tabla_foda_full = create_table_pdf(datos_foda_full, col_widths=[1*inch, 2.5*inch, 2.5*inch, 0.8*inch])
        story.append(tabla_foda_full)
        story.append(Spacer(1, 0.1*inch))
        
        # Gráfico FODA
        if puntajes_foda is not None and not puntajes_foda.empty:
            grafico_foda = generar_grafico_foda_radar_mejorado(puntajes_foda)
            if grafico_foda:
                story.append(Image(grafico_foda, width=5*inch, height=5*inch))
                story.append(Paragraph(
                    "Figura A.2. Posicionamiento estratégico - Matriz FODA cruzada.",
                    ParagraphStyle(name='Caption', parent=styles['BodyTextEnhanced'], alignment=TA_CENTER, fontSize=10, firstLineIndent=0)
                ))
        story.append(PageBreak())
    
    # Anexo B: Dashboards y Visualizaciones
    story.append(Paragraph("Anexo B. Dashboards de Análisis Estratégico", styles['Heading2Enhanced']))
    
    story.append(Paragraph(
        "Este anexo presenta las visualizaciones completas del análisis estratégico.",
        styles['BodyTextEnhanced']
    ))
    story.append(Spacer(1, 0.2*inch))
    
    # B.1 Distribución de Estrategias
    if not df_estrategias.empty:
        story.append(Paragraph("B.1 Distribución de Estrategias por Cuadrante", styles['Heading3Enhanced']))
        
        fig, ax = plt.subplots(figsize=(7, 4))
        est_counts = df_estrategias['cuadrante'].value_counts()
        colors_est = {'FO': '#2ca02c', 'FA': '#d62728', 'DO': '#ff7f0e', 'DA': '#9467bd'}
        bars = ax.bar(est_counts.index, est_counts.values, 
                     color=[colors_est.get(x, '#1f77b4') for x in est_counts.index],
                     edgecolor='black', linewidth=0.5)
        ax.set_title('Distribución de Estrategias por Cuadrante FODA', fontsize=11, fontweight='bold')
        ax.set_ylabel('Número de Estrategias')
        
        for bar in bars:
            height = bar.get_height()
            ax.text(bar.get_x() + bar.get_width()/2., height, f'{int(height)}', 
                   ha='center', va='bottom')
        
        plt.tight_layout()
        buf_est = BytesIO()
        plt.savefig(buf_est, format='PNG', dpi=150, bbox_inches='tight', facecolor='white')
        plt.close()
        buf_est.seek(0)
        
        story.append(Image(buf_est, width=5*inch, height=3*inch))
        story.append(Spacer(1, 0.2*inch))
    
    # B.2 Proyección Financiera Detallada
    if not df_proy.empty:
        story.append(Paragraph("B.2 Proyección Financiera Detallada", styles['Heading3Enhanced']))
        
        grafico_proy = generar_grafico_proyeccion_mejorado(df_proy)
        if grafico_proy:
            story.append(Image(grafico_proy, width=6*inch, height=4*inch))
            story.append(Spacer(1, 0.1*inch))
        
        datos_proy = [['Año', 'Ingresos Proyectados', 'Costos Proyectados', 'Utilidad Neta Proyectada']]
        for _, row in df_proy.iterrows():
            datos_proy.append([
                str(int(row['anio'])),
                f"${row['ingresos_proyectados']:,.0f}",
                f"${row['costos_proyectados']:,.0f}",
                f"${row['utilidad_neta_proyectada']:,.0f}"
            ])
        
        tabla_proy = create_table_pdf(datos_proy, col_widths=[1*inch, 1.8*inch, 1.8*inch, 1.8*inch])
        story.append(tabla_proy)
        story.append(PageBreak())
    
    # B.3 Análisis Costo-Beneficio Detallado
    if not df_cb.empty:
        story.append(Paragraph("B.3 Análisis Costo-Beneficio Detallado", styles['Heading3Enhanced']))
        
        datos_cb = df_cb.iloc[0]
        
        indicadores_cb = [
            ['Indicador', 'Valor', 'Umbral Aceptable', 'Estado'],
            ['Relación C-B ($)', f"{float(datos_cb['relacion_costo_beneficio_dolares']):.2f}", '≥ 1.0', 
             'Aceptable' if datos_cb['relacion_costo_beneficio_dolares'] >= 1 else 'No aceptable'],
            ['Payback (años)', 
             f"{float(datos_cb['payback_periodo_anios']):.1f}" if datos_cb['payback_periodo_anios'] else 'N/A',
             '≤ 5', 
             'Aceptable' if datos_cb['payback_periodo_anios'] and datos_cb['payback_periodo_anios'] <= 5 else 'Revisar'],
            ['VPN Total', f"${float(datos_cb['vpn_total']):,.2f}", '> 0', 
             'Positivo' if datos_cb['vpn_total'] > 0 else 'Negativo'],
            ['Inversión Total', f"${float(datos_cb['inversion_total']):,.2f}", '-', 'Requerimiento'],
            ['Beneficio/Unidad', f"${float(datos_cb['beneficio_por_unidad']):,.2f}", '≥ Costo/Unidad', 
             'Rentable' if datos_cb['relacion_cb_unidades'] >= 1 else 'No rentable'],
        ]
        
        tabla_cb = create_table_pdf(indicadores_cb, col_widths=[1.8*inch, 1.5*inch, 1.5*inch, 1.5*inch])
        story.append(tabla_cb)
        story.append(PageBreak())
    
    # Anexo C: Operativización Completa
    if not df_oper.empty:
        story.append(Paragraph("Anexo C. Cuadro de Operativización Completo", styles['Heading2Enhanced']))
        
        # Tabla completa de operativización (todas las actividades)
        datos_oper_full = [['N°', 'Estrategia', 'Actividad', 'Plazo', 'Responsable', 'Costo']]
        for idx, row in df_oper.iterrows():
            datos_oper_full.append([
                str(idx + 1),
                row['estrategia_nombre'][:30] + '...' if len(row['estrategia_nombre']) > 30 else row['estrategia_nombre'],
                row['descripcion_actividad'][:40] + '...' if len(row['descripcion_actividad']) > 40 else row['descripcion_actividad'],
                row['plazo'] or 'Pendiente',
                row['responsable'] or 'Sin asignar',
                f"${row['costo']:,.2f}"
            ])
        
        # Dividir en múltiples tablas si es muy largo
        chunk_size = 40
        for i in range(0, len(datos_oper_full), chunk_size):
            chunk = datos_oper_full[i:i+chunk_size]
            if i == 0:
                tabla_oper_full = create_table_pdf(chunk, col_widths=[0.5*inch, 2*inch, 2.5*inch, 0.8*inch, 1.2*inch, 1*inch])
            else:
                # Sin encabezado para continuación
                tabla_oper_full = create_table_pdf(chunk, col_widths=[0.5*inch, 2*inch, 2.5*inch, 0.8*inch, 1.2*inch, 1*inch])
            story.append(tabla_oper_full)
            if i + chunk_size < len(datos_oper_full):
                story.append(PageBreak())
    
    # Anexo D: CMI Completo
    if not df_cmi.empty:
        story.append(PageBreak())
        story.append(Paragraph("Anexo D. Cuadro de Mando Integral Completo", styles['Heading2Enhanced']))
        
        datos_cmi_full = [['Estrategia', 'Perspectiva', 'KPIs', 'Fórmulas', 'Frecuencia', 'LI', 'LC', 'LS']]
        for _, row in df_cmi.iterrows():
            datos_cmi_full.append([
                row['Estrategia'][:40] + '...' if len(row['Estrategia']) > 40 else row['Estrategia'],
                row['Perspectiva'],
                row['KPIs'][:35] + '...' if len(row['KPIs']) > 35 else row['KPIs'],
                row['Formulas'][:25] + '...' if len(row['Formulas']) > 25 else row['Formulas'],
                row['Frecuencia'],
                str(row['LI']),
                str(row['LC']),
                str(row['LS'])
            ])
        
        tabla_cmi_full = create_table_pdf(datos_cmi_full, col_widths=[1.8*inch, 1*inch, 1.8*inch, 1.2*inch, 0.8*inch, 0.6*inch, 0.6*inch, 0.6*inch])
        story.append(tabla_cmi_full)
    
    # Construir el PDF con manejo de errores mejorado
    try:
        from reportlab.platypus import PageTemplate

        # Crear PageTemplate con Frame para limitar área de contenido
        page_template = PageTemplate(
            id='content_page', 
            frames=content_frame, 
            onPage=header_footer
        )
        doc.addPageTemplates([page_template])

        # Construir documento con manejo de errores por elemento
        try:
            doc.build(story)
        except Exception as build_error:
            # Si falla, intentar identificar qué elemento causó el problema
            st.error(f"Error en construcción del PDF: {build_error}")
            
            # Intentar construir con story mínima para diagnóstico
            story_minima = [
                Paragraph("DOCUMENTO DE PRUEBA - ERROR EN GENERACION", styles['Heading1Enhanced']),
                Paragraph(f"Error detectado: {str(build_error)[:200]}", styles['BodyTextEnhanced']),
                Paragraph("Verifique que los textos de la IA no contengan caracteres especiales.", styles['BodyTextEnhanced'])
            ]
            doc.build(story_minima)
            
        pdf_buffer.seek(0)
        return pdf_buffer
        
    except Exception as e:
        st.error(f"Error fatal al generar PDF: {e}")
        import traceback
        st.error(traceback.format_exc())
        return None
        

def generar_word_completo_mejorada(empresa_id, version, elaborado, revisado, aprobado):
    """
    Genera el documento Word completo con formato profesional
    """
    from datetime import datetime
    
    empresa = get_datos_empresa(empresa_id)
    if not empresa:
        return None
    
    # Crear documento
    doc = Document()
    
    # Configurar márgenes
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # PORTADA
    doc.add_heading('PLAN ESTRATÉGICO', 0)
    doc.add_heading(empresa.get('nombre', 'EMPRESA').upper(), level=1)
    
    # Información del documento
    doc.add_paragraph(f'Versión: {version}')
    doc.add_paragraph(f'Fecha: {datetime.now().strftime("%d/%m/%Y")}')
    doc.add_paragraph()
    doc.add_paragraph(f'Elaborado por: {elaborado}')
    doc.add_paragraph(f'Revisado por: {revisado}')
    doc.add_paragraph(f'Aprobado por: {aprobado}')
    
    doc.add_page_break()
    
    # 1. INTRODUCCIÓN
    doc.add_heading('1. INTRODUCCIÓN Y FUNDAMENTOS', level=1)
    
    doc.add_heading('1.1 Datos Generales de la Empresa', level=2)
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Light Grid Accent 1'
    datos = [
        ('Nombre', empresa.get('nombre', 'N/A')),
        ('Giro', empresa.get('giro', 'N/A')),
        ('Fecha', datetime.now().strftime('%d/%m/%Y')),
        ('Versión', version),
        ('Elaborado por', elaborado)
    ]
    for i, (campo, valor) in enumerate(datos):
        table.rows[i].cells[0].text = campo
        table.rows[i].cells[1].text = str(valor)
    
    # Cultura organizacional
    if empresa.get('mision'):
        doc.add_heading('Misión', level=3)
        doc.add_paragraph(empresa['mision'])
    
    if empresa.get('vision'):
        doc.add_heading('Visión', level=3)
        doc.add_paragraph(empresa['vision'])
    
    if empresa.get('valores'):
        doc.add_heading('Valores', level=3)
        doc.add_paragraph(empresa['valores'])
    
    doc.add_page_break()
    
    # Obtener datos de matrices
    df_pest = get_datos_tabla('matrices', empresa_id, tipo_matriz_filter='PEST')
    df_foda = get_datos_tabla('foda_cruzado', empresa_id)
    df_estrategias = get_datos_tabla('estrategias_generadas', empresa_id)
    df_oper = get_datos_tabla('operativizacion', empresa_id)
    
    # 2. ANÁLISIS SITUACIONAL
    doc.add_heading('2. ANÁLISIS SITUACIONAL', level=1)
    
    if not df_pest.empty:
        doc.add_heading('2.1 Análisis PEST', level=2)
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Light List Accent 1'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Categoría'
        hdr_cells[1].text = 'Factor'
        hdr_cells[2].text = 'Tipo FODA'
        hdr_cells[3].text = 'Puntaje'
        hdr_cells[4].text = 'Ponderado'
        
        for _, row in df_pest.head(20).iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = str(row['categoria'])
            row_cells[1].text = str(row['factor'])[:50]
            row_cells[2].text = str(row['tipo_foda'])
            row_cells[3].text = str(row['puntaje'])
            row_cells[4].text = f"{row['valor_ponderado']:.2f}"
    
    if not df_foda.empty:
        doc.add_heading('2.2 Matriz FODA Cruzado', level=2)
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Light List Accent 1'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Cuadrante'
        hdr_cells[1].text = 'Factor Fila'
        hdr_cells[2].text = 'Factor Columna'
        hdr_cells[3].text = 'Impacto'
        
        for _, row in df_foda.head(20).iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = str(row['cuadrante'])
            row_cells[1].text = str(row['factor_fila'])[:40]
            row_cells[2].text = str(row['factor_columna'])[:40]
            row_cells[3].text = str(row['impacto'])
    
    doc.add_page_break()
    
    # 3. ESTRATEGIAS
    doc.add_heading('3. ESTRATEGIAS', level=1)
    
    if not df_estrategias.empty:
        for idx, row in df_estrategias.iterrows():
            doc.add_heading(f"Estrategia {idx + 1}: [{row['cuadrante']}]", level=2)
            doc.add_paragraph(f"Descripción: {row['estrategia']}")
            doc.add_paragraph(f"Plan asignado: {row['plan_asignado']}")
            doc.add_paragraph(f"Importancia: {row['importancia']}")
            doc.add_paragraph(f"Actividades: {row['actividades']}")
            doc.add_paragraph()
    else:
        doc.add_paragraph("No se han generado estrategias.")
    
    doc.add_page_break()
    
    # 4. PLAN DE ACCIÓN
    doc.add_heading('4. PLAN DE ACCIÓN', level=1)
    
    if not df_oper.empty:
        doc.add_heading('4.1 Operativización', level=2)
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Light List Accent 1'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Estrategia'
        hdr_cells[1].text = 'Actividad'
        hdr_cells[2].text = 'Plazo'
        hdr_cells[3].text = 'Responsable'
        hdr_cells[4].text = 'Costo'
        
        total_costo = 0
        for _, row in df_oper.head(30).iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = str(row['estrategia_nombre'])[:30]
            row_cells[1].text = str(row['descripcion_actividad'])[:40]
            row_cells[2].text = str(row['plazo'] or 'Pendiente')
            row_cells[3].text = str(row['responsable'] or 'Sin asignar')
            costo = float(row['costo'] or 0)
            row_cells[4].text = f"${costo:,.2f}"
            total_costo += costo
        
        doc.add_paragraph()
        doc.add_paragraph(f"**Inversión Total: ${total_costo:,.2f}**").bold = True
    else:
        doc.add_paragraph("No hay datos de operativización.")
    
    # Guardar en buffer
    word_buffer = io.BytesIO()
    doc.save(word_buffer)
    word_buffer.seek(0)
    return word_buffer

def generar_word_completo_mejorado(empresa_id, version, elaborado, revisado, aprobado):
    """
    Genera el documento Word completo con el mismo diseño profesional que el PDF.
    Incluye encabezados/pies de página, colores corporativos, tablas estilizadas y gráficos.
    """
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from io import BytesIO
    from datetime import datetime
    
    empresa = get_datos_empresa(empresa_id)
    if not empresa:
        return None
    
    # Obtener todos los datos necesarios (igual que en PDF)
    df_pest = get_datos_tabla('matrices', empresa_id, tipo_matriz_filter='PEST')
    df_foda = get_datos_tabla('foda_cruzado', empresa_id)
    df_estrategias = get_datos_tabla('estrategias_generadas', empresa_id)
    df_oper = get_datos_tabla('operativizacion', empresa_id)
    df_pg = get_datos_tabla('perdida_ganancia', empresa_id)
    df_proy = get_datos_tabla('proyeccion_financiera', empresa_id)
    df_cb = get_datos_tabla('analisis_costo_beneficio', empresa_id)
    df_made = get_datos_tabla('matriz_marketing', empresa_id, tipo_matriz_filter='MADE')
    df_madi = get_datos_tabla('matriz_marketing', empresa_id, tipo_matriz_filter='MADI')
    
    # Calcular análisis FODA
    analisis_foda_df, resumen_foda, estrategia_principal, puntajes_foda = analizar_foda(df_foda)
    pest_total = df_pest['valor_ponderado'].sum() if not df_pest.empty else 0
    total_costo = pd.to_numeric(df_oper['costo'], errors='coerce').fillna(0).sum() if not df_oper.empty else 0
    
    fecha_actual = datetime.now().strftime('%d/%m/%Y')
    
    # Crear documento
    doc = Document()
    
    # Configurar márgenes (1 pulgada = 2.54 cm en todos los lados - formato APA)
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
    
    # =========================================================================
    # CONFIGURAR ESTILOS PERSONALIZADOS (Colores corporativos)
    # =========================================================================
    
    # Colores corporativos
    COLOR_PRIMARY = RGBColor(30, 58, 95)      # Azul marino #1e3a5f
    COLOR_SECONDARY = RGBColor(201, 162, 39)  # Dorado #c9a227
    COLOR_TEXT = RGBColor(31, 41, 55)         # Gris oscuro #1f2937
    COLOR_WHITE = RGBColor(255, 255, 255)
    
    # Modificar estilo Normal
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Helvetica'
    style_normal.font.size = Pt(11)
    style_normal.font.color.rgb = COLOR_TEXT
    
    # Crear/Modificar estilos de título
    def set_heading_style(doc, level, font_size, color, bold=True):
        """Configura estilo de título"""
        style_name = f'Heading {level}'
        try:
            style = doc.styles[style_name]
        except KeyError:
            style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        
        style.font.name = 'Helvetica'
        style.font.size = Pt(font_size)
        style.font.color.rgb = color
        style.font.bold = bold
        style.paragraph_format.space_before = Pt(12 if level > 1 else 24)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.keep_with_next = True
        return style
    
    # Configurar jerarquía de títulos
    set_heading_style(doc, 1, 16, COLOR_PRIMARY)      # Título principal
    set_heading_style(doc, 2, 14, COLOR_PRIMARY)      # Secciones
    set_heading_style(doc, 3, 12, COLOR_TEXT)         # Subsecciones
    set_heading_style(doc, 4, 11, COLOR_TEXT, bold=True)  # Sub-subsecciones
    
    # =========================================================================
    # FUNCIONES AUXILIARES PARA FORMATO
    # =========================================================================
    
    def add_custom_heading(doc, text, level=1):
        """Agrega título con formato personalizado"""
        p = doc.add_heading(text, level=level)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        return p
    
    def add_formatted_paragraph(doc, text, bold=False, italic=False, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
        """Agrega párrafo con formato"""
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Helvetica'
        run.font.size = Pt(11)
        run.font.color.rgb = COLOR_TEXT
        run.bold = bold
        run.italic = italic
        p.alignment = alignment
        p.paragraph_format.first_line_indent = Cm(1.25) if alignment == WD_ALIGN_PARAGRAPH.JUSTIFY else Cm(0)
        p.paragraph_format.space_after = Pt(8)
        return p
    
    def add_bullet_point(doc, text, indent_level=0):
        """Agrega viñeta con formato"""
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent = Cm(1.25 + (indent_level * 0.5))
        p.paragraph_format.first_line_indent = Cm(-0.5)
        run = p.add_run(text)
        run.font.name = 'Helvetica'
        run.font.size = Pt(11)
        run.font.color.rgb = COLOR_TEXT
        return p
    
    def set_cell_shading(cell, color):
        """Aplica color de fondo a celda"""
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), color)
        cell._tc.get_or_add_tcPr().append(shading_elm)
    
    def set_cell_border(cell, **kwargs):
        """Configura bordes de celda"""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        
        for edge in ('top', 'left', 'bottom', 'right'):
            edge_data = kwargs.get(edge)
            if edge_data:
                tag = f'w:{edge}'
                element = OxmlElement(tag)
                for key in ["sz", "val", "color", "space"]:
                    if key in edge_data:
                        element.set(qn(f'w:{key}'), str(edge_data[key]))
                tcBorders.append(element)
        
        tcPr.append(tcBorders)
    
    def create_professional_table(doc, headers, data, col_widths=None):
        """Crea tabla con estilo profesional"""
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        
        # Configurar ancho de columnas si se especifica
        if col_widths:
            for i, width in enumerate(col_widths):
                for cell in table.columns[i].cells:
                    cell.width = Cm(width)
        
        # Encabezado con color corporativo
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            set_cell_shading(hdr_cells[i], '1e3a5f')  # Azul marino
            set_cell_border(hdr_cells[i], 
                          top={"sz": "12", "val": "single", "color": "1e3a5f"},
                          bottom={"sz": "12", "val": "single", "color": "c9a227"})  # Línea dorada abajo
            
            # Formato texto encabezado
            for paragraph in hdr_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Helvetica'
                    run.font.size = Pt(10)
                    run.font.bold = True
                    run.font.color.rgb = COLOR_WHITE
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Datos
        for row_data in data:
            row_cells = table.add_row().cells
            for i, cell_text in enumerate(row_data):
                row_cells[i].text = str(cell_text)
                # Alternar colores de fondo (zebra striping)
                if len(table.rows) % 2 == 0:
                    set_cell_shading(row_cells[i], 'f8fafc')  # Gris muy claro
                
                # Formato
                for paragraph in row_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Helvetica'
                        run.font.size = Pt(9)
                        run.font.color.rgb = COLOR_TEXT
        
        doc.add_paragraph()  # Espacio después de tabla
        return table
    
    def add_page_break(doc):
        """Agrega salto de página"""
        doc.add_page_break()
    
    # =========================================================================
    # ENCABEZADO Y PIE DE PÁGINA
    # =========================================================================
    
    section = doc.sections[0]
    
    # Encabezado
    header = section.header
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Tabla en encabezado para layout profesional
    header_table = header.add_table(1, 3, width=Inches(6))
    header_table.autofit = False
    
    # Celda izquierda: Logo (si existe)
    logo_cell = header_table.rows[0].cells[0]
    logo_cell.width = Inches(1)
    
    logo_bytes_data = empresa.get('logo')
    if logo_bytes_data:
        try:
            if isinstance(logo_bytes_data, str):
                import base64
                logo_bytes = base64.b64decode(logo_bytes_data)
            else:
                logo_bytes = logo_bytes_data
            
            # Verificar que sea imagen válida
            from PIL import Image as PILImage
            test_img = PILImage.open(BytesIO(logo_bytes))
            test_img.verify()
            
            logo_para = logo_cell.paragraphs[0]
            logo_run = logo_para.add_run()
            logo_run.add_picture(BytesIO(logo_bytes), width=Inches(0.8))
            logo_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        except:
            pass
    
    # Celda centro: Nombre empresa
    center_cell = header_table.rows[0].cells[1]
    center_cell.width = Inches(4)
    center_para = center_cell.paragraphs[0]
    center_run = center_para.add_run(empresa.get('nombre', 'EMPRESA').upper())
    center_run.font.name = 'Helvetica'
    center_run.font.size = Pt(12)
    center_run.font.bold = True
    center_run.font.color.rgb = COLOR_PRIMARY
    center_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Celda derecha: Versión y fecha
    right_cell = header_table.rows[0].cells[2]
    right_cell.width = Inches(1)
    right_para = right_cell.paragraphs[0]
    right_run = right_para.add_run(f"Versión {version}\n{fecha_actual}")
    right_run.font.name = 'Helvetica'
    right_run.font.size = Pt(8)
    right_run.font.color.rgb = COLOR_TEXT
    right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Línea decorativa debajo del encabezado
    header_border = header.add_paragraph()
    header_border.paragraph_format.space_before = Pt(2)
    header_border_run = header_border.add_run("_" * 80)
    header_border_run.font.color.rgb = COLOR_SECONDARY
    
    # Pie de página
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Tabla en pie de página
    footer_table = footer.add_table(1, 3, width=Inches(6))
    
    # Elaborado/Revisado/Aprobado
    footer_table.rows[0].cells[0].text = f"Elaborado: {elaborado}"
    footer_table.rows[0].cells[1].text = f"Revisado: {revisado}"
    footer_table.rows[0].cells[2].text = f"Aprobado: {aprobado}"
    
    for cell in footer_table.rows[0].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Helvetica'
                run.font.size = Pt(8)
                run.font.color.rgb = COLOR_TEXT
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Línea decorativa y número de página
    footer_line = footer.add_paragraph()
    footer_line_run = footer_line.add_run("─" * 30)
    footer_line_run.font.color.rgb = COLOR_SECONDARY
    footer_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    page_num = footer.add_paragraph()
    page_num_run = page_num.add_run("Página ")
    page_num_run.font.name = 'Helvetica'
    page_num_run.font.size = Pt(9)
    page_num_run.font.bold = True
    page_num_run.font.color.rgb = COLOR_PRIMARY
    
    # Agregar campo de número de página (se actualiza automáticamente en Word)
    page_num_run2 = page_num.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.text = "PAGE"
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    
    page_num_run2._r.append(fldChar1)
    page_num_run2._r.append(instrText)
    page_num_run2._r.append(fldChar2)
    
    page_num.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # =========================================================================
    # CONTENIDO DEL DOCUMENTO
    # =========================================================================
    
    # ============================================
    # PARTE 1: RESUMEN EJECUTIVO
    # ============================================
    
    add_custom_heading(doc, "RESUMEN EJECUTIVO", level=1)
    
    intro_text = (f"El presente documento constituye el Plan Estratégico de {empresa.get('nombre', 'la empresa')}, "
                  f"elaborado con fecha {fecha_actual}. Este resumen ejecutivo presenta los hallazgos más relevantes "
                  f"del diagnóstico estratégico y las recomendaciones prioritarias para la alta dirección.")
    add_formatted_paragraph(doc, intro_text)
    
    # Diagnóstico clave
    add_custom_heading(doc, "Diagnóstico Estratégico Clave", level=2)
    
    if estrategia_principal:
        estrategia_text = (f"Estrategia Principal Recomendada: {estrategia_principal}. "
                          f"Esta postura estratégica se determina a partir del análisis FODA cruzado y representa "
                          f"la orientación prioritaria para el período de planificación.")
        add_formatted_paragraph(doc, estrategia_text, bold=True)
    
    # Gráfico FODA (insertar imagen si existe)
    if puntajes_foda is not None and not puntajes_foda.empty:
        grafico_foda = generar_grafico_foda_radar_mejorado(puntajes_foda)
        if grafico_foda:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(grafico_foda, width=Inches(3.5))
            
            caption = doc.add_paragraph()
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_run = caption.add_run("Figura 1. Posicionamiento estratégico según análisis FODA cruzado.")
            caption_run.font.name = 'Helvetica'
            caption_run.font.size = Pt(10)
            caption_run.font.italic = True
            caption_run.font.color.rgb = RGBColor(107, 114, 128)  # Gris
    
    # Análisis PEST resumido
    if not df_pest.empty:
        add_custom_heading(doc, "Factores Críticos del Entorno", level=2)
        
        pest_criticos = df_pest.nlargest(3, 'valor_ponderado')
        for _, row in pest_criticos.iterrows():
            add_bullet_point(doc, f"{row['categoria']}: {row['factor'][:100]}...")
        
        # Gráfico PEST
        grafico_pest = generar_grafico_barras_pest_mejorado(df_pest)
        if grafico_pest:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(grafico_pest, width=Inches(4))
            
            caption = doc.add_paragraph()
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_run = caption.add_run("Figura 2. Distribución de factores PEST por impacto.")
            caption_run.font.name = 'Helvetica'
            caption_run.font.size = Pt(10)
            caption_run.font.italic = True
            caption_run.font.color.rgb = RGBColor(107, 114, 128)
    
    # Estrategias priorizadas
    if not df_estrategias.empty:
        add_custom_heading(doc, "Estrategias Prioritarias", level=2)
        
        estrategias_alta = df_estrategias[df_estrategias['importancia'].isin(['Alta', 'Media Alta'])].head(5)
        for _, row in estrategias_alta.iterrows():
            add_bullet_point(doc, f"{row['cuadrante']} - {row['plan_asignado']}: {row['estrategia'][:150]}...")
    
    # Indicadores financieros
    if not df_cb.empty:
        add_custom_heading(doc, "Viabilidad Financiera", level=2)
        
        datos_cb = df_cb.iloc[0]
        
        # Tabla resumen financiero
        headers = ['Indicador', 'Valor', 'Interpretación']
        data = [
            ['Relación Costo-Beneficio', f"{float(datos_cb['relacion_costo_beneficio_dolares']):.2f}", 
             'Rentable' if datos_cb['relacion_costo_beneficio_dolares'] >= 1 else 'No rentable'],
            ['Periodo de Recuperación', 
             f"{float(datos_cb['payback_periodo_anios']):.1f} años" if datos_cb['payback_periodo_anios'] else 'N/A',
             'Aceptable' if datos_cb['payback_periodo_anios'] and datos_cb['payback_periodo_anios'] <= 5 else 'Revisar'],
            ['Inversión Total', f"${float(datos_cb['inversion_total']):,.2f}", 'Requerimiento de capital'],
        ]
        create_professional_table(doc, headers, data, col_widths=[5, 3.5, 4])
        
        # Gráfico de proyección
        if not df_proy.empty:
            grafico_proy = generar_grafico_proyeccion_mejorado(df_proy)
            if grafico_proy:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(grafico_proy, width=Inches(5))
                
                caption = doc.add_paragraph()
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption_run = caption.add_run("Figura 3. Proyección financiera del plan estratégico.")
                caption_run.font.name = 'Helvetica'
                caption_run.font.size = Pt(10)
                caption_run.font.italic = True
                caption_run.font.color.rgb = RGBColor(107, 114, 128)
    
    # Recomendación final
    add_custom_heading(doc, "Recomendación Ejecutiva", level=2)
    
    if not df_cb.empty:
        puntos_positivos = sum([
            datos_cb['relacion_costo_beneficio_dolares'] >= 1,
            datos_cb['payback_periodo_anios'] is not None and datos_cb['payback_periodo_anios'] <= 5,
            datos_cb['relacion_cb_unidades'] >= 1
        ])
        
        if puntos_positivos >= 2:
            recomendacion = ("APROBAR EL PROYECTO. El análisis integral indica que el Plan Estratégico es "
                           "viable financieramente y alineado con la posición competitiva de la empresa. "
                           "Se recomienda iniciar la implementación priorizando las estrategias de alta importancia.")
        elif puntos_positivos == 1:
            recomendacion = ("EVALUAR CON PRECAUCIÓN. El proyecto presenta riesgos moderados. "
                           "Se sugiere revisar los supuestos críticos y desarrollar planes de contingencia "
                           "antes de aprobar la inversión total.")
        else:
            recomendacion = ("RECHAZAR O REPLANTEAR. El análisis revela riesgos significativos. "
                           "Se recomienda revisar la estrategia, reducir el alcance inicial o buscar "
                           "alternativas de menor inversión.")
        
        add_formatted_paragraph(doc, recomendacion, bold=True)
    
    add_page_break(doc)
    
    # ============================================
    # PARTE 2: PLAN ESTRATÉGICO
    # ============================================
    
    # Portada interna
    doc.add_paragraph()
    doc.add_paragraph()
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("PLAN ESTRATÉGICO")
    title_run.font.name = 'Helvetica'
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = COLOR_PRIMARY
    
    doc.add_paragraph()
    empresa_para = doc.add_paragraph()
    empresa_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    empresa_run = empresa_para.add_run(empresa.get('nombre', 'EMPRESA').upper())
    empresa_run.font.name = 'Helvetica'
    empresa_run.font.size = Pt(18)
    empresa_run.font.bold = True
    empresa_run.font.color.rgb = COLOR_PRIMARY
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Logo grande si existe
    if logo_bytes_data:
        try:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(BytesIO(logo_bytes), width=Inches(2))
        except:
            pass
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Metadatos
    meta_data = [
        f"Versión: {version}",
        f"Fecha: {fecha_actual}",
        f"Elaborado por: {elaborado}",
        f"Revisado por: {revisado}",
        f"Aprobado por: {aprobado}"
    ]
    
    for meta in meta_data:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(meta)
        run.font.name = 'Helvetica'
        run.font.size = Pt(11)
        run.font.color.rgb = COLOR_TEXT
    
    add_page_break(doc)
    
    # 1. INTRODUCCIÓN Y FUNDAMENTOS
    add_custom_heading(doc, "1. INTRODUCCIÓN Y FUNDAMENTOS", level=1)
    
    add_custom_heading(doc, "1.1 Datos Generales de la Empresa", level=2)
    
    headers = ['Campo', 'Información']
    data = [
        ['Nombre de la Empresa', empresa.get('nombre', 'N/A')],
        ['Giro del Negocio', empresa.get('giro', 'N/A')],
        ['Fecha del Plan', fecha_actual],
        ['Versión del Documento', version],
    ]
    create_professional_table(doc, headers, data, col_widths=[5, 7])
    
    # Logo
    if logo_bytes_data:
        try:
            add_custom_heading(doc, "Logo de la Empresa", level=3)
            p = doc.add_paragraph()
            run = p.add_run()
            run.add_picture(BytesIO(logo_bytes), width=Inches(2))
        except:
            pass
    
    # 1.2 Cultura Organizacional
    add_custom_heading(doc, "1.2 Elementos Orientadores de la Cultura Organizacional", level=2)
    
    if empresa.get('mision'):
        add_custom_heading(doc, "Misión", level=3)
        add_formatted_paragraph(doc, empresa['mision'])
    
    if empresa.get('vision'):
        add_custom_heading(doc, "Visión", level=3)
        add_formatted_paragraph(doc, empresa['vision'])
    
    if empresa.get('valores'):
        add_custom_heading(doc, "Valores y Principios", level=3)
        valores = empresa['valores']
        if len(valores) > 200:
            lineas_valores = [v.strip() for v in valores.replace('\n', ',').split(',') if v.strip()]
            for val in lineas_valores[:5]:
                add_bullet_point(doc, val)
        else:
            add_formatted_paragraph(doc, valores)
    
    if empresa.get('objetivo_plan'):
        add_custom_heading(doc, "Objetivo del Plan Estratégico", level=3)
        add_formatted_paragraph(doc, empresa['objetivo_plan'])
    
    # Organigrama
    organigrama_bytes = None
    org_bytes_data = empresa.get('organigrama')
    if org_bytes_data:
        try:
            if isinstance(org_bytes_data, str):
                import base64
                organigrama_bytes = base64.b64decode(org_bytes_data)
            else:
                organigrama_bytes = org_bytes_data
            
            from PIL import Image as PILImage
            test_img = PILImage.open(BytesIO(organigrama_bytes))
            test_img.verify()
            
            add_custom_heading(doc, "1.3 Organigrama de la Empresa", level=2)
            p = doc.add_paragraph()
            run = p.add_run()
            run.add_picture(BytesIO(organigrama_bytes), width=Inches(6))
            
            caption = doc.add_paragraph()
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_run = caption.add_run("Figura 1.1. Estructura organizacional de la empresa.")
            caption_run.font.name = 'Helvetica'
            caption_run.font.size = Pt(10)
            caption_run.font.italic = True
            caption_run.font.color.rgb = RGBColor(107, 114, 128)
        except Exception as e:
            add_custom_heading(doc, "1.3 Organigrama de la Empresa", level=2)
            add_formatted_paragraph(doc, "[El organigrama no pudo ser cargado - formato de imagen no válido]")
    
    add_page_break(doc)
    
    # 2. ANÁLISIS SITUACIONAL
    add_custom_heading(doc, "2. ANÁLISIS SITUACIONAL", level=1)
    add_formatted_paragraph(doc, 
        "El análisis situacional examina tanto los factores internos como externos que afectan "
        "a la organización, permitiendo identificar fortalezas, debilidades, oportunidades y amenazas.")
    
    # 2.1 Diagnóstico Interno
    add_custom_heading(doc, "2.1 Diagnóstico Interno", level=2)
    
    if not df_made.empty:
        add_custom_heading(doc, "2.1.1 Análisis de Marketing Interno (MADE)", level=3)
        add_formatted_paragraph(doc, 
            "La matriz MADE evalúa las variables internas de marketing: Producto, Precio, Plaza y Promoción.")
        
        # Tabla resumen MADE
        headers = ['Variable', 'Factor', 'Rating', 'Ponderación']
        data = []
        for _, row in df_made.head(10).iterrows():
            data.append([
                str(row.get('variable', ''))[:20],
                str(row.get('factor', ''))[:30],
                str(row.get('rating', '')),
                f"{row.get('weight_percent', '')}%"
            ])
        create_professional_table(doc, headers, data, col_widths=[3, 6, 2, 2.5])
        
        # Análisis de IA
        analisis_made = empresa.get('analisis_made', '')
        if analisis_made:
            add_custom_heading(doc, "Análisis de Marketing Interno", level=3)
            # Parsear y agregar contenido estructurado
            for linea in str(analisis_made).split('\n'):
                if linea.strip():
                    if linea.strip().startswith('-') or linea.strip().startswith('•'):
                        add_bullet_point(doc, linea.strip()[1:].strip())
                    elif linea[0].isdigit() and '.' in linea[:3]:
                        add_custom_heading(doc, linea.strip(), level=4)
                    else:
                        add_formatted_paragraph(doc, linea.strip())
    
    # 2.2 Diagnóstico Externo
    add_custom_heading(doc, "2.2 Diagnóstico Externo", level=2)
    
    if not df_madi.empty:
        add_custom_heading(doc, "2.2.1 Análisis de Marketing Externo (MADI)", level=3)
        add_formatted_paragraph(doc, 
            "La matriz MADI evalúa las variables externas de marketing que impactan en la posición competitiva.")
        
        headers = ['Variable', 'Factor', 'Rating', 'Ponderación']
        data = []
        for _, row in df_madi.head(10).iterrows():
            data.append([
                str(row.get('variable', ''))[:20],
                str(row.get('factor', ''))[:30],
                str(row.get('rating', '')),
                f"{row.get('weight_percent', '')}%"
            ])
        create_professional_table(doc, headers, data, col_widths=[3, 6, 2, 2.5])
        
        analisis_madi = empresa.get('analisis_madi', '')
        if analisis_madi:
            add_custom_heading(doc, "Análisis de Marketing Externo", level=3)
            for linea in str(analisis_madi).split('\n'):
                if linea.strip():
                    if linea.strip().startswith('-') or linea.strip().startswith('•'):
                        add_bullet_point(doc, linea.strip()[1:].strip())
                    elif linea[0].isdigit() and '.' in linea[:3]:
                        add_custom_heading(doc, linea.strip(), level=4)
                    else:
                        add_formatted_paragraph(doc, linea.strip())
    
    # Análisis PEST
    if not df_pest.empty:
        add_custom_heading(doc, "2.2.2 Análisis del Entorno PEST", level=3)
        add_formatted_paragraph(doc, 
            "El análisis PEST examina los factores Políticos, Económicos, Sociales y Tecnológicos.")
        
        headers = ['Categoría', 'Factor', 'Tipo FODA', 'Puntaje', 'Ponderado']
        data = []
        for _, row in df_pest.head(15).iterrows():
            factor_text = str(row.get('factor', ''))
            data.append([
                str(row.get('categoria', '')),
                (factor_text[:35] + '...') if len(factor_text) > 35 else factor_text,
                str(row.get('tipo_foda', '')),
                str(row.get('puntaje', '')),
                f"{row.get('valor_ponderado', 0):.2f}"
            ])
        create_professional_table(doc, headers, data, col_widths=[2.5, 6, 2, 1.5, 2])
        
        # Gráfico PEST
        grafico_pest = generar_grafico_barras_pest_mejorado(df_pest)
        if grafico_pest:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(grafico_pest, width=Inches(5))
            
            caption = doc.add_paragraph()
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_run = caption.add_run("Figura 2.1. Análisis PEST - Distribución por categoría.")
            caption_run.font.name = 'Helvetica'
            caption_run.font.size = Pt(10)
            caption_run.font.italic = True
            caption_run.font.color.rgb = RGBColor(107, 114, 128)
        
        analisis_pest = empresa.get('analisis_pest', '')
        if analisis_pest:
            add_custom_heading(doc, "Interpretación del Análisis PEST", level=3)
            for linea in str(analisis_pest).split('\n'):
                if linea.strip():
                    if linea.strip().startswith('-') or linea.strip().startswith('•'):
                        add_bullet_point(doc, linea.strip()[1:].strip())
                    elif linea[0].isdigit() and '.' in linea[:3]:
                        add_custom_heading(doc, linea.strip(), level=4)
                    else:
                        add_formatted_paragraph(doc, linea.strip())
    
    # Matriz FODA Cruzado
    if not df_foda.empty:
        add_custom_heading(doc, "2.3 Matriz FODA Cruzado", level=2)
        add_formatted_paragraph(doc, 
            "El análisis FODA cruzado identifica estrategias a partir de la combinación de "
            "fortalezas, debilidades, oportunidades y amenazas.")
        
        headers = ['Cuadrante', 'Factor Fila', 'Factor Columna', 'Impacto']
        data = []
        for _, row in df_foda.head(20).iterrows():
            data.append([
                str(row.get('cuadrante', '')),
                str(row.get('factor_fila', ''))[:30] + ('...' if len(str(row.get('factor_fila', ''))) > 30 else ''),
                str(row.get('factor_columna', ''))[:30] + ('...' if len(str(row.get('factor_columna', ''))) > 30 else ''),
                str(row.get('impacto', ''))
            ])
        create_professional_table(doc, headers, data, col_widths=[2.5, 5, 5, 1.5])
        
        # Gráfico FODA
        if puntajes_foda is not None and not puntajes_foda.empty:
            grafico_foda = generar_grafico_foda_radar_mejorado(puntajes_foda)
            if grafico_foda:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(grafico_foda, width=Inches(4))
                
                caption = doc.add_paragraph()
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption_run = caption.add_run("Figura 2.2. Posicionamiento estratégico según FODA cruzado.")
                caption_run.font.name = 'Helvetica'
                caption_run.font.size = Pt(10)
                caption_run.font.italic = True
                caption_run.font.color.rgb = RGBColor(107, 114, 128)
        
        # Postura estratégica
        if analisis_foda_df is not None:
            add_custom_heading(doc, "Postura Estratégica Recomendada", level=3)
            postura_text = (f"Basado en el análisis cruzado, la estrategia principal recomendada es "
                          f"{estrategia_principal}. La distribución de puntajes por estrategia es:")
            add_formatted_paragraph(doc, postura_text, bold=True)
            
            headers = ['Estrategia', 'Puntaje Total']
            data = []
            for _, row in analisis_foda_df.iterrows():
                data.append([str(row['Estrategia']), str(row['Puntaje Total'])])
            create_professional_table(doc, headers, data, col_widths=[7.5, 2.5])
        
        analisis_foda_texto = empresa.get('analisis_foda', '')
        if analisis_foda_texto:
            add_custom_heading(doc, "Interpretación del Análisis FODA", level=3)
            for linea in str(analisis_foda_texto).split('\n'):
                if linea.strip():
                    if linea.strip().startswith('-') or linea.strip().startswith('•'):
                        add_bullet_point(doc, linea.strip()[1:].strip())
                    elif linea[0].isdigit() and '.' in linea[:3]:
                        add_custom_heading(doc, linea.strip(), level=4)
                    else:
                        add_formatted_paragraph(doc, linea.strip())
    
    add_page_break(doc)
    
    # 3. ESTRATEGIAS
    add_custom_heading(doc, "3. ESTRATEGIAS", level=1)
    add_formatted_paragraph(doc, 
        "Las estrategias representan las acciones específicas diseñadas para alcanzar los objetivos "
        "organizacionales, derivadas del análisis situacional.")
    
    if not df_estrategias.empty:
        estrategia_intro = (f"Se han formulado {len(df_estrategias)} estrategias distribuidas en los cuatro cuadrantes "
                          f"del análisis FODA cruzado. La estrategia principal recomendada es {estrategia_principal}.")
        add_formatted_paragraph(doc, estrategia_intro)
        
        add_custom_heading(doc, "3.1 Resumen de Estrategias por Cuadrante", level=2)
        
        resumen_est = df_estrategias.groupby('cuadrante').agg({
            'estrategia': 'count',
            'plan_asignado': lambda x: ', '.join(x.unique())
        }).reset_index()
        resumen_est.columns = ['Cuadrante', 'N° Estrategias', 'Planes Asignados']
        
        headers = ['Cuadrante', 'N° Estrategias', 'Planes Asignados']
        data = []
        for _, row in resumen_est.iterrows():
            data.append([row['Cuadrante'], str(row['N° Estrategias']), row['Planes Asignados']])
        create_professional_table(doc, headers, data, col_widths=[3, 3, 6])
        
        add_custom_heading(doc, "3.2 Estrategias Detalladas", level=2)
        
        for idx, row in df_estrategias.iterrows():
            add_custom_heading(doc, f"Estrategia {idx + 1}: [{row['cuadrante']}] {row['estrategia']}", level=3)
            add_formatted_paragraph(doc, f"Plan asignado: {row['plan_asignado']}", bold=False)
            add_formatted_paragraph(doc, f"Importancia: {row['importancia']}", bold=False)
            add_formatted_paragraph(doc, f"Actividades clave: {row['actividades']}")
    else:
        add_formatted_paragraph(doc, "No se han generado estrategias en el sistema.")
    
    add_page_break(doc)
    
    # 4. PLAN DE ACCIÓN
    add_custom_heading(doc, "4. PLAN DE ACCIÓN", level=1)
    add_formatted_paragraph(doc, 
        "El plan de acción detalla las actividades específicas, responsables, tiempos y recursos "
        "necesarios para implementar las estrategias formuladas.")
    
    # Planes Maestros
    planes_maestros = empresa.get('analisis_operativo', '')
    if planes_maestros and str(planes_maestros).strip():
        add_custom_heading(doc, "4.1 Planes Funcionales Detallados", level=2)
        
        # Intentar dividir por secciones si están marcadas
        import re
        secciones_planes = re.split(r'\n(?=\d+\.\s+PLAN\s+)', str(planes_maestros))
        
        if len(secciones_planes) > 1:
            for seccion in secciones_planes:
                if seccion.strip():
                    for linea in seccion.split('\n'):
                        if linea.strip():
                            if linea.strip().startswith('==='):
                                continue
                            elif linea[0].isdigit() and '.' in linea[:5]:
                                add_custom_heading(doc, linea.strip(), level=3)
                            elif linea.strip().startswith('-') or linea.strip().startswith('•'):
                                add_bullet_point(doc, linea.strip()[1:].strip())
                            else:
                                add_formatted_paragraph(doc, linea.strip())
        else:
            for linea in str(planes_maestros).split('\n'):
                if linea.strip():
                    if linea[0].isdigit() and '.' in linea[:5]:
                        add_custom_heading(doc, linea.strip(), level=3)
                    elif linea.strip().startswith('-') or linea.strip().startswith('•'):
                        add_bullet_point(doc, linea.strip()[1:].strip())
                    else:
                        add_formatted_paragraph(doc, linea.strip())
        
        add_page_break(doc)
    
    # Operativización
    add_custom_heading(doc, "4.2 Operativización y Presupuesto", level=2)
    
    if not df_oper.empty:
        oper_intro = (f"La operativización detalla {len(df_oper)} actividades derivadas de las estrategias formuladas, "
                     f"con una inversión total estimada de ${total_costo:,.2f}.")
        add_formatted_paragraph(doc, oper_intro, bold=True)
        
        add_custom_heading(doc, "4.2.1 Resumen de Inversión por Plan", level=3)
        
        presupuesto_plan = df_oper.groupby('plan_asignado').agg({
            'costo': 'sum',
            'descripcion_actividad': 'count'
        }).reset_index()
        presupuesto_plan.columns = ['Plan', 'Costo Total', 'N° Actividades']
        
        headers = ['Plan Estratégico', 'Costo Total', 'N° Actividades']
        data = []
        for _, row in presupuesto_plan.iterrows():
            data.append([row['Plan'], f"${row['Costo Total']:,.2f}", str(int(row['N° Actividades']))])
        create_professional_table(doc, headers, data, col_widths=[7.5, 3, 2.5])
        
        add_custom_heading(doc, "4.2.2 Cuadro de Operativización Detallado", level=3)
        
        headers = ['Estrategia', 'Actividad', 'Plazo', 'Responsable', 'Costo']
        data = []
        for _, row in df_oper.head(30).iterrows():
            data.append([
                (row['estrategia_nombre'][:25] + '...') if len(row['estrategia_nombre']) > 25 else row['estrategia_nombre'],
                (row['descripcion_actividad'][:35] + '...') if len(row['descripcion_actividad']) > 35 else row['descripcion_actividad'],
                row['plazo'] or 'Pendiente',
                row['responsable'] or 'Sin asignar',
                f"${row['costo']:,.2f}"
            ])
        create_professional_table(doc, headers, data, col_widths=[4.5, 5.5, 2, 2.5, 2])
        
        if len(df_oper) > 30:
            add_formatted_paragraph(doc, 
                f"Nota: Se muestran 30 de {len(df_oper)} actividades. El detalle completo está en los anexos.",
                italic=True)
        
        add_custom_heading(doc, "4.2.3 Actividades de Mayor Inversión", level=3)
        
        actividades_criticas = df_oper.nlargest(5, 'costo')
        for _, row in actividades_criticas.iterrows():
            add_bullet_point(doc, 
                f"${row['costo']:,.2f} - {row['descripcion_actividad'][:80]}... ({row['responsable'] or 'Sin asignar'})")
    else:
        add_formatted_paragraph(doc, "No se ha completado la operativización de estrategias.")
    
    add_page_break(doc)
    
    # 5. EVALUACIÓN Y CONTROL
    add_custom_heading(doc, "5. EVALUACIÓN Y CONTROL", level=1)
    add_formatted_paragraph(doc, 
        "La evaluación y control permiten monitorear el desempeño de las estrategias mediante "
        "indicadores clave de desempeño (KPIs) y sistemas de alerta temprana.")
    
    add_custom_heading(doc, "5.1 Cuadro de Mando Integral (CMI)", level=2)
    
    if not df_estrategias.empty:
        add_formatted_paragraph(doc, 
            "El Cuadro de Mando Integral traduce las estrategias en indicadores medibles "
            "desde cuatro perspectivas: Financiera, Cliente, Procesos Internos, y Aprendizaje y Crecimiento.")
        
        # Intentar obtener CMI guardado
        cmi_guardado = empresa.get('analisis_cmi', '')
        df_cmi = pd.DataFrame()
        
        if cmi_guardado and '|' in str(cmi_guardado):
            try:
                df_cmi = pd.read_csv(io.StringIO(cmi_guardado), sep="|")
            except:
                pass
        
        if not df_cmi.empty:
            add_custom_heading(doc, "5.1.1 Indicadores Clave por Perspectiva", level=3)
            
            headers = ['Estrategia', 'Perspectiva', 'KPI', 'Frecuencia', 'Límites']
            data = []
            for _, row in df_cmi.head(15).iterrows():
                limites = f"LI:{row.get('LI','')} LC:{row.get('LC','')} LS:{row.get('LS','')}"
                data.append([
                    (row['Estrategia'][:35] + '...') if len(row['Estrategia']) > 35 else row['Estrategia'],
                    row['Perspectiva'],
                    (row['KPIs'][:30] + '...') if len(row['KPIs']) > 30 else row['KPIs'],
                    row['Frecuencia'],
                    limites
                ])
            create_professional_table(doc, headers, data, col_widths=[5, 3, 4.5, 2, 3])
            
            if len(df_cmi) > 15:
                add_formatted_paragraph(doc, 
                    f"Nota: Se muestran 15 de {len(df_cmi)} indicadores. El detalle completo está en los anexos.",
                    italic=True)
    
    add_custom_heading(doc, "5.2 Semaforización Estratégica", level=2)
    
    analisis_semaforo = empresa.get('analisis_semaforo_resumen', '')
    if analisis_semaforo:
        add_formatted_paragraph(doc, analisis_semaforo)
    else:
        add_formatted_paragraph(doc, 
            "El sistema de semaforización evalúa el estado de cada estrategia considerando alineación "
            "con objetivos, recursos asignados y contexto externo. Las estrategias se clasifican en: "
            "Óptimas (implementar según plan), Atención (requieren ajustes), "
            "y Críticas (necesitan revisión inmediata).")
    
    add_page_break(doc)
    
    # ============================================
    # PARTE 3: ANEXOS
    # ============================================
    
    add_custom_heading(doc, "ANEXOS", level=1)
    
    # Anexo A: Análisis Detallados
    add_custom_heading(doc, "Anexo A. Análisis Detallados de Matrices", level=2)
    
    if not df_made.empty:
        add_custom_heading(doc, "A.1 Matriz MADE (Marketing Interno) - Datos Completos", level=3)
        
        headers = ['Variable', 'Factor', 'Producto', 'Precio', 'Plaza', 'Promoción', 'Rating', 'Peso %', 'Valor']
        data = []
        for _, row in df_made.iterrows():
            data.append([
                str(row.get('variable', ''))[:15],
                str(row.get('factor', ''))[:25],
                str(row.get('producto', ''))[:10],
                str(row.get('precio', ''))[:10],
                str(row.get('plaza', ''))[:10],
                str(row.get('promocion', ''))[:10],
                str(row.get('rating', '')),
                f"{row.get('weight_percent', '')}%",
                str(row.get('valor', ''))
            ])
        create_professional_table(doc, headers, data, col_widths=[2.2, 3.8, 1.8, 1.8, 1.8, 1.8, 1.5, 1.8, 1.5])
        add_page_break(doc)
    
    if not df_madi.empty:
        add_custom_heading(doc, "A.2 Matriz MADI (Marketing Externo) - Datos Completos", level=3)
        
        headers = ['Variable', 'Factor', 'Producto', 'Precio', 'Plaza', 'Promoción', 'Rating', 'Peso %', 'Valor']
        data = []
        for _, row in df_madi.iterrows():
            data.append([
                str(row.get('variable', ''))[:15],
                str(row.get('factor', ''))[:25],
                str(row.get('producto', ''))[:10],
                str(row.get('precio', ''))[:10],
                str(row.get('plaza', ''))[:10],
                str(row.get('promocion', ''))[:10],
                str(row.get('rating', '')),
                f"{row.get('weight_percent', '')}%",
                str(row.get('valor', ''))
            ])
        create_professional_table(doc, headers, data, col_widths=[2.2, 3.8, 1.8, 1.8, 1.8, 1.8, 1.5, 1.8, 1.5])
        add_page_break(doc)
    
    if not df_pest.empty:
        add_custom_heading(doc, "A.3 Análisis PEST Completo", level=3)
        
        headers = ['Categoría', 'Factor', 'Tipo FODA', 'Puntaje', 'Importancia', 'Ponderado']
        data = []
        for _, row in df_pest.iterrows():
            factor_text = str(row.get('factor', ''))
            data.append([
                row['categoria'],
                (factor_text[:40] + '...') if len(factor_text) > 40 else factor_text,
                str(row.get('tipo_foda', '')),
                str(row.get('puntaje', '')),
                f"{row['importancia']}%",
                f"{row['valor_ponderado']:.2f}"
            ])
        create_professional_table(doc, headers, data, col_widths=[2.5, 7.5, 2.5, 2, 2.5, 2.5])
        
        if grafico_pest:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(grafico_pest, width=Inches(6))
            
            caption = doc.add_paragraph()
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_run = caption.add_run("Figura A.1. Análisis PEST - Distribución de factores por categoría.")
            caption_run.font.name = 'Helvetica'
            caption_run.font.size = Pt(10)
            caption_run.font.italic = True
            caption_run.font.color.rgb = RGBColor(107, 114, 128)
        
        add_page_break(doc)
    
    if not df_foda.empty:
        add_custom_heading(doc, "A.4 Matriz FODA Cruzado Completa", level=3)
        
        headers = ['Cuadrante', 'Factor Fila', 'Factor Columna', 'Impacto']
        data = []
        for _, row in df_foda.iterrows():
            data.append([
                row['cuadrante'],
                str(row['factor_fila']),
                str(row['factor_columna']),
                str(row['impacto'])
            ])
        create_professional_table(doc, headers, data, col_widths=[2.5, 6.25, 6.25, 2])
        
        if puntajes_foda is not None and not puntajes_foda.empty:
            grafico_foda = generar_grafico_foda_radar_mejorado(puntajes_foda)
            if grafico_foda:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(grafico_foda, width=Inches(5))
                
                caption = doc.add_paragraph()
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption_run = caption.add_run("Figura A.2. Posicionamiento estratégico - Matriz FODA cruzada.")
                caption_run.font.name = 'Helvetica'
                caption_run.font.size = Pt(10)
                caption_run.font.italic = True
                caption_run.font.color.rgb = RGBColor(107, 114, 128)
        
        add_page_break(doc)
    
    # Anexo B: Dashboards
    add_custom_heading(doc, "Anexo B. Dashboards de Análisis Estratégico", level=2)
    add_formatted_paragraph(doc, "Este anexo presenta las visualizaciones completas del análisis estratégico.")
    
    if not df_estrategias.empty:
        add_custom_heading(doc, "B.1 Distribución de Estrategias por Cuadrante", level=3)
        
        # Crear gráfico de barras simple
        import matplotlib.pyplot as plt
        import numpy as np
        
        fig, ax = plt.subplots(figsize=(7, 4))
        est_counts = df_estrategias['cuadrante'].value_counts()
        colors_est = {'FO': '#2ca02c', 'FA': '#d62728', 'DO': '#ff7f0e', 'DA': '#9467bd'}
        bars = ax.bar(est_counts.index, est_counts.values, 
                     color=[colors_est.get(x, '#1f77b4') for x in est_counts.index],
                     edgecolor='black', linewidth=0.5)
        ax.set_title('Distribución de Estrategias por Cuadrante FODA', fontsize=11, fontweight='bold')
        ax.set_ylabel('Número de Estrategias')
        
        for bar in bars:
            height = bar.get_height()
            ax.text(bar.get_x() + bar.get_width()/2., height, f'{int(height)}', 
                   ha='center', va='bottom')
        
        plt.tight_layout()
        buf_est = BytesIO()
        plt.savefig(buf_est, format='PNG', dpi=150, bbox_inches='tight', facecolor='white')
        plt.close()
        buf_est.seek(0)
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(buf_est, width=Inches(5))
    
    if not df_proy.empty:
        add_custom_heading(doc, "B.2 Proyección Financiera Detallada", level=3)
        
        grafico_proy = generar_grafico_proyeccion_mejorado(df_proy)
        if grafico_proy:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(grafico_proy, width=Inches(6))
        
        headers = ['Año', 'Ingresos Proyectados', 'Costos Proyectados', 'Utilidad Neta Proyectada']
        data = []
        for _, row in df_proy.iterrows():
            data.append([
                str(int(row['anio'])),
                f"${row['ingresos_proyectados']:,.0f}",
                f"${row['costos_proyectados']:,.0f}",
                f"${row['utilidad_neta_proyectada']:,.0f}"
            ])
        create_professional_table(doc, headers, data, col_widths=[2.5, 4.5, 4.5, 4.5])
        add_page_break(doc)
    
    if not df_cb.empty:
        add_custom_heading(doc, "B.3 Análisis Costo-Beneficio Detallado", level=3)
        
        datos_cb = df_cb.iloc[0]
        
        headers = ['Indicador', 'Valor', 'Umbral Aceptable', 'Estado']
        data = [
            ['Relación C-B ($)', f"{float(datos_cb['relacion_costo_beneficio_dolares']):.2f}", '≥ 1.0', 
             'Aceptable' if datos_cb['relacion_costo_beneficio_dolares'] >= 1 else 'No aceptable'],
            ['Payback (años)', 
             f"{float(datos_cb['payback_periodo_anios']):.1f}" if datos_cb['payback_periodo_anios'] else 'N/A',
             '≤ 5', 
             'Aceptable' if datos_cb['payback_periodo_anios'] and datos_cb['payback_periodo_anios'] <= 5 else 'Revisar'],
            ['VPN Total', f"${float(datos_cb['vpn_total']):,.2f}", '> 0', 
             'Positivo' if datos_cb['vpn_total'] > 0 else 'Negativo'],
            ['Inversión Total', f"${float(datos_cb['inversion_total']):,.2f}", '-', 'Requerimiento'],
            ['Beneficio/Unidad', f"${float(datos_cb['beneficio_por_unidad']):,.2f}", '≥ Costo/Unidad', 
             'Rentable' if datos_cb['relacion_cb_unidades'] >= 1 else 'No rentable'],
        ]
        create_professional_table(doc, headers, data, col_widths=[4.5, 3.75, 3.75, 3.75])
        add_page_break(doc)
    
    # Anexo C: Operativización Completa
    if not df_oper.empty:
        add_custom_heading(doc, "Anexo C. Cuadro de Operativización Completo", level=2)
        
        # Dividir en chunks para no sobrecargar
        chunk_size = 40
        datos_oper_full = [['N°', 'Estrategia', 'Actividad', 'Plazo', 'Responsable', 'Costo']]
        
        for idx, row in df_oper.iterrows():
            datos_oper_full.append([
                str(idx + 1),
                (row['estrategia_nombre'][:30] + '...') if len(row['estrategia_nombre']) > 30 else row['estrategia_nombre'],
                (row['descripcion_actividad'][:40] + '...') if len(row['descripcion_actividad']) > 40 else row['descripcion_actividad'],
                row['plazo'] or 'Pendiente',
                row['responsable'] or 'Sin asignar',
                f"${row['costo']:,.2f}"
            ])
        
        for i in range(0, len(datos_oper_full), chunk_size):
            chunk = datos_oper_full[i:i+chunk_size]
            if i == 0:
                create_professional_table(doc, chunk[0], chunk[1:], col_widths=[1.25, 5, 6.25, 2, 3, 2.5])
            else:
                # Continuación sin encabezado
                table = doc.add_table(rows=len(chunk), cols=len(chunk[0]))
                table.style = 'Table Grid'
                for row_idx, row_data in enumerate(chunk):
                    for col_idx, cell_text in enumerate(row_data):
                        cell = table.rows[row_idx].cells[col_idx]
                        cell.text = str(cell_text)
                        # Formato
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.name = 'Helvetica'
                                run.font.size = Pt(9)
            if i + chunk_size < len(datos_oper_full):
                add_page_break(doc)
    
    # Anexo D: CMI Completo
    if not df_cmi.empty:
        add_page_break(doc)
        add_custom_heading(doc, "Anexo D. Cuadro de Mando Integral Completo", level=2)
        
        headers = ['Estrategia', 'Perspectiva', 'KPIs', 'Fórmulas', 'Frecuencia', 'LI', 'LC', 'LS']
        data = []
        for _, row in df_cmi.iterrows():
            data.append([
                (row['Estrategia'][:40] + '...') if len(row['Estrategia']) > 40 else row['Estrategia'],
                row['Perspectiva'],
                (row['KPIs'][:35] + '...') if len(row['KPIs']) > 35 else row['KPIs'],
                (row['Formulas'][:25] + '...') if len(row['Formulas']) > 25 else row['Formulas'],
                row['Frecuencia'],
                str(row['LI']),
                str(row['LC']),
                str(row['LS'])
            ])
        create_professional_table(doc, headers, data, col_widths=[4.5, 2.5, 4.5, 3, 2, 1.5, 1.5, 1.5])
    
    # Guardar en buffer
    word_buffer = BytesIO()
    doc.save(word_buffer)
    word_buffer.seek(0)
    return word_buffer

def mostrar_ultimo_analisis_guardado(empresa_data, tipo_analisis):
    """
    Muestra el último análisis guardado desde el diccionario de datos de la empresa.
    """
    contenido = empresa_data.get(f'analisis_{tipo_analisis}')
    if contenido and str(contenido).strip():
        st.markdown("---")
        st.markdown(f"**📄 Último análisis de {tipo_analisis.upper()} guardado:**")
        if tipo_analisis == 'cmi' and '|' in str(contenido):
            try:
                df_view = pd.read_csv(io.StringIO(contenido), sep="|")
                st.table(df_view)
            except Exception as e:
                st.text_area(f"contenido_guardado_{tipo_analisis}", value=contenido, height=200, disabled=True, label_visibility="collapsed")
        else:
            st.text_area(f"contenido_guardado_{tipo_analisis}", value=contenido, height=200, disabled=True, label_visibility="collapsed")

def aplicacion_principal():
    with st.sidebar:
        st.header("Estratega Pro UG-UCE")
        empresas_df = get_empresas()

        # Mostrar empresas con indicador de permiso
        opciones_empresas = []
        empresa_info = {}

        if not empresas_df.empty:
            for _, row in empresas_df.iterrows():
                es_propietario = row.get('es_propietario', False)
                permiso = row.get('permiso', 'lector')

                if es_propietario:
                    label = f"👑 {row['nombre']} (Propietario)"
                elif permiso == 'editor':
                    label = f"✏️ {row['nombre']} (Editor)"
                else:
                    label = f"👁️ {row['nombre']} (Lector)"

                opciones_empresas.append(label)
                empresa_info[label] = {
                    'id': row['id'],
                    'nombre': row['nombre'],
                    'es_propietario': es_propietario,
                    'permiso': permiso
                }

        empresa_seleccionada_label = st.selectbox("Selecciona una Empresa", opciones_empresas, index=None, placeholder="Elige una opción")

        empresa_id = None
        empresa_seleccionada = None
        es_propietario = False
        permiso_actual = None

        if empresa_seleccionada_label and empresa_seleccionada_label in empresa_info:
            info = empresa_info[empresa_seleccionada_label]
            empresa_id = info['id']
            empresa_seleccionada = info['nombre']
            es_propietario = info['es_propietario']
            permiso_actual = info['permiso']

            # Mostrar badge de permiso
            if es_propietario:
                st.success("👑 Eres el propietario de esta empresa")
            elif permiso_actual == 'editor':
                st.info("✏️ Tienes permiso de Editor")
            else:
                st.info("👁️ Tienes permiso de Lector (solo ver)")

        st.divider()

        # Sección de crear empresa (solo para propietarios)
        with st.expander("➕ Crear Nueva Empresa"):
            with st.form("new_empresa_form"):
                new_empresa_name = st.text_input("Nombre de la nueva empresa")
                if st.form_submit_button("Crear"):
                    if new_empresa_name and supabase and st.session_state.get("user"):
                        try:
                            user_id = st.session_state.user.id
                            supabase.table('empresas').insert({
                                "nombre": new_empresa_name,
                                "propietario_id": user_id
                            }).execute()
                            st.success(f"Empresa '{new_empresa_name}' creada en la nube.")
                            st.rerun()
                        except Exception as e:
                            st.error(f"Error al crear la empresa: {e}")
                    else:
                        st.warning("El nombre no puede estar vacío.")

        # Opciones solo para propietarios
        if empresa_id and es_propietario:
            st.divider()

            # Compartir empresa
            with st.expander("🔗 Compartir Empresa"):
                st.write("Comparte esta empresa con otros usuarios:")

                with st.form("form_compartir"):
                    email_compartir = st.text_input("Email del usuario")
                    permiso_compartir = st.selectbox("Permiso", ["lector", "editor"], 
                                                    format_func=lambda x: "👁️ Lector (solo ver)" if x == "lector" else "✏️ Editor (puede modificar)")

                    if st.form_submit_button("Compartir"):
                        if email_compartir:
                            exito, mensaje = compartir_empresa(empresa_id, email_compartir, permiso_compartir)
                            if exito:
                                st.success(mensaje)
                                st.rerun()
                            else:
                                st.error(mensaje)
                        else:
                            st.warning("Ingresa un email")

                # Mostrar usuarios con quienes se compartió
                usuarios_comp = get_usuarios_compartidos(empresa_id)
                
                if usuarios_comp:
                    st.write("---")
                    st.markdown("**👥 Usuarios con acceso compartido**")
                    
                    for uc in usuarios_comp:
                        with st.container():
                            # Fila principal: Email, Rol, Eliminar
                            col1, col2, col3 = st.columns([4, 2, 1])
                            
                            with col1:
                                icono = "✏️" if uc['permiso'] == 'editor' else "👁️"
                                # Mostrar email completo, no truncado
                                email_display = uc['email']
                                st.markdown(f"{icono} **{email_display}**")
                            
                            with col2:
                                # Badge del rol
                                color_badge = "#2196f3" if uc['permiso'] == 'editor' else "#757575"
                                st.markdown(
                                    f"<span style='background-color: {color_badge}; color: white; "
                                    f"padding: 4px 12px; border-radius: 12px; font-size: 11px;'>"
                                    f"{uc['permiso'].upper()}</span>",
                                    unsafe_allow_html=True
                                )
                            
                            with col3:
                                user_id = uc['usuario_id']
                                if st.button("🗑️", key=f"del_{user_id}", help="Eliminar acceso"):
                                    try:
                                        supabase.table('empresas_compartidas').delete().eq(
                                            'empresa_id', empresa_id
                                        ).eq('usuario_compartido_id', user_id).execute()
                                        st.success("✓ Eliminado")
                                        st.rerun()
                                    except Exception as e:
                                        st.error(f"Error: {str(e)[:50]}")
                            
                            # Fila de cambio de rol
                            col_cambio1, col_cambio2 = st.columns([3, 2])
                            
                            with col_cambio1:
                                user_id = uc['usuario_id']
                                rol_actual = uc['permiso']
                                
                                nuevo_rol = st.selectbox(
                                    "Cambiar a:",
                                    ["lector", "editor"],
                                    index=0 if rol_actual == 'lector' else 1,
                                    key=f"rol_{user_id}",
                                    label_visibility="collapsed"
                                )
                            
                            with col_cambio2:
                                user_id = uc['usuario_id']
                                rol_actual = uc['permiso']
                                
                                if nuevo_rol != rol_actual:
                                    if st.button("💾 Guardar", key=f"save_{user_id}", type="primary"):
                                        try:
                                            supabase.table('empresas_compartidas').update({
                                                'permiso': nuevo_rol
                                                # También actualizamos el email por si acaso
                                            }).eq('empresa_id', empresa_id).eq('usuario_compartido_id', user_id).execute()
                                            st.success(f"✓ {nuevo_rol}")
                                            st.rerun()
                                        except Exception as e:
                                            st.error(f"Error: {str(e)[:50]}")
                                else:
                                    st.caption("✓ Actual")
                        
                        st.divider()
                else:
                    st.info("ℹ️ Esta empresa no está compartida con ningún usuario todavía.")
                    
            
            # Eliminar empresa (solo propietario) - CON CONFIRMACIÓN
            if "confirmar_eliminacion" not in st.session_state:
                st.session_state.confirmar_eliminacion = False
            
            if not st.session_state.confirmar_eliminacion:
                if st.button("❌ Eliminar Empresa", type="primary"):
                    st.session_state.confirmar_eliminacion = True
                    st.rerun()
            else:
                # Mostrar diálogo de confirmación
                st.error(f"⚠️ ¿ESTÁS SEGURO de eliminar '{empresa_seleccionada}'?")
                st.warning("🚨 Esta acción NO se puede deshacer. Se perderán todos los datos permanentemente.")
                
                col_confirm1, col_confirm2 = st.columns(2)
                with col_confirm1:
                    if st.button("✅ SÍ, Eliminar Definitivamente", type="primary", use_container_width=True):
                        if supabase:
                            try:
                                # Primero eliminar registros compartidos
                                supabase.table('empresas_compartidas').delete().eq('empresa_id', empresa_id).execute()
                                # Luego eliminar empresa
                                supabase.table('empresas').delete().eq('id', empresa_id).execute()
                                st.success(f"Empresa '{empresa_seleccionada}' eliminada.")
                                # Limpiar estado
                                st.session_state.confirmar_eliminacion = False
                                st.rerun()
                            except Exception as e:
                                st.error(f"Error al eliminar la empresa: {e}")
                
                with col_confirm2:
                    if st.button("❌ NO, Cancelar", type="secondary", use_container_width=True):
                        st.session_state.confirmar_eliminacion = False
                        st.info("Eliminación cancelada.")
                        st.rerun()
                        
        # Mensaje para no propietarios
        elif empresa_id and not es_propietario:
            st.divider()
            st.info("ℹ️ Solo el propietario puede compartir o eliminar esta empresa")
        
        # BOTÓN DE CERRAR SESIÓN (AL FINAL DEL SIDEBAR)
        st.divider()
        st.subheader("Sesión")
        
        # Mostrar usuario actual si existe
        if st.session_state.get("user"):
            st.caption(f"👤 {st.session_state.user.email}")
        
        if st.button("🚪 Cerrar Sesión", type="secondary", use_container_width=True):
            # Limpiar toda la sesión
            keys_to_delete = [k for k in st.session_state.keys()]
            for key in keys_to_delete:
                del st.session_state[key]
            st.success("Sesión cerrada correctamente")
            st.rerun()
                
    if not empresa_id:
        st.info("👈 Por favor, selecciona o crea una empresa en el menú lateral para comenzar.")
        st.stop()


    empresa_data = get_datos_empresa(empresa_id)
    if not empresa_data:
        st.error("No se pudieron cargar los datos de la empresa. Verifica tus permisos.")
        st.stop()

    # Verificar permisos usando la información ya obtenida en get_empresas
    es_propietario = empresa_data.get('propietario_id') == st.session_state.user.id

    # Si no es propietario, verificar si tiene acceso compartido
    if not es_propietario:
        try:
            res = supabase.table('empresas_compartidas').select('permiso').eq('empresa_id', empresa_id).eq('usuario_compartido_id', st.session_state.user.id).single().execute()
            if not res.data:
                st.error("No tienes permiso para ver esta empresa.")
                st.stop()
            permiso_usuario = res.data['permiso']
            es_editor = (permiso_usuario == 'editor')
        except:
            st.error("No tienes permiso para ver esta empresa.")
            st.stop()
    else:
        es_editor = False

    puede_editar = es_propietario or es_editor

    tab1, tab2, tab3, tab4, tab5, tab6, tab7, tab8, tab9 = st.tabs([
    "1. Introducción", 
    "2. Diagnóstico Situacional", 
    "3. Estrategia", 
    "4. Planes", 
    "5. CMI/Indicadores", 
    "6. Semaforización", 
    "7. Operativización/Presupuesto", 
    "8. Dashboard de Análisis", 
    "9. Resumen y Conclusiones"
])
    
    # --- PESTAÑA 1: INTRODUCCION ---
    with tab1:
        st.header("Introduccion y Cultura Organizacional")
        
        with st.form("form_intro"):
            st.subheader("Datos Generales")
            nombre = st.text_input("Nombre de la Empresa", empresa_data.get('nombre', ''), disabled=not puede_editar)
            giro = st.text_input("Giro del Negocio", empresa_data.get('giro', ''), disabled=not puede_editar)
            
            # LOGO
            st.divider()
            st.subheader("Logo de la Empresa")
            
            logo_actual = empresa_data.get('logo')
            if logo_actual:
                try:
                    import base64
                    logo_bytes = base64.b64decode(logo_actual)
                    st.image(logo_bytes, width=150, caption="Logo guardado")
                except Exception as e:
                    st.info("No hay logo guardado")
            else:
                st.info("No hay logo guardado. Sube uno nuevo.")
            
            logo_file = st.file_uploader("Subir Logo", type=['png', 'jpg', 'jpeg'], disabled=not puede_editar)
            
            # CULTURA ORGANIZACIONAL
            st.divider()
            st.subheader("Cultura Organizacional")
            objetivo_plan = st.text_area("Objetivo del Plan", empresa_data.get('objetivo_plan', ''), disabled=not puede_editar)
            mision = st.text_area("Mision", empresa_data.get('mision', ''), disabled=not puede_editar)
            vision = st.text_area("Vision", empresa_data.get('vision', ''), disabled=not puede_editar)
            obj_gen = st.text_area("Objetivo General", empresa_data.get('obj_general', ''), disabled=not puede_editar)
            obj_esp = st.text_area("Objetivos Especificos", empresa_data.get('obj_especificos', ''), disabled=not puede_editar)
            politicas = st.text_area("Politicas", empresa_data.get('politicas', ''), disabled=not puede_editar)
            valores = st.text_area("Valores", empresa_data.get('valores', ''), disabled=not puede_editar)
            
            # ORGANIGRAMA (AL FINAL)
            st.divider()
            st.subheader("Organigrama")
            
            org_actual = empresa_data.get('organigrama')
            if org_actual:
                try:
                    import base64
                    org_bytes = base64.b64decode(org_actual)
                    st.image(org_bytes, width=600, caption="Organigrama guardado")
                except Exception as e:
                    st.info("No hay organigrama guardado")
            else:
                st.info("No hay organigrama. Sube uno.")
            
            org_file = st.file_uploader("Subir Organigrama", type=['png', 'jpg', 'jpeg'], disabled=not puede_editar)
            
            # BOTON GUARDAR
            submitted = st.form_submit_button("Guardar Todo", disabled=not puede_editar)
            
            if submitted:
                update_data = {
                    "nombre": nombre, 
                    "giro": giro, 
                    "objetivo_plan": objetivo_plan, 
                    "mision": mision, 
                    "vision": vision, 
                    "obj_general": obj_gen, 
                    "obj_especificos": obj_esp, 
                    "politicas": politicas, 
                    "valores": valores
                }
                
                # Logo nuevo
                if logo_file is not None:
                    try:
                        import base64
                        update_data['logo'] = base64.b64encode(logo_file.getvalue()).decode('utf-8')
                    except Exception as e:
                        st.error(f"Error logo: {e}")
                
                # Organigrama nuevo
                if org_file is not None:
                    try:
                        import base64
                        update_data['organigrama'] = base64.b64encode(org_file.getvalue()).decode('utf-8')
                    except Exception as e:
                        st.error(f"Error organigrama: {e}")
                
                # Guardar
                try:
                    supabase.table('empresas').update(update_data).eq('id', empresa_id).execute()
                    st.success("Guardado!")
                    st.rerun()
                except Exception as e:
                    st.error(f"Error al guardar: {e}")
                    
    # --- PESTAÑA 2: DIAGNÓSTICO SITUACIONAL ---
    with tab2:
        st.header("Diagnóstico Situacional (Análisis de Matrices)")

        def procesar_made_madi(data_str, tipo):
            if isinstance(data_str, pd.DataFrame):
                df = data_str.copy()
            else:
                try:
                    df = pd.read_csv(StringIO(data_str), sep='\t', header=0)
                except Exception as e:
                    st.error(f"Error al leer los datos pegados. Error: {e}")
                    return pd.DataFrame()
            def normalizar_nombre(nombre):
                nombre_sin_tildes = unicodedata.normalize('NFKD', str(nombre)).encode('ascii', 'ignore').decode('utf-8')
                return nombre_sin_tildes.lower().replace(' ', '_').replace('%', '_percent').replace('°', '')
            df.columns = [normalizar_nombre(c) for c in df.columns]
            mapeo_columnas = {
                'n': 'n_temp', 'variable': 'variable', 'factor': 'factor', 'producto': 'producto',
                'precio': 'precio', 'plaza': 'plaza', 'promocion': 'promocion', 'rating': 'rating',
                'weight__percent': 'weight_percent', 'weight_percent': 'weight_percent', 
                'valor': 'valor', 'total': 'total'
            }
            df.rename(columns=mapeo_columnas, inplace=True)
            numeric_cols = ['rating', 'weight_percent', 'valor', 'total']
            for col in numeric_cols:
                if col in df.columns:
                    df[col] = pd.to_numeric(df[col], errors='coerce')
            df.replace({np.nan: None}, inplace=True)
            columnas_finales_bd = ['variable', 'factor', 'producto', 'precio', 'plaza', 'promocion', 'rating', 'weight_percent', 'valor', 'total']
            for col in columnas_finales_bd:
                if col not in df.columns:
                    df[col] = None
            return df[columnas_finales_bd]

        def display_and_edit_matrix(tipo_matriz, analisis_propio_data):
            df_db = get_datos_tabla('matriz_marketing', empresa_id, tipo_matriz_filter=tipo_matriz)
            if not df_db.empty:
                st.info("Puedes editar los datos directamente en la tabla.")
                df_display = df_db.rename(columns={
                    'variable': 'Variable', 'factor': 'Factor', 'producto': 'Producto', 'precio': 'Precio',
                    'plaza': 'Plaza', 'promocion': 'Promoción', 'rating': 'Rating', 'weight_percent': 'Weight %',
                    'valor': 'Valor', 'total': 'Total'
                })
                columnas_a_mostrar = [c for c in df_display.columns if c not in ['id', 'empresa_id', 'tipo_matriz']]
                edited_df_display = st.data_editor(df_display[columnas_a_mostrar], key=f"editor_{tipo_matriz}", num_rows="dynamic", use_container_width=True, disabled=not puede_editar)
                if st.button(f"💾 Guardar Cambios en {tipo_matriz}", disabled=not puede_editar, key=f"save_{tipo_matriz}"):
                    try:
                        df_to_save = edited_df_display.rename(columns={
                            'Variable': 'variable', 'Factor': 'factor', 'Producto': 'producto', 'Precio': 'precio',
                            'Plaza': 'plaza', 'Promoción': 'promocion', 'Rating': 'rating', 'Weight %': 'weight_percent',
                            'Valor': 'valor', 'Total': 'total'
                        })
                        df_to_save = procesar_made_madi(df_to_save.to_csv(sep='\t', index=False), tipo_matriz)
                        supabase.table('matriz_marketing').delete().eq('empresa_id', empresa_id).eq('tipo_matriz', tipo_matriz).execute()
                        df_to_save['empresa_id'] = empresa_id
                        df_to_save['tipo_matriz'] = tipo_matriz
                        supabase.table('matriz_marketing').insert(df_to_save.to_dict(orient='records')).execute()
                        st.success(f"Cambios en {tipo_matriz} guardados."); 
                        st.rerun()
                    except Exception as e:
                        st.error(f"Error al guardar {tipo_matriz}: {e}")
                total_score = pd.to_numeric(df_db['total'], errors='coerce').sum()
                st.metric(f"Puntaje Total {tipo_matriz}", f"{total_score}")
            else:
                st.info(f"Aún no hay datos para la Matriz {tipo_matriz}. Pega los datos desde Excel para comenzar.")
        
        diag_tab1, diag_tab2, diag_tab3, diag_tab4, diag_tab5 = st.tabs([
            "Matriz MADE", "Matriz MADI", "Matriz de Posicionamiento", "Matriz PEST", "Matriz FODA Numérico"
        ])

        # --- MADE ---
        with diag_tab1:
            st.subheader("Análisis de Marketing Interno (MADE)")
            # 1. Entrada de datos
            with st.expander("📋 Pegar datos de MADE desde Excel"):
                made_paste_data = st.text_area("Pega tus datos de MADE aquí", height=200, key="paste_MADE")
                if st.button("Procesar y Guardar Datos de MADE", key="process_made", disabled=not puede_editar):
                    if made_paste_data:
                        try:
                            df_made = procesar_made_madi(made_paste_data, 'MADE')
                            supabase.table('matriz_marketing').delete().eq('empresa_id', empresa_id).eq('tipo_matriz', 'MADE').execute()
                            df_made['empresa_id'] = empresa_id
                            df_made['tipo_matriz'] = 'MADE'
                            supabase.table('matriz_marketing').insert(df_made.to_dict(orient='records')).execute()
                            st.success(f"¡{len(df_made)} filas importadas a MADE exitosamente!"); 
                            st.rerun()
                        except Exception as e:
                            st.error(f"Error al procesar datos de MADE: {e}")
            # 2. Mostrar datos existentes
            df_made_actual = get_datos_tabla('matriz_marketing', empresa_id, tipo_matriz_filter='MADE')
            if not df_made_actual.empty:
                st.write("**Datos Actuales:**")
                st.dataframe(df_made_actual.drop(columns=['id', 'empresa_id', 'tipo_matriz'], errors='ignore'), use_container_width=True)
            
            st.divider()
            # 3. Análisis con IA
            st.subheader("📊 Análisis de MADE")
            if st.button("🤖 Generar Análisis MADE con IA", key="gen_made_ia", disabled=not puede_editar):
                with st.spinner("Analizando Marketing Interno..."):
                    contexto = df_made_actual.to_string() if not df_made_actual.empty else "Sin datos numéricos"
                    analisis_generado = generar_analisis_ia("MADE (Marketing Interno)", contexto)
                    st.session_state['made_analisis_temp'] = analisis_generado
                    st.rerun()
            
            # 4. Formulario para guardar análisis
            with st.form("form_analisis_made"):
                analisis_made_val = st.session_state.get('made_analisis_temp', empresa_data.get('analisis_made', ''))
                analisis_made_text = st.text_area("Análisis de Marketing Interno (MADE)", value=analisis_made_val, height=200, disabled=not puede_editar)
                if st.form_submit_button("💾 Guardar Análisis MADE", disabled=not puede_editar):
                    guardar_analisis_db(empresa_id, 'made', analisis_made_text)
            
            # 5. Mostrar último análisis
            mostrar_ultimo_analisis_guardado(empresa_data, 'made')

            # --- MADI ---
        with diag_tab2:
            st.subheader("Análisis de Marketing Externo (MADI)")
            with st.expander("📋 Pegar datos de MADI desde Excel"):
                madi_paste_data = st.text_area("Pega tus datos de MADI aquí", height=200, key="paste_MADI")
                if st.button("Procesar y Guardar Datos de MADI", key="process_madi", disabled=not puede_editar):
                    if madi_paste_data:
                        try:
                            df_madi = procesar_made_madi(madi_paste_data, 'MADI')
                            supabase.table('matriz_marketing').delete().eq('empresa_id', empresa_id).eq('tipo_matriz', 'MADI').execute()
                            df_madi['empresa_id'] = empresa_id
                            df_madi['tipo_matriz'] = 'MADI'
                            supabase.table('matriz_marketing').insert(df_madi.to_dict(orient='records')).execute()
                            st.success(f"¡{len(df_madi)} filas importadas a MADI exitosamente!"); 
                            st.rerun()
                        except Exception as e:
                            st.error(f"Error al procesar datos de MADI: {e}")
                    else:
                        st.warning("El área de texto está vacía. Pega tus datos para procesar.")
            
            # Mostrar datos existentes (fuera del expander)
            df_madi_actual = get_datos_tabla('matriz_marketing', empresa_id, tipo_matriz_filter='MADI')
            if not df_madi_actual.empty:
                st.write("**Datos Actuales:**")
                st.dataframe(df_madi_actual.drop(columns=['id', 'empresa_id', 'tipo_matriz'], errors='ignore'), use_container_width=True)
            
            st.divider()
            # Análisis con IA
            st.subheader("📊 Análisis de MADI")
            if st.button("🤖 Generar Análisis MADI con IA", key="gen_madi_ia", disabled=not puede_editar):
                with st.spinner("Analizando Marketing Externo..."):
                    contexto = df_madi_actual.to_string() if not df_madi_actual.empty else "Sin datos numéricos"
                    analisis_generado = generar_analisis_ia("MADI (Marketing Externo)", contexto)
                    st.session_state['madi_analisis_temp'] = analisis_generado
                    st.rerun()
            
            with st.form("form_analisis_madi"):
                analisis_madi_val = st.session_state.get('madi_analisis_temp', empresa_data.get('analisis_madi', ''))
                analisis_madi_text = st.text_area("Análisis de Marketing Externo (MADI)", value=analisis_madi_val, height=200, disabled=not puede_editar)
                if st.form_submit_button("💾 Guardar Análisis MADI", disabled=not puede_editar):
                    guardar_analisis_db(empresa_id, 'madi', analisis_madi_text)
            
            mostrar_ultimo_analisis_guardado(empresa_data, 'madi')

        # --- POSICIONAMIENTO ---
        with diag_tab3:
            st.subheader("Matriz de Posicionamiento")
            coord_x_val = float(empresa_data.get('posicionamiento_x') or 0)
            coord_y_val = float(empresa_data.get('posicionamiento_y') or 0)
            
            # 1. Entrada de datos
            with st.form("form_posicionamiento"):
                coord_x = st.number_input("Coordenada X", value=coord_x_val, disabled=not puede_editar)
                coord_y = st.number_input("Coordenada Y", value=coord_y_val, disabled=not puede_editar)
                if st.form_submit_button("Guardar Coordenadas", disabled=not puede_editar):
                    try:
                        supabase.table('empresas').update({
                            "posicionamiento_x": coord_x, 
                            "posicionamiento_y": coord_y
                        }).eq('id', empresa_id).execute()
                        st.success("Coordenadas guardadas."); 
                        st.rerun()
                    except Exception as e:
                        st.error(f"Error al guardar coordenadas: {e}")
            
            # 2. Visualización
            fig, ax = plt.subplots()
            ax.axhline(0, color='gray', lw=1); ax.axvline(0, color='gray', lw=1)
            ax.plot(coord_x_val, coord_y_val, 'ro', markersize=10)
            ax.set_title("Matriz de Posicionamiento"); ax.set_xlabel("Eje X"); ax.set_ylabel("Eje Y")
            ax.grid(True, which='both', linestyle='--', linewidth=0.5)
            st.pyplot(fig)
            
            st.divider()
            # 3. Análisis IA
            st.subheader("📊 Análisis de Posicionamiento")
            if st.button("🤖 Generar Análisis de Posicionamiento con IA", key="gen_pos_ia", disabled=not puede_editar):
                with st.spinner("Analizando posición competitiva..."):
                    contexto = f"Coordenadas X: {coord_x_val}, Y: {coord_y_val}"
                    analisis_generado = generar_analisis_ia("Posicionamiento Competitivo", contexto)
                    st.session_state['posicionamiento_analisis_temp'] = analisis_generado
                    st.rerun()
            
            with st.form("form_analisis_pos"):
                analisis_pos_val = st.session_state.get('posicionamiento_analisis_temp', empresa_data.get('analisis_posicionamiento', ''))
                analisis_pos_text = st.text_area("Análisis de Posicionamiento Competitivo", value=analisis_pos_val, height=200, disabled=not puede_editar)
                if st.form_submit_button("💾 Guardar Análisis", disabled=not puede_editar):
                    guardar_analisis_db(empresa_id, 'posicionamiento', analisis_pos_text)
            
            mostrar_ultimo_analisis_guardado(empresa_data, 'posicionamiento')

        # --- PEST ---
        with diag_tab4:
            st.subheader("Análisis PEST")
            with st.expander("📋 Pegar datos desde Excel"):
                pest_paste_data = st.text_area("Pega tus datos PEST aquí (columnas: categoria, factor, tipo_foda, puntaje, importancia)", height=200, key="pest_input")
                if st.button("Procesar y Guardar Datos PEST", key="process_pest", disabled=not puede_editar):
                    if pest_paste_data:
                        try:
                            df_pasted = pd.read_csv(StringIO(pest_paste_data), sep='\t', header=0)
                            df_pasted.columns = ['categoria', 'factor', 'tipo_foda', 'puntaje', 'importancia']
                            df_pasted['valor_ponderado'] = pd.to_numeric(df_pasted['puntaje'], errors='coerce') * (pd.to_numeric(df_pasted['importancia'], errors='coerce') / 100.0)
                            df_pasted['empresa_id'] = empresa_id
                            df_pasted['tipo_matriz'] = 'PEST'
                            supabase.table('matrices').delete().eq('empresa_id', empresa_id).eq('tipo_matriz', 'PEST').execute()
                            supabase.table('matrices').insert(df_pasted.to_dict(orient='records')).execute()
                            st.success(f"¡{len(df_pasted)} filas importadas a PEST exitosamente!"); 
                            st.rerun()
                        except Exception as e:
                            st.error(f"Error al procesar los datos: {e}.")
            
            df_pest = get_datos_tabla('matrices', empresa_id, tipo_matriz_filter='PEST')
            if not df_pest.empty:
                st.write("**Datos PEST Actuales:**")
                edited_pest = st.data_editor(df_pest.drop(columns=['id', 'empresa_id', 'tipo_matriz'], errors='ignore'), num_rows="dynamic", key="editor_pest", use_container_width=True, disabled=not puede_editar)
                if st.button("💾 Guardar Cambios en PEST", key="save_pest_changes", disabled=not puede_editar):
                    try:
                        supabase.table('matrices').delete().eq('empresa_id', empresa_id).eq('tipo_matriz', 'PEST').execute()
                        if not edited_pest.empty:
                            edited_pest['empresa_id'] = empresa_id
                            edited_pest['tipo_matriz'] = 'PEST'
                            supabase.table('matrices').insert(edited_pest.to_dict(orient='records')).execute()
                        st.success("Cambios en PEST guardados."); 
                        st.rerun()
                    except Exception as e:
                        st.error(f"Error al guardar PEST: {e}")
                total_ponderado = pd.to_numeric(df_pest['valor_ponderado'], errors='coerce').sum()
                st.metric("Puntaje Ponderado Total PEST", f"{total_ponderado:.2f}")
            
            st.divider()
            st.subheader("📊 Análisis PEST")
            if st.button("🤖 Generar Análisis PEST con IA", key="gen_pest_ia", disabled=not puede_editar):
                with st.spinner("Analizando entorno PEST..."):
                    contexto = df_pest.to_string() if not df_pest.empty else "Sin datos numéricos"
                    analisis_generado = generar_analisis_ia("PEST (Político, Económico, Social, Tecnológico)", contexto)
                    st.session_state['pest_analisis_temp'] = analisis_generado
                    st.rerun()
            
            with st.form("form_analisis_pest"):
                analisis_pest_val = st.session_state.get('pest_analisis_temp', empresa_data.get('analisis_pest', ''))
                analisis_pest_text = st.text_area("Análisis del Entorno PEST", value=analisis_pest_val, height=200, disabled=not puede_editar)
                if st.form_submit_button("💾 Guardar Análisis PEST", disabled=not puede_editar):
                    guardar_analisis_db(empresa_id, 'pest', analisis_pest_text)
            
            mostrar_ultimo_analisis_guardado(empresa_data, 'pest')

        # --- FODA ---
        with diag_tab5:
            st.subheader("Análisis FODA Cruzado (Numérico)")
            with st.expander("📋 Pegar datos de FODA Cruzado desde Excel"):
                foda_paste_data = st.text_area("Pega tus datos de FODA aquí (columnas: cuadrante, factor_fila, factor_columna, impacto)", height=200, key="foda_paste")
                if st.button("Procesar y Guardar Datos FODA", key="process_foda", disabled=not puede_editar):
                    if foda_paste_data:
                        try:
                            df_foda_pasted = pd.read_csv(StringIO(foda_paste_data), sep='\t', header=0)
                            df_foda_pasted.columns = ['cuadrante', 'factor_fila', 'factor_columna', 'impacto']
                            df_foda_pasted['impacto'] = pd.to_numeric(df_foda_pasted['impacto'], errors='coerce').fillna(0).astype(int)
                            df_foda_pasted['empresa_id'] = empresa_id
                            supabase.table('foda_cruzado').delete().eq('empresa_id', empresa_id).execute()
                            supabase.table('foda_cruzado').insert(df_foda_pasted.to_dict(orient='records')).execute()
                            st.success(f"¡{len(df_foda_pasted)} filas importadas a FODA Cruzado!"); 
                            st.rerun()
                        except Exception as e:
                            st.error(f"Error al procesar los datos: {e}.")
            
            df_foda = get_datos_tabla('foda_cruzado', empresa_id)
            if not df_foda.empty:
                st.write("**Datos FODA Actuales:**")
                edited_foda = st.data_editor(df_foda.drop(columns=['id', 'empresa_id'], errors='ignore'), num_rows="dynamic", key="editor_foda", use_container_width=True, disabled=not puede_editar)
                if st.button("💾 Guardar Cambios en FODA", key="save_foda_changes", disabled=not puede_editar):
                    try:
                        supabase.table('foda_cruzado').delete().eq('empresa_id', empresa_id).execute()
                        if not edited_foda.empty:
                            edited_foda['empresa_id'] = empresa_id
                            supabase.table('foda_cruzado').insert(edited_foda.to_dict(orient='records')).execute()
                        st.success("Cambios en FODA guardados."); 
                        st.rerun()
                    except Exception as e:
                        st.error(f"Error al guardar FODA: {e}")
                
                analisis_df, _, estrategia_principal, puntajes_foda = analizar_foda(df_foda)
                if analisis_df is not None:
                    st.subheader("🎯 Postura Competitiva Sugerida")
                    st.info(f"Estrategia Principal: {estrategia_principal}")
                    col1, col2 = st.columns(2)
                    with col1:
                        st.dataframe(analisis_df, use_container_width=True)
                    with col2:
                        grafico_foda = generar_grafico_foda_radar_mejorado(puntajes_foda)
                        if grafico_foda: 
                            st.image(grafico_foda)
            
            st.divider()
            st.subheader("📊 Análisis FODA")
            if st.button("🤖 Generar Análisis FODA con IA", key="gen_foda_ia", disabled=not puede_editar):
                with st.spinner("Analizando matrices FODA..."):
                    contexto = df_foda.to_string() if not df_foda.empty else "Sin datos numéricos"
                    analisis_generado = generar_analisis_ia("FODA Cruzado", contexto)
                    st.session_state['foda_analisis_temp'] = analisis_generado
                    st.rerun()
            
            with st.form("form_analisis_foda"):
                analisis_foda_val = st.session_state.get('foda_analisis_temp', empresa_data.get('analisis_foda', ''))
                analisis_foda_text = st.text_area("Análisis FODA Cruzado", value=analisis_foda_val, height=200, disabled=not puede_editar)
                if st.form_submit_button("💾 Guardar Análisis FODA", disabled=not puede_editar):
                    guardar_analisis_db(empresa_id, 'foda', analisis_foda_text)
            
            mostrar_ultimo_analisis_guardado(empresa_data, 'foda')

# --- PESTAÑA 3: ESTRATEGIA (CORREGIDA CON GENERACIÓN IA) ---
    with tab3:
        st.header("🎯 Formulación de Estrategias")
    
    df_estrategias = get_datos_tabla('estrategias_generadas', empresa_id)
    df_foda_estrategia = get_datos_tabla('foda_cruzado', empresa_id)
    
    if df_estrategias.empty:
        st.info("No hay estrategias generadas. Utiliza el botón de abajo para generarlas automáticamente con IA basándote en el análisis FODA.")
        
        if st.button("🤖 Generar Estrategias con IA (3 por cuadrante = 12 estrategias)", disabled=not puede_editar, type="primary"):
            if df_foda_estrategia.empty:
                st.error("Primero debes cargar los datos del FODA Cruzado en la pestaña anterior.")
            else:
                with st.spinner("Generando 12 estrategias estratégicas (3 por cuadrante) con 5 actividades cada una..."):
                    # Obtener factores por cuadrante para contexto
                    contexto_foda = df_foda_estrategia.to_string()
                    
                    prompt_estrategias = (
                        "Basado en el siguiente análisis FODA Cruzado:\n"
                        f"{contexto_foda}\n\n"
                        "Genera exactamente 3 estrategias para cada uno de los 4 cuadrantes (FO, FA, DO, DA), total 12 estrategias.\n"
                        "Para cada estrategia proporciona:\n"
                        "1. Cuadrante (FO, FA, DO, o DA)\n"
                        "2. Estrategia: Descripción clara y específica de la estrategia\n"
                        "3. Importancia: Selecciona una de (Alta, Media Alta, Media Baja, Baja)\n"
                        "4. Actividades: Lista de EXACTAMENTE 5 actividades clave separadas por punto y coma (;) para implementar la estrategia\n"
                        "5. Plan Asignado: Selecciona uno de (Plan Administrativo, Plan Operativo, Plan Tecnológico, Plan Financiero, Plan de Monitoreo y control, Plan de Mejora, Plan de Contingencia)\n\n"
                        "Formato de salida EXACTO (una estrategia por línea):\n"
                        "CUADRANTE|ESTRATEGIA|IMPORTANCIA|ACTIVIDAD|PLAN_ASIGNADO\n\n"
                        "Ejemplo:\n"
                        "FO|Expandir mercado en nuevas regiones utilizando fortalezas tecnológicas|Alta|Investigar mercados potenciales;Adaptar producto a nuevas necesidades;Lanzar campaña marketing digital;Capacitar equipo de ventas;Establecer alianzas locales|Plan Operativo\n"
                        "FO|Alianza estratégica con proveedores clave|Media Alta|Identificar proveedores potenciales;Negociar contratos marco;Implementar integración de sistemas;Capacitar personal en nuevos procesos;Evaluar desempeño de proveedores|Plan Administrativo\n"
                        "FA|Programa de retención de clientes ante nueva competencia|Alta|Analizar tasa de churn actual;Crear programa fidelización;Capacitar equipo de servicio al cliente;Implementar encuestas satisfacción;Diseñar promociones exclusivas|Plan de Mejora\n\n"
                        "Genera exactamente 12 líneas (3 por cada cuadrante FO, FA, DO, DA). Cada estrategia debe tener EXACTAMENTE 5 actividades separadas por punto y coma (;). No uses encabezados."
                    )                        
                    resultado = generar_analisis(prompt_estrategias)
                    
                    # Parsear resultado con mejor manejo de errores
                    estrategias_list = []
                    lineas = [l.strip() for l in resultado.split('\n') if l.strip()]
                    
                    st.write("Debug - Respuesta de IA:", resultado[:500] if len(resultado) > 500 else resultado)
                    
                    for linea in lineas[:12]:  # Máximo 12 estrategias
                        partes = linea.split('|')
                        if len(partes) >= 5:
                            # Asegurar que hay 5 actividades
                            actividades = partes[3].strip()
                            # Contar actividades separadas por ;
                            num_actividades = len([a for a in actividades.split(';') if a.strip()])
                            
                            estrategias_list.append({
                                'empresa_id': empresa_id,
                                'cuadrante': partes[0].strip().upper(),
                                'estrategia': partes[1].strip(),
                                'importancia': partes[2].strip(),
                                'actividades': actividades,
                                'plan_asignado': partes[4].strip()
                            })
                        elif len(partes) == 4:
                            # Intentar detectar si falta el plan asignado
                            estrategias_list.append({
                                'empresa_id': empresa_id,
                                'cuadrante': partes[0].strip().upper(),
                                'estrategia': partes[1].strip(),
                                'importancia': partes[2].strip(),
                                'actividades': partes[3].strip(),
                                'plan_asignado': 'Plan Operativo'  # Default
                            })
                    
                    if estrategias_list:
                        try:
                            # Guardar en Supabase
                            supabase.table('estrategias_generadas').delete().eq('empresa_id', empresa_id).execute()
                            supabase.table('estrategias_generadas').insert(estrategias_list).execute()
                            
                            total_actividades = sum([len([a for a in est['actividades'].split(';') if a.strip()]) for est in estrategias_list])
                            st.success(f"✅ {len(estrategias_list)} estrategias generadas y guardadas exitosamente.")
                            st.info(f"📋 Total de actividades generadas: {total_actividades} (objetivo: 60 actividades)")
                            st.rerun()
                        except Exception as e:
                            st.error(f"Error al guardar estrategias: {e}")
                    else:
                        st.error("No se pudieron parsear las estrategias generadas. Intenta nuevamente.")
                        st.write("Respuesta completa de la IA para debugging:")
                        st.code(resultado)
    else:
        # Mostrar estrategias existentes en editor
        st.success(f"Se encontraron {len(df_estrategias)} estrategias generadas.")
        
        # Calcular total de actividades
        total_actividades = 0
        for _, est in df_estrategias.iterrows():
            acts = str(est.get('actividades', ''))
            total_actividades += len([a for a in acts.split(';') if a.strip()])
        
        st.info(f"📋 Total de actividades: {total_actividades} (objetivo: 60 actividades = 12 estrategias × 5 actividades)")
        
        st.write("**Editar Estrategias:**")
        
        edited_df = st.data_editor(
            df_estrategias.drop(columns=['id', 'empresa_id'], errors='ignore'), 
            num_rows="dynamic", 
            key="editor_estrategias", 
            use_container_width=True,
            disabled=not puede_editar,
            column_config={
                "cuadrante": st.column_config.SelectboxColumn("Cuadrante", options=["FO", "FA", "DO", "DA"]),
                "importancia": st.column_config.SelectboxColumn("Importancia", options=["Alta", "Media Alta", "Media Baja", "Baja"]),
                "plan_asignado": st.column_config.SelectboxColumn("Plan Asignado", 
                    options=["Plan Administrativo", "Plan Operativo", "Plan Tecnológico", "Plan Financiero", 
                            "Plan de Monitoreo y control", "Plan de Mejora", "Plan de Contingencia"]),
                "estrategia": st.column_config.TextColumn("Estrategia", width="large"),
                "actividades": st.column_config.TextColumn("Actividades (separar con ;)", width="large")
            }
        )
        
        col1, col2, col3 = st.columns(3)
        with col1:
            if st.button("💾 Guardar Cambios", disabled=not puede_editar):
                try:
                    supabase.table('estrategias_generadas').delete().eq('empresa_id', empresa_id).execute()
                    if not edited_df.empty:
                        df_to_save = edited_df.copy()
                        df_to_save['empresa_id'] = empresa_id
                        supabase.table('estrategias_generadas').insert(df_to_save.to_dict(orient='records')).execute()
                    st.success("Estrategias actualizadas."); 
                    st.rerun()
                except Exception as e:
                    st.error(f"Error al guardar estrategias: {e}")
        with col2:
            if st.button("🗑️ Eliminar Todas", type="secondary", disabled=not puede_editar):
                try:
                    supabase.table('estrategias_generadas').delete().eq('empresa_id', empresa_id).execute()
                    st.success("Estrategias eliminadas."); 
                    st.rerun()
                except Exception as e:
                    st.error(f"Error al eliminar: {e}")
        with col3:
            if st.button("🔄 Regenerar con IA", disabled=not puede_editar):
                # Eliminar actuales y regenerar
                try:
                    supabase.table('estrategias_generadas').delete().eq('empresa_id', empresa_id).execute()
                    st.rerun()
                except:
                    pass
                        
    # --- PESTAÑA 4: PLANES ESTRATÉGICOS ---
    with tab4:
        st.header("📋 Planes Estratégicos Funcionales")
        
        # Importar re al inicio de la pestaña para evitar UnboundLocalError
        import re
        
        # Obtener datos necesarios para contextualizar
        df_foda_temp = get_datos_tabla('foda_cruzado', empresa_id)
        analisis_df_temp, _, estrategia_principal, puntajes_foda = analizar_foda(df_foda_temp)
        df_estrategias_planes = get_datos_tabla('estrategias_generadas', empresa_id)
        empresa_datos = get_datos_empresa(empresa_id)
        
        # Calcular puntaje PEST para contexto
        df_pest_temp = get_datos_tabla('matrices', empresa_id, tipo_matriz_filter='PEST')
        pest_total = df_pest_temp['valor_ponderado'].sum() if not df_pest_temp.empty else 0
        
        # Determinar contexto estratégico
        es_ofensiva = "Ofensiva" in str(estrategia_principal)
        es_adaptativa = "Adaptativa" in str(estrategia_principal)
        es_defensiva = "Defensiva" in str(estrategia_principal)
        es_supervivencia = "Supervivencia" in str(estrategia_principal)
        
        contexto_empresa = {
            'nombre': empresa_datos.get('nombre', 'La empresa'),
            'giro': empresa_datos.get('giro', 'su sector'),
            'estrategia_principal': estrategia_principal or 'No definida',
            'pest_total': pest_total,
            'postura': 'crecimiento' if (es_ofensiva or es_adaptativa) else 'consolidación/defensa'
        }
        
        st.info(f"""
        **Contexto Estratégico Detectado:**
        - Estrategia Principal: **{contexto_empresa['estrategia_principal']}**
        - Postura: **{contexto_empresa['postura'].upper()}**
        - Entorno PEST: **{'Favorable' if pest_total > 2.5 else 'Desafiante'}** ({pest_total:.2f})
        """)
        
        col1, col2 = st.columns([2, 1])
        with col1:
            if st.button("🤖 Generar 7 Planes Funcionales Profesionales", disabled=not puede_editar, type="primary"):
                with st.spinner("Elaborando planes estratégicos detallados con IA. Esto puede tomar 1-2 minutos..."):
                    
                    # Preparar contexto de estrategias generadas
                    contexto_estrategias = ""
                    if not df_estrategias_planes.empty:
                        for idx, row in df_estrategias_planes.iterrows():
                            contexto_estrategias += f"\n{idx+1}. [{row['cuadrante']}] {row['estrategia']} -> {row['plan_asignado']}"
                    
                    # Construir prompt como lista de líneas para evitar problemas con triple comillas
                    lineas_prompt = [
                        "Actúa como un consultor senior de estrategia empresarial con 20 años de experiencia.",
                        f"Elabora 7 PLANES FUNCIONALES ESTRATÉGICOS de alto nivel para {contexto_empresa['nombre']},",
                        f"empresa del sector {contexto_empresa['giro']}.",
                        "",
                        "CONTEXTO ESTRATÉGICO:",
                        f"- Estrategia principal FODA: {contexto_empresa['estrategia_principal']}",
                        f"- Postura estratégica: {contexto_empresa['postura']}",
                        f"- Entorno PEST score: {contexto_empresa['pest_total']:.2f}",
                        f"- Estrategias generadas: {contexto_estrategias if contexto_estrategias else 'No hay estrategias previas'}",
                        "",
                        "ESTRUCTURA REQUERIDA PARA CADA PLAN (7 planes totales):",
                        "",
                        "1. PLAN ADMINISTRATIVO",
                        "2. PLAN OPERATIVO",
                        "3. PLAN TECNOLÓGICO",
                        "4. PLAN FINANCIERO",
                        "5. PLAN DE MONITOREO Y CONTROL",
                        "6. PLAN DE MEJORA",
                        "7. PLAN DE CONTINGENCIA",
                        "",
                        "PARA CADA PLAN DEBES INCLUIR EXACTAMENTE:",
                        "",
                        "=== [NOMBRE DEL PLAN] ===",
                        "",
                        "1. FUNDAMENTO ESTRATÉGICO",
                        "[Explicación de por qué este plan es crítico para la empresa en su contexto actual, 3-4 párrafos profundos]",
                        "",
                        "2. OBJETIVO GENERAL DEL PLAN",
                        "[Objetivo SMART específico]",
                        "",
                        "3. OBJETIVOS ESPECÍFICOS (mínimo 3)",
                        "- Objetivo 1",
                        "- Objetivo 2",
                        "- Objetivo 3",
                        "",
                        "4. ESTRATEGIAS DE IMPLEMENTACIÓN (mínimo 4 estrategias concretas)",
                        "A. [Nombre estrategia 1]",
                        "   - Descripción detallada",
                        "   - Acciones clave",
                        "B. [Nombre estrategia 2]",
                        "   - Descripción detallada",
                        "   - Acciones clave",
                        "[C continuar...]",
                        "",
                        "5. KPIs Y METAS (mínimo 5 KPIs por plan)",
                        "- KPI 1: [Nombre] | Meta: [X] | Frecuencia: [mensual/trimestral]",
                        "- KPI 2: [Nombre] | Meta: [X] | Frecuencia: [mensual/trimestral]",
                        "[Continuar...]",
                        "",
                        "6. RECURSOS REQUERIDOS",
                        "- Humanos: [Detalle]",
                        "- Financieros: [Presupuesto estimado]",
                        "- Tecnológicos: [Infraestructura]",
                        "- Temporales: [Cronograma]",
                        "",
                        "7. RESPONSABLES Y GOBIERNO",
                        "- Responsable principal: [Rol]",
                        "- Comité de seguimiento: [Miembros]",
                        "- Frecuencia de revisión: [Semanal/Mensual]",
                        "",
                        "8. RIESGOS Y MITIGACIÓN",
                        "- Riesgo 1: [Descripción] | Mitigación: [Acción]",
                        "- Riesgo 2: [Descripción] | Mitigación: [Acción]",
                        "",
                        "9. ALINEACIÓN CON ESTRATEGIA FODA",
                        "[Cómo este plan contribuye específicamente a FO, FA, DO o DA]",
                        "",
                        "REQUISITOS DE CALIDAD:",
                        "- Lenguaje ejecutivo y profesional",
                        "- Contenido específico, no genérico",
                        "- Cada plan debe tener 800-1200 palabras mínimo",
                        "- Los KPIs deben ser cuantificables y realistas",
                        "- Las estrategias deben ser accionables",
                        "- Considerar el contexto PEST y FODA proporcionado",
                        "- NO usar frases vacías como 'mejorar procesos' sin especificar cómo",
                        "",
                        "GENERA LOS 7 PLANES COMPLETOS AHORA:"
                    ]
                    
                    prompt_planes_maestros = "\n".join(lineas_prompt)
                    
                    planes_generados = generar_analisis(prompt_planes_maestros)
                    st.session_state['planes_maestros_generados'] = planes_generados
                    st.success("✅ Planes funcionales generados con éxito")
                    st.rerun()
        
        with col2:
            st.caption("""
            **Los 7 Planes Funcionales:**
            1. 📊 Administrativo
            2. ⚙️ Operativo
            3. 💻 Tecnológico
            4. 💰 Financiero
            5. 📈 Monitoreo y Control
            6. 🚀 Mejora
            7. 🛡️ Contingencia
            """)
        
        # Mostrar y editar planes
        planes_actuales = st.session_state.get('planes_maestros_generados', empresa_data.get('analisis_operativo', ''))
        
        if planes_actuales:
            st.divider()
            st.subheader("📄 Planes Estratégicos Funcionales - Documento Maestro")
            
            # Mostrar en tabs los 7 planes si se detectan en el texto
            planes_str = str(planes_actuales)
            if "===" in planes_str:
                # Intentar dividir por planes
                try:
                    # Patrón más flexible para detectar planes
                    patron_planes = r'={2,}\s*(PLAN\s+[A-ZÁÉÍÓÚÑa-záéíóúñ]+(?:\s+[A-ZÁÉÍÓÚÑa-záéíóúñ]+)*)\s*={2,}'
                    planes_secciones = re.split(patron_planes, planes_str, flags=re.IGNORECASE)
                    
                    if len(planes_secciones) > 1:
                        # Crear tabs para cada plan detectado
                        nombres_planes = []
                        contenidos_planes = []
                        
                        for i in range(1, len(planes_secciones), 2):
                            if i < len(planes_secciones):
                                nombre = planes_secciones[i].strip()
                                contenido = planes_secciones[i+1] if i+1 < len(planes_secciones) else ""
                                # Limpiar nombre
                                nombre_limpio = re.sub(r'[^\w\s-]', '', nombre).strip()
                                if nombre_limpio and len(nombre_limpio) < 60:
                                    nombres_planes.append(nombre_limpio)
                                    contenidos_planes.append(contenido)
                        
                        if nombres_planes:
                            tabs_planes = st.tabs(nombres_planes)
                            for tab, nombre, contenido in zip(tabs_planes, nombres_planes, contenidos_planes):
                                with tab:
                                    st.markdown(f"### {nombre}")
                                    st.markdown(contenido)
                                    
                                    # Botón para copiar contenido específico
                                    key_segura = re.sub(r'[^\w]', '_', nombre)[:30]
                                    if st.button(f"📋 Copiar contenido", key=f"copy_{key_segura}"):
                                        st.code(contenido, language='markdown')
                        else:
                            st.info("Se generaron los planes pero no se pudieron separar automáticamente. Revisa el documento completo abajo.")
                    else:
                        st.info("No se detectaron separadores de planes. Mostrando documento completo.")
                        
                except Exception as e:
                    st.warning(f"No se pudieron dividir los planes automáticamente: {e}")
                    st.info("Mostrando documento completo para edición manual.")
            
            # Editor completo
            with st.form("form_planes_maestros"):
                st.write("**Editar Documento Completo de Planes:**")
                planes_editados = st.text_area(
                    "Planes Estratégicos Funcionales", 
                    value=planes_str, 
                    height=800, 
                    disabled=not puede_editar,
                    help="Edite los 7 planes funcionales. Use === NOMBRE DEL PLAN === para separar secciones."
                )
                
                col_save, col_regen = st.columns(2)
                with col_save:
                    submitted_save = st.form_submit_button("💾 Guardar Planes Maestros", disabled=not puede_editar)
                    if submitted_save:
                        guardar_analisis_db(empresa_id, 'operativo', planes_editados)
                        st.session_state['planes_maestros_generados'] = planes_editados
                
                with col_regen:
                    submitted_regen = st.form_submit_button("🔄 Regenerar Todo", disabled=not puede_editar)
                    if submitted_regen:
                        if 'planes_maestros_generados' in st.session_state:
                            del st.session_state['planes_maestros_generados']
                        st.rerun()
            
            # Vista previa estructurada
            with st.expander("📊 Ver Resumen Estructurado de Planes"):
                try:
                    # Extraer KPIs mencionados
                    kpis_encontrados = re.findall(r'KPI\s*\d*[:.-]\s*([^\n|]+)', planes_str, re.IGNORECASE)
                    if kpis_encontrados:
                        st.write("**KPIs Detectados (primeros 20):**")
                        for i, kpi in enumerate(kpis_encontrados[:20], 1):
                            texto_kpi = kpi.strip()[:100]
                            st.write(f"{i}. {texto_kpi}...")
                    
                    # Contar estrategias
                    estrategias_count = len(re.findall(r'^[A-Z][.-]\s+', planes_str, re.MULTILINE))
                    col_m1, col_m2, col_m3 = st.columns(3)
                    with col_m1:
                        st.metric("Estrategias Detectadas", estrategias_count)
                    
                    # Contar planes
                    planes_count = len(re.findall(r'={2,}\s*PLAN', planes_str, re.IGNORECASE))
                    with col_m2:
                        st.metric("Planes Detectados", planes_count if planes_count > 0 else "No detectado")
                    
                    # Contar objetivos
                    objetivos_count = len(re.findall(r'^\s*[-•]\s*Objetivo', planes_str, re.MULTILINE | re.IGNORECASE))
                    with col_m3:
                        st.metric("Objetivos Detectados", objetivos_count)
                        
                except Exception as e:
                    st.error(f"Error al analizar el contenido: {e}")
        
        else:
            st.info("""
            👆 **Haz clic en "Generar 7 Planes Funcionales Profesionales"** para crear:
            
            **Plan Administrativo**: Estructura organizacional, gestión del talento, cultura corporativa
            
            **Plan Operativo**: Procesos productivos, cadena de suministro, calidad, logística
            
            **Plan Tecnológico**: Infraestructura digital, transformación digital, ciberseguridad, innovación
            
            **Plan Financiero**: Presupuesto, flujo de caja, inversión, rentabilidad, control financiero
            
            **Plan de Monitoreo y Control**: Dashboards, indicadores, auditorías, seguimiento estratégico
            
            **Plan de Mejora**: Metodologías (Kaizen, Six Sigma), optimización continua, capacitación
            
            **Plan de Contingencia**: Gestión de riesgos, continuidad del negocio, planes de respuesta
            """)
        
        # Mostrar último guardado
        mostrar_ultimo_analisis_guardado(empresa_data, 'operativo')        
        
    # --- PESTAÑA 5: CMI/INDICADORES ---
    with tab5:
        st.header("CMI / Indicadores")
        df_estrategias_cmi = get_datos_tabla('estrategias_generadas', empresa_id)
        
        if not df_estrategias_cmi.empty:
            # Verificar si ya existe un CMI guardado
            cmi_guardado = empresa_data.get('analisis_cmi', '')
            df_cmi_actual = pd.DataFrame()
            
            # Si hay CMI guardado, convertirlo a DataFrame
            if cmi_guardado and '|' in str(cmi_guardado):
                try:
                    df_cmi_actual = pd.read_csv(io.StringIO(cmi_guardado), sep="|")
                    st.success("📋 CMI cargado desde datos guardados")
                except Exception as e:
                    st.warning(f"Error al cargar CMI guardado: {e}")
                    df_cmi_actual = pd.DataFrame()
            
            # Botón para regenerar con IA
            col1, col2 = st.columns([1, 2])
            with col1:
                if st.button("🤖 Regenerar CMI con IA", disabled=not puede_editar, type="primary"):
                    with st.spinner("Generando Cuadro de Mando Integral..."):
                        df_cmi_generado = generar_cuadro_de_mando_ia(df_estrategias_cmi)
                        if not df_cmi_generado.empty:
                            st.session_state['cmi_df_temp'] = df_cmi_generado
                            st.success("CMI generado por IA. Edita la tabla y guarda los cambios.")
                            st.rerun()
                        else:
                            st.error("No se pudo generar el CMI. Verifica las estrategias.")
            
            with col2:
                if 'cmi_df_temp' in st.session_state:
                    st.info("💡 Estás viendo el CMI generado por IA. Edita los valores y guarda.")
                elif not df_cmi_actual.empty:
                    st.info("💡 Puedes editar directamente la tabla y guardar los cambios.")
                else:
                    st.info("👆 Genera el CMI con IA para comenzar")
            
            # Determinar qué DataFrame mostrar (temporal de IA o guardado)
            if 'cmi_df_temp' in st.session_state:
                df_cmi_editar = st.session_state['cmi_df_temp'].copy()
            elif not df_cmi_actual.empty:
                df_cmi_editar = df_cmi_actual.copy()
            else:
                df_cmi_editar = pd.DataFrame()
            
            # Mostrar editor de tabla si hay datos
            if not df_cmi_editar.empty:
                st.divider()
                st.subheader("📊 Cuadro de Mando Integral - Editable")
                
                # Definir columnas esperadas
                columnas_cmi = ['Estrategia', 'Perspectiva', 'KPIs', 'Formulas', 'Frecuencia', 'LI', 'LC', 'LS']
                
                # Asegurar que todas las columnas existan
                for col in columnas_cmi:
                    if col not in df_cmi_editar.columns:
                        df_cmi_editar[col] = ''
                
                # Reordenar columnas
                df_cmi_editar = df_cmi_editar[columnas_cmi]
                
                # Editor de datos
                edited_cmi = st.data_editor(
                    df_cmi_editar,
                    num_rows="dynamic",
                    key="editor_cmi",
                    use_container_width=True,
                    disabled=not puede_editar,
                    column_config={
                        'Estrategia': st.column_config.TextColumn("Estrategia", width="large"),
                        'Perspectiva': st.column_config.SelectboxColumn(
                            "Perspectiva", 
                            options=['Financiera', 'Cliente', 'Procesos', 'Aprendizaje y Control'],
                            width="medium"
                        ),
                        'KPIs': st.column_config.TextColumn("KPIs", width="large"),
                        'Formulas': st.column_config.TextColumn("Fórmulas", width="medium"),
                        'Frecuencia': st.column_config.SelectboxColumn(
                            "Frecuencia",
                            options=['Diaria', 'Semanal', 'Mensual', 'Trimestral', 'Semestral', 'Anual'],
                            width="small"
                        ),
                        'LI': st.column_config.TextColumn("Límite Inferior", width="small"),
                        'LC': st.column_config.TextColumn("Límite Control", width="small"),
                        'LS': st.column_config.TextColumn("Límite Superior", width="small")
                    },
                    hide_index=True
                )
                
                # Botones de acción
                col_guardar, col_cancelar, col_nuevo = st.columns(3)
                
                with col_guardar:
                    if st.button("💾 Guardar CMI", type="primary", disabled=not puede_editar):
                        try:
                            # Validar que no esté vacío
                            if edited_cmi.empty:
                                st.error("No se puede guardar un CMI vacío")
                            else:
                                # Convertir a formato pipe-separated para guardar
                                cmi_texto = edited_cmi.to_csv(sep="|", index=False)
                                
                                # Guardar en base de datos
                                guardar_analisis_db(empresa_id, 'cmi', cmi_texto)
                                
                                # Limpiar temporal
                                if 'cmi_df_temp' in st.session_state:
                                    del st.session_state['cmi_df_temp']
                                
                                st.success("✅ CMI guardado correctamente.")
                                st.rerun()
                        except Exception as e:
                            st.error(f"Error al guardar CMI: {e}")
                
                with col_cancelar:
                    if st.button("❌ Cancelar", disabled=not puede_editar):
                        if 'cmi_df_temp' in st.session_state:
                            del st.session_state['cmi_df_temp']
                        st.rerun()
                
                with col_nuevo:
                    if st.button("➕ Agregar Fila", disabled=not puede_editar):
                        nueva_fila = pd.DataFrame([{
                            'Estrategia': '',
                            'Perspectiva': 'Procesos',
                            'KPIs': '',
                            'Formulas': '',
                            'Frecuencia': 'Mensual',
                            'LI': '',
                            'LC': '',
                            'LS': ''
                        }])
                        # Actualizar el temporal con la nueva fila
                        if 'cmi_df_temp' in st.session_state:
                            st.session_state['cmi_df_temp'] = pd.concat([st.session_state['cmi_df_temp'], nueva_fila], ignore_index=True)
                        else:
                            st.session_state['cmi_df_temp'] = pd.concat([edited_cmi, nueva_fila], ignore_index=True)
                        st.rerun()
                
                # Vista previa de gráfico o resumen
                st.divider()
                with st.expander("📈 Ver Resumen del CMI"):
                    # Contar KPIs por perspectiva
                    if 'Perspectiva' in edited_cmi.columns and not edited_cmi.empty:
                        resumen_perspectiva = edited_cmi.groupby('Perspectiva').agg({
                            'KPIs': 'count',
                            'Estrategia': 'nunique'
                        }).reset_index()
                        resumen_perspectiva.columns = ['Perspectiva', 'N° KPIs', 'N° Estrategias']
                        
                        col_res1, col_res2 = st.columns(2)
                        with col_res1:
                            st.dataframe(resumen_perspectiva, use_container_width=True, hide_index=True)
                        with col_res2:
                            # Gráfico simple de distribución
                            if not resumen_perspectiva.empty:
                                fig = px.pie(
                                    resumen_perspectiva, 
                                    names='Perspectiva', 
                                    values='N° KPIs',
                                    title='Distribución de KPIs por Perspectiva',
                                    hole=0.4
                                )
                                st.plotly_chart(fig, use_container_width=True)
            
            else:
                st.info("👆 Haz clic en 'Regenerar CMI con IA' para generar el cuadro automáticamente basado en tus estrategias.")
        
        else:
            st.warning("No hay estrategias disponibles para generar el CMI. Genera estrategias primero en la pestaña anterior.")    

    # --- PESTAÑA 6: ANÁLISIS DE SEMAFORIZACIÓN ESTRATÉGICA (NUEVA) ---
    with tab6:
        st.header("🚦 Análisis de Semaforización Estratégica")
        st.markdown("""
        *Análisis inteligente del estado de tus estrategias basado en el CMI, FODA y contexto empresarial*
        """)
        
        import re
        
        # Obtener todos los datos necesarios automáticamente
        df_cmi = get_datos_tabla('estrategias_generadas', empresa_id)
        df_foda = get_datos_tabla('foda_cruzado', empresa_id)
        df_pest = get_datos_tabla('matrices', empresa_id, tipo_matriz_filter='PEST')
        empresa_info = get_datos_empresa(empresa_id)
        
        # Verificar si hay datos suficientes
        if df_cmi.empty:
            st.warning("⚠️ No hay estrategias ni CMI configurado. Genera primero el CMI en la pestaña anterior.")
            st.stop()
        
        # Calcular métricas de contexto automáticamente
        analisis_foda_df, resumen_foda, estrategia_principal, puntajes_foda = analizar_foda(df_foda)
        pest_total = df_pest['valor_ponderado'].sum() if not df_pest.empty else 0
        
        # Contexto completo para la IA
        contexto_empresarial = {
            'nombre': empresa_info.get('nombre', 'La empresa'),
            'giro': empresa_info.get('giro', 'No especificado'),
            'estrategia_foda': estrategia_principal or 'No definida',
            'pest_score': pest_total,
            'num_estrategias': len(df_cmi),
            'postura': 'ofensiva' if 'Ofensiva' in str(estrategia_principal) else 
                      'defensiva' if 'Defensiva' in str(estrategia_principal) else 
                      'mixta'
        }
        
        # Mostrar contexto detectado
        with st.expander("📋 Contexto Estratégico Detectado (Automático)"):
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("Estrategia Principal", contexto_empresarial['estrategia_foda'])
            with col2:
                st.metric("Entorno PEST", f"{contexto_empresarial['pest_score']:.2f}")
            with col3:
                st.metric("Estrategias a Analizar", contexto_empresarial['num_estrategias'])
        
        st.divider()
        
        # Botón para generar análisis completo con IA
        if st.button("🤖 Generar Análisis Completo de Semaforización", type="primary"):
            with st.spinner("Analizando todas las estrategias con IA... Esto puede tomar 1-2 minutos"):
                
                resultados_semaforo = []
                
                # Procesar cada estrategia del CMI
                for idx, estrategia in df_cmi.iterrows():
                    
                    # Preparar prompt enriquecido con TODO el contexto disponible
                    prompt_analisis = f"""Actúa como un Director de Planeación Estratégica senior analizando el desempeño de una estrategia específica.

CONTEXTO EMPRESARIAL COMPLETO:
- Empresa: {contexto_empresarial['nombre']} ({contexto_empresarial['giro']})
- Postura Estratégica General: {contexto_empresarial['estrategia_foda']}
- Entorno PEST: {'Favorable' if contexto_empresarial['pest_score'] > 2.5 else 'Desafiante'} ({contexto_empresarial['pest_score']:.2f})
- Postura: {contexto_empresarial['postura']}

ESTRATEGIA A ANALIZAR:
- Nombre: {estrategia['estrategia']}
- Cuadrante FODA: {estrategia['cuadrante']}
- Importancia: {estrategia['importancia']}
- Plan Funcional: {estrategia['plan_asignado']}
- Actividades Clave: {estrategia['actividades']}

INSTRUCCIÓN:
Basándote ÚNICAMENTE en la información anterior (sin pedir datos adicionales), realiza un análisis experto que determine:

1. COLOR DEL SEMÁFORO (🔴🟡🟢):
   - Analiza la coherencia entre la importancia de la estrategia, su cuadrante FODA y el plan asignado
   - Considera: ¿Es una estrategia crítica para la postura {contexto_empresarial['postura']} de la empresa?
   - Evalúa: ¿Las actividades propuestas son suficientes y coherentes?
   - Determina: Dado el contexto PEST, ¿esta estrategia está bien posicionada?

2. JUSTIFICACIÓN DEL COLOR (máximo 150 palabras):
   Explica POR QUÉ cae en ese color considerando:
   - Alineación con la estrategia FODA principal
   - Factores PEST que afectan específicamente esta estrategia
   - Coherencia entre importancia y recursos asignados (plan funcional)
   - Viabilidad de las actividades propuestas
   - Riesgos implícitos en el cuadrante (FO=explotar, FA=defender, DO=reforzar, DA=sobrevivir)

3. DIAGNÓSTICO PROFUNDO:
   - Fortalezas de esta estrategia específica
   - Debilidades o riesgos detectados
   - Oportunidades de mejora inmediata

4. ACCIONES RECOMENDADAS:
   - Acción inmediata (esta semana)
   - Acción corto plazo (este mes)
   - Acción de seguimiento (trimestral)

5. IMPACTO ESTRATÉGICO:
   - ¿Qué pasa si esta estrategia falla?
   - ¿Qué otras estrategias se ven afectadas?
   - ¿Es crítica para el éxito del plan {estrategia['plan_asignado']}?

FORMATO DE RESPUESTA OBLIGATORIO:
COLOR: [🔴 ROJO / 🟡 AMARILLO / 🟢 VERDE]

JUSTIFICACIÓN:
[Texto explicativo coherente y específico]

DIAGNÓSTICO:
- Fortalezas: [lista]
- Debilidades: [lista]
- Oportunidades: [lista]

ACCIONES:
- Inmediata: [acción específica]
- Corto plazo: [acción específica]
- Seguimiento: [acción específica]

IMPACTO: [Alto/Medio/Bajo] - [explicación breve]

Genera el análisis ahora:"""

                    # Llamar a la IA
                    respuesta_ia = generar_analisis(prompt_analisis)
                    
                    # Parsear respuesta
                    color_detectado = '🟡'  # default
                    if 'COLOR:' in respuesta_ia or '🔴' in respuesta_ia or '🟢' in respuesta_ia or '🟡' in respuesta_ia:
                        if '🔴' in respuesta_ia or 'ROJO' in respuesta_ia.upper():
                            color_detectado = '🔴'
                        elif '🟢' in respuesta_ia or 'VERDE' in respuesta_ia.upper():
                            color_detectado = '🟢'
                        else:
                            color_detectado = '🟡'
                    
                    resultados_semaforo.append({
                        'estrategia_id': estrategia.get('id', idx),
                        'cuadrante': estrategia['cuadrante'],
                        'estrategia_nombre': estrategia['estrategia'],
                        'plan_asignado': estrategia['plan_asignado'],
                        'importancia': estrategia['importancia'],
                        'color': color_detectado,
                        'analisis_completo': respuesta_ia
                    })
                
                # Guardar en session state
                st.session_state['resultados_semaforo'] = resultados_semaforo
                st.success(f"✅ Análisis completado para {len(resultados_semaforo)} estrategias")
                st.rerun()
        
        # Mostrar resultados si existen
        if 'resultados_semaforo' in st.session_state:
            resultados = st.session_state['resultados_semaforo']
            
            # Dashboard resumen
            st.subheader("📊 Dashboard de Semaforización")
            
            conteo_colores = {'🔴': 0, '🟡': 0, '🟢': 0}
            for r in resultados:
                conteo_colores[r['color']] = conteo_colores.get(r['color'], 0) + 1
            
            total = len(resultados)
            
            col1, col2, col3, col4 = st.columns(4)
            with col1:
                st.metric("🟢 Óptimas", f"{conteo_colores['🟢']} ({conteo_colores['🟢']/total*100:.0f}%)")
            with col2:
                st.metric("🟡 Atención", f"{conteo_colores['🟡']} ({conteo_colores['🟡']/total*100:.0f}%)")
            with col3:
                st.metric("🔴 Críticas", f"{conteo_colores['🔴']} ({conteo_colores['🔴']/total*100:.0f}%)")
            with col4:
                riesgo = "ALTO" if conteo_colores['🔴'] > total*0.3 else "MEDIO" if conteo_colores['🔴'] > total*0.1 else "BAJO"
                st.metric("Nivel de Riesgo", riesgo)
            
            # Gráfico de distribución
            fig_colores = go.Figure(data=[go.Pie(
                labels=['Óptimas', 'Atención', 'Críticas'],
                values=[conteo_colores['🟢'], conteo_colores['🟡'], conteo_colores['🔴']],
                marker_colors=['#2ecc71', '#f1c40f', '#e74c3c'],
                hole=0.4
            )])
            fig_colores.update_layout(title="Distribución de Estrategias por Estado")
            st.plotly_chart(fig_colores, use_container_width=True)
            
            st.divider()
            
            # Detalle por estrategia
            st.subheader("📋 Análisis Detallado por Estrategia")
            
            # Filtros
            col_filtro1, col_filtro2 = st.columns(2)
            with col_filtro1:
                filtro_color = st.multiselect("Filtrar por color", ['🔴', '🟡', '🟢'], default=['🔴', '🟡', '🟢'])
            with col_filtro2:
                filtro_plan = st.multiselect("Filtrar por plan", 
                    df_cmi['plan_asignado'].unique().tolist() if not df_cmi.empty else [],
                    default=df_cmi['plan_asignado'].unique().tolist() if not df_cmi.empty else [])
            
            # Mostrar estrategias filtradas
            for resultado in resultados:
                if resultado['color'] not in filtro_color:
                    continue
                if resultado['plan_asignado'] not in filtro_plan:
                    continue
                
                with st.container():
                    # Encabezado con color
                    color_bg = {'🔴': '#ffebee', '🟡': '#fffde7', '🟢': '#e8f5e9'}[resultado['color']]
                    
                    st.markdown(f"""
                    <div style='background-color: {color_bg}; padding: 15px; border-radius: 10px; border-left: 5px solid {"#e74c3c" if resultado["color"]=="🔴" else "#f1c40f" if resultado["color"]=="🟡" else "#2ecc71"}; margin-bottom: 10px;'>
                        <h4>{resultado['color']} {resultado['cuadrante']} | {resultado['estrategia_nombre'][:60]}...</h4>
                        <p><strong>Plan:</strong> {resultado['plan_asignado']} | <strong>Importancia:</strong> {resultado['importancia']}</p>
                    </div>
                    """, unsafe_allow_html=True)
                    
                    with st.expander("Ver análisis completo de IA"):
                        st.markdown(resultado['analisis_completo'])
                        
                        # Botón para regenerar análisis específico
                        if st.button("🔄 Regenerar análisis", key=f"regen_{resultado['estrategia_id']}"):
                            # Eliminar este resultado específico y regenerar
                            st.session_state['regenerar_estrategia'] = resultado['estrategia_id']
                            st.rerun()
                    
                    st.divider()
            
            # Análisis ejecutivo general
            with st.expander("📄 Ver Análisis Ejecutivo General"):
                prompt_resumen = f"""Como Director Estratégico, genera un resumen ejecutivo de máximo 400 palabras basado en este análisis de semaforización:

DISTRIBUCIÓN DE ESTRATEGIAS:
- 🟢 Óptimas: {conteo_colores['🟢']} de {total}
- 🟡 Atención: {conteo_colores['🟡']} de {total}  
- 🔴 Críticas: {conteo_colores['🔴']} de {total}

CONTEXTO: {contexto_empresarial['nombre']} con postura {contexto_empresarial['postura']}

Incluye:
1. Diagnóstico general del portafolio de estrategias
2. Patrones detectados (¿algún plan funcional con problemas?)
3. Prioridades de acción inmediatas
4. Recomendación estratégica final

Lenguaje ejecutivo y directo."""
                
                if st.button("📊 Generar Resumen Ejecutivo con IA"):
                    with st.spinner("Generando..."):
                        resumen_ejecutivo = generar_analisis(prompt_resumen)
                        st.markdown("### 📋 Análisis Generado:")
                        st.info(resumen_ejecutivo)
                        
                        # Botón para guardar en BD
                        if st.button("💾 Guardar Resumen en Empresa"):
                            try:
                                supabase.table('empresas').update({
                                    'analisis_semaforo_resumen': resumen_ejecutivo
                                }).eq('id', empresa_id).execute()
                                st.success("Resumen guardado")
                            except Exception as e:
                                st.error(f"Error al guardar: {e}")
        
        else:
            # Vista previa/informativa antes de generar
            st.info("""
            ### 🎯 ¿Qué analizará la IA automáticamente?
            
            **Sin que tengas que ingresar nada más**, el sistema evaluará cada estrategia considerando:
            
            1. **🎯 Alineación FODA**: ¿La estrategia apoya la postura estratégica general?
            
            2. **🌍 Contexto PEST**: ¿El entorno externo favorece o dificulta esta estrategia?
            
            3. **⚖️ Coherencia**: ¿Hay coherencia entre la importancia declarada y el plan asignado?
            
            4. **📋 Viabilidad**: ¿Las actividades propuestas son suficientes y realistas?
            
            5. **⚠️ Riesgos implícitos**: Cada cuadrante FODA tiene riesgos característicos
            
            ### 📊 Output del análisis:
            - **Color del semáforo** para cada estrategia
            - **Justificación detallada** del por qué ese color
            - **Diagnóstico** de fortalezas/debilidades
            - **Acciones recomendadas** inmediatas
            - **Impacto estratégico** si falla
            
            Haz clic en **"Generar Análisis Completo"** arriba para comenzar.
            """)
            
            # Mostrar preview de estrategias a analizar
            if not df_cmi.empty:
                st.subheader("Estrategias que se analizarán:")
                preview_df = df_cmi[['cuadrante', 'estrategia', 'plan_asignado', 'importancia']].copy()
                preview_df.columns = ['Cuadrante', 'Estrategia', 'Plan', 'Importancia']
                st.dataframe(preview_df, use_container_width=True, hide_index=True)
    
    # --- PESTAÑA 7: OPERATIVIZACIÓN/PRESUPUESTO ---
    with tab7:
        st.header("Operativización / Presupuesto")
        
        # Obtener estrategias generadas
        df_estrategias_oper = get_datos_tabla('estrategias_generadas', empresa_id)
        
        if df_estrategias_oper.empty:
            st.warning("No hay estrategias disponibles. Genera estrategias primero en la pestaña 'Estrategia'.")
            st.stop()
        
        st.info(f"📋 Se encontraron {len(df_estrategias_oper)} estrategias para operativizar.")
        
        # Botón para generar/poblar el cuadro de operativización automáticamente
        col1, col2 = st.columns([1, 2])
        with col1:
            if st.button("🔄 Generar/Actualizar Cuadro de Operativización", type="primary", disabled=not puede_editar):
                with st.spinner("Desglosando estrategias en actividades..."):
                    try:
                        # Obtener operativización existente para preservar datos ya ingresados
                        df_oper_existente = get_datos_tabla('operativizacion', empresa_id)
                        datos_existentes = {}
                        if not df_oper_existente.empty:
                            for _, row in df_oper_existente.iterrows():
                                key = f"{row.get('estrategia_id', '')}_{row.get('numero_actividad', '')}"
                                datos_existentes[key] = {
                                    'plazo': row.get('plazo', ''),
                                    'responsable': row.get('responsable', ''),
                                    'costo': row.get('costo', 0)
                                }
                        
                        # Eliminar operativización anterior
                        supabase.table('operativizacion').delete().eq('empresa_id', empresa_id).execute()
                        
                        nuevas_actividades = []
                        
                        for _, estrategia in df_estrategias_oper.iterrows():
                            # Dividir las actividades de la estrategia
                            actividades_texto = estrategia.get('actividades', '')
                            import re
                            lista_actividades = re.split(r'[,;\n]+', actividades_texto)
                            lista_actividades = [act.strip() for act in lista_actividades if act.strip()]
                            
                            if not lista_actividades:
                                lista_actividades = ["Implementación general de la estrategia"]
                            
                            for num_act, actividad in enumerate(lista_actividades, 1):
                                key = f"{estrategia['id']}_{num_act}"
                                datos_previos = datos_existentes.get(key, {})
                                
                                nueva_act = {
                                    'empresa_id': empresa_id,
                                    'estrategia_id': estrategia['id'],
                                    'cuadrante': estrategia['cuadrante'],
                                    'estrategia_nombre': estrategia['estrategia'],
                                    'plan_asignado': estrategia['plan_asignado'],
                                    'numero_actividad': num_act,
                                    'descripcion_actividad': actividad,
                                    'plazo': datos_previos.get('plazo', ''),
                                    'responsable': datos_previos.get('responsable', ''),
                                    'costo': datos_previos.get('costo', 0),
                                    'importancia': estrategia.get('importancia', 'Media')
                                }
                                nuevas_actividades.append(nueva_act)
                        
                        if nuevas_actividades:
                            supabase.table('operativizacion').insert(nuevas_actividades).execute()
                            st.success(f"✅ {len(nuevas_actividades)} actividades generadas de {len(df_estrategias_oper)} estrategias.")
                            st.rerun()
                    
                    except Exception as e:
                        st.error(f"Error al generar operativización: {e}")
        
        with col2:
            st.caption("Este botón desglosa cada estrategia en sus actividades específicas para que solo edites plazo, responsable y costo.")
        
        st.divider()
        
        # Mostrar y editar el cuadro de operativización
        st.subheader("📋 Cuadro de Operativización por Actividades")
        
        df_oper = get_datos_tabla('operativizacion', empresa_id)
        
        if not df_oper.empty:
            # Calcular totales
            total_costo = pd.to_numeric(df_oper['costo'], errors='coerce').fillna(0).sum()
            total_actividades = len(df_oper)
            
            # Métricas resumen
            col_met1, col_met2, col_met3, col_met4 = st.columns(4)
            with col_met1:
                st.metric("💰 Costo Total", f"${total_costo:,.2f}")
            with col_met2:
                st.metric("📊 Total Actividades", total_actividades)
            with col_met3:
                st.metric("🎯 Estrategias", df_oper['estrategia_id'].nunique())
            with col_met4:
                promedio_costo = total_costo / total_actividades if total_actividades > 0 else 0
                st.metric("💵 Promedio por Actividad", f"${promedio_costo:,.2f}")
            
            st.divider()
            
            # Preparar datos para el editor
            df_oper['ref_estrategia'] = df_oper['cuadrante'] + " - " + df_oper['estrategia_nombre'].str[:50] + "..."
            
            # Ordenar por cuadrante y número de actividad
            orden_cuadrantes = {'FO': 1, 'FA': 2, 'DO': 3, 'DA': 4}
            df_oper['orden_cuadrante'] = df_oper['cuadrante'].map(orden_cuadrantes)
            df_oper = df_oper.sort_values(['orden_cuadrante', 'estrategia_id', 'numero_actividad'])
            
            # Seleccionar y renombrar columnas para mostrar
            columnas_mostrar = [
                'ref_estrategia', 'numero_actividad', 'descripcion_actividad', 
                'plazo', 'responsable', 'costo', 'plan_asignado', 'importancia'
            ]
            
            df_display = df_oper[columnas_mostrar].copy()
            df_display.columns = [
                'Estrategia (Referencia)', 'N°', 'Actividad', 
                'Plazo', 'Responsable', 'Costo ($)', 'Plan Asignado', 'Importancia'
            ]
            
            # Editor de datos
            st.write("**Edita directamente el Plazo, Responsable y Costo de cada actividad:**")
            
            edited_df = st.data_editor(
                df_display,
                num_rows="fixed",
                key="editor_oper_actividades",
                use_container_width=True,
                disabled=not puede_editar,
                column_config={
                    'Estrategia (Referencia)': st.column_config.TextColumn("Estrategia", disabled=True, width="medium"),
                    'N°': st.column_config.NumberColumn("N°", disabled=True, width="small"),
                    'Actividad': st.column_config.TextColumn("Descripción de la Actividad", disabled=True, width="large"),
                    'Plazo': st.column_config.TextColumn("⏱️ Plazo", width="medium"),
                    'Responsable': st.column_config.TextColumn("👤 Responsable", width="medium"),
                    'Costo ($)': st.column_config.NumberColumn("💰 Costo ($)", min_value=0, step=100, format="$%.2f", width="medium"),
                    'Plan Asignado': st.column_config.TextColumn("Plan", disabled=True, width="small"),
                    'Importancia': st.column_config.TextColumn("Importancia", disabled=True, width="small")
                },
                hide_index=True
            )
            
            # Botón para guardar cambios
            if st.button("💾 Guardar Cambios del Cuadro de Operativización", type="primary", disabled=not puede_editar):
                try:
                    # Reconstruir el dataframe completo con IDs
                    df_completo = df_oper.copy()
                    
                    # Actualizar solo las columnas editables
                    df_completo['plazo'] = edited_df['Plazo'].values
                    df_completo['responsable'] = edited_df['Responsable'].values
                    df_completo['costo'] = edited_df['Costo ($)'].values
                    
                    # Eliminar datos anteriores
                    supabase.table('operativizacion').delete().eq('empresa_id', empresa_id).execute()
                    
                    # Insertar datos actualizados
                    columnas_guardar = [
                        'empresa_id', 'estrategia_id', 'cuadrante', 'estrategia_nombre',
                        'plan_asignado', 'numero_actividad', 'descripcion_actividad',
                        'plazo', 'responsable', 'costo', 'importancia'
                    ]
                    
                    datos_guardar = df_completo[columnas_guardar].to_dict(orient='records')
                    supabase.table('operativizacion').insert(datos_guardar).execute()
                    
                    st.success("✅ Cuadro de operativización guardado correctamente.")
                    st.rerun()
                    
                except Exception as e:
                    st.error(f"Error al guardar: {e}")
            
            st.divider()
            
            # Vista resumen por estrategia
            with st.expander("📊 Ver Resumen por Estrategia"):
                resumen = df_oper.groupby(['cuadrante', 'estrategia_nombre']).agg({
                    'costo': 'sum',
                    'descripcion_actividad': 'count',
                    'plazo': lambda x: list(x.unique()),
                    'responsable': lambda x: list(x.dropna().unique())
                }).reset_index()
                
                resumen.columns = ['Cuadrante', 'Estrategia', 'Costo Total', 'N° Actividades', 'Plazos', 'Responsables']
                
                for _, row in resumen.iterrows():
                    with st.container():
                        col_r1, col_r2, col_r3 = st.columns([3, 1, 1])
                        with col_r1:
                            st.write(f"**{row['Cuadrante']}**: {row['Estrategia'][:60]}...")
                        with col_r2:
                            st.write(f"💰 ${row['Costo Total']:,.2f}")
                        with col_r3:
                            st.write(f"📋 {int(row['N° Actividades'])} act.")
                    st.divider()
            
            # Vista resumen por plan asignado
            with st.expander("📈 Ver Distribución por Plan"):
                fig = px.pie(
                    df_oper, 
                    names='plan_asignado', 
                    values='costo',
                    title='Distribución de Costos por Plan Estratégico',
                    hole=0.4
                )
                st.plotly_chart(fig, use_container_width=True)
                
                tabla_planes = df_oper.groupby('plan_asignado').agg({
                    'costo': 'sum',
                    'descripcion_actividad': 'count'
                }).reset_index()
                tabla_planes.columns = ['Plan', 'Costo Total', 'N° Actividades']
                st.dataframe(tabla_planes, use_container_width=True, hide_index=True)
        
        else:
            st.info("👆 Haz clic en 'Generar/Actualizar Cuadro de Operativización' para crear el cuadro con todas las estrategias y sus actividades.")
        
        st.divider()
        
        # Estado de Pérdidas y Ganancias del ÚLTIMO AÑO
        st.subheader("💰 Estado de Pérdidas y Ganancias (Último Año Real)")
        
        with st.expander("📊 Ingresar Datos del Último Año"):
            with st.form("form_pygn"):
                st.write("**Ingresa los datos reales del último año fiscal:**")
                
                col1, col2 = st.columns(2)
                with col1:
                    ingresos_ventas = st.number_input("Ingresos por Ventas", min_value=0.0, step=1000.0, key="pg_ingresos")
                    costos_ventas = st.number_input("Costos de Ventas (COGS)", min_value=0.0, step=1000.0, key="pg_costos")
                    gastos_operativos = st.number_input("Gastos Operativos", min_value=0.0, step=1000.0, key="pg_gop")
                with col2:
                    gastos_admin = st.number_input("Gastos Administrativos", min_value=0.0, step=1000.0, key="pg_gadm")
                    gastos_ventas_marketing = st.number_input("Gastos de Ventas y Marketing", min_value=0.0, step=1000.0, key="pg_gvta")
                    otros_ingresos = st.number_input("Otros Ingresos", min_value=0.0, step=1000.0, key="pg_oting")
                
                impuestos = st.number_input("Impuestos sobre la Renta", min_value=0.0, step=1000.0, key="pg_imp")
                utilidad_retenida = st.number_input("Utilidad Retenida del Año Anterior", min_value=0.0, step=1000.0, key="pg_uti_ret")
                
                submitted_pg = st.form_submit_button("💾 Guardar Estado P&G", disabled=not puede_editar)
            
            if submitted_pg:
                try:
                    utilidad_bruta = ingresos_ventas - costos_ventas
                    utilidad_operativa = utilidad_bruta - gastos_operativos - gastos_admin - gastos_ventas_marketing
                    utilidad_antes_impuestos = utilidad_operativa + otros_ingresos
                    utilidad_neta = utilidad_antes_impuestos - impuestos
                    
                    datos_pg = {
                        'empresa_id': empresa_id,
                        'ingresos_ventas': ingresos_ventas,
                        'costos_ventas': costos_ventas,
                        'gastos_operativos': gastos_operativos,
                        'gastos_administrativos': gastos_admin,
                        'gastos_ventas': gastos_ventas_marketing,
                        'otros_ingresos': otros_ingresos,
                        'impuestos': impuestos,
                        'utilidad_retenida_anterior': utilidad_retenida,
                        'utilidad_bruta': utilidad_bruta,
                        'utilidad_operativa': utilidad_operativa,
                        'utilidad_antes_impuestos': utilidad_antes_impuestos,
                        'utilidad_neta': utilidad_neta
                    }
                    
                    supabase.table('perdida_ganancia').delete().eq('empresa_id', empresa_id).execute()
                    supabase.table('perdida_ganancia').insert(datos_pg).execute()
                    st.success("✅ Estado de Pérdidas y Ganancias guardado.")
                    st.rerun()
                except Exception as e:
                    st.error(f"Error al guardar P&G: {e}")
        
        df_pg = get_datos_tabla('perdida_ganancia', empresa_id)
        if not df_pg.empty:
            st.write("**Resumen del Último Año:**")
            datos_pg = df_pg.iloc[0]
            
            col1, col2, col3, col4 = st.columns(4)
            with col1:
                st.metric("Ingresos", f"${datos_pg['ingresos_ventas']:,.2f}")
            with col2:
                st.metric("Utilidad Bruta", f"${datos_pg['utilidad_bruta']:,.2f}")
            with col3:
                st.metric("Utilidad Operativa", f"${datos_pg['utilidad_operativa']:,.2f}")
            with col4:
                st.metric("Utilidad Neta", f"${datos_pg['utilidad_neta']:,.2f}")
        
        st.divider()
        
        # PROYECCIÓN con Análisis Costo-Beneficio
        st.subheader("📈 Proyección y Análisis Costo-Beneficio")
        
        if not df_pg.empty and not df_oper.empty:
            datos_base = df_pg.iloc[0]
            
            with st.form("form_proyeccion"):
                st.write("**Parámetros de Proyección:**")
                
                col1, col2, col3 = st.columns(3)
                with col1:
                    crecimiento_ventas = st.slider("Crecimiento en Ventas (%)", -50, 100, 10, key="proj_crec") / 100
                    reduccion_costos = st.slider("Reducción de Costos (%)", 0, 50, 5, key="proj_red") / 100
                with col2:
                    periodos_proyeccion = st.selectbox("Período de Proyección", ["1 año", "2 años", "3 años", "5 años"], key="proj_per")
                    unidades_proyectadas = st.number_input("Unidades Totales Proyectadas", min_value=1, value=1000, key="proj_unid")
                with col3:
                    inversion_total = st.number_input("Inversión Total Requerida ($)", min_value=0.0, value=float(total_costo), key="proj_inv")
                    tasa_descuento = st.slider("Tasa de Descuento Anual (%)", 0, 30, 10, key="proj_tasa") / 100
                
                submitted_proj = st.form_submit_button("🚀 Calcular Proyección y Costo-Beneficio", disabled=not puede_editar)
            
            # El procesamiento va FUERA del with st.form()
            if submitted_proj:
                try:
                    anios = int(periodos_proyeccion.split()[0])
                    
                    proyecciones = []
                    ingreso_actual = datos_base['ingresos_ventas']
                    costo_actual = datos_base['costos_ventas']
                    
                    flujos_futuros = []
                    
                    for anio_num in range(1, anios + 1):
                        ingreso_proy = ingreso_actual * ((1 + crecimiento_ventas) ** anio_num)
                        costo_proy = costo_actual * ((1 - reduccion_costos) ** anio_num) * ((1 + crecimiento_ventas) ** anio_num)
                        utilidad_bruta_proy = ingreso_proy - costo_proy
                        gastos_proy = (datos_base['gastos_operativos'] + datos_base['gastos_administrativos'] + datos_base['gastos_ventas']) * ((1 + crecimiento_ventas * 0.5) ** anio_num)
                        utilidad_neta_proy = utilidad_bruta_proy - gastos_proy - (utilidad_bruta_proy * 0.25)
                        
                        flujos_futuros.append(utilidad_neta_proy)
                        
                        proyecciones.append({
                            'anio': anio_num,
                            'ingresos_proyectados': ingreso_proy,
                            'costos_proyectados': costo_proy,
                            'utilidad_neta_proyectada': utilidad_neta_proy
                        })
                    
                    df_proyeccion = pd.DataFrame(proyecciones)
                    
                    # Guardar proyección
                    supabase.table('proyeccion_financiera').delete().eq('empresa_id', empresa_id).execute()
                    for _, row in df_proyeccion.iterrows():
                        row_dict = {
                            'empresa_id': empresa_id,
                            'anio': int(row['anio']),
                            'ingresos_proyectados': float(row['ingresos_proyectados']),
                            'costos_proyectados': float(row['costos_proyectados']),
                            'utilidad_neta_proyectada': float(row['utilidad_neta_proyectada'])
                        }
                        supabase.table('proyeccion_financiera').insert(row_dict).execute()
                    
                    # Calcular Costo-Beneficio
                    vpn = sum([f / ((1 + tasa_descuento) ** (i+1)) for i, f in enumerate(flujos_futuros)])
                    
                    if inversion_total > 0:
                        relacion_cb_dolares = vpn / inversion_total
                    else:
                        relacion_cb_dolares = float('inf') if vpn > 0 else 0
                    
                    flujo_acumulado = 0
                    payback_anios = None
                    for i, flujo in enumerate(flujos_futuros):
                        flujo_acumulado += flujo
                        if flujo_acumulado >= inversion_total:
                            exceso = flujo_acumulado - inversion_total
                            fraccion_anio = 1 - (exceso / flujo) if flujo > 0 else 0
                            payback_anios = i + fraccion_anio
                            break
                    
                    if unidades_proyectadas > 0 and inversion_total > 0:
                        beneficio_por_unidad = vpn / unidades_proyectadas
                        costo_por_unidad = inversion_total / unidades_proyectadas
                        relacion_cb_unidades = beneficio_por_unidad / costo_por_unidad if costo_por_unidad > 0 else 0
                    else:
                        beneficio_por_unidad = 0
                        costo_por_unidad = 0
                        relacion_cb_unidades = 0
                    
                    # Guardar análisis CB
                    analisis_cb = {
                        'empresa_id': empresa_id,
                        'inversion_total': float(inversion_total),
                        'vpn_total': float(vpn),
                        'relacion_costo_beneficio_dolares': float(relacion_cb_dolares),
                        'payback_periodo_anios': float(payback_anios) if payback_anios else None,
                        'unidades_proyectadas': int(unidades_proyectadas),
                        'beneficio_por_unidad': float(beneficio_por_unidad),
                        'costo_por_unidad': float(costo_por_unidad),
                        'relacion_cb_unidades': float(relacion_cb_unidades),
                        'tasa_descuento': float(tasa_descuento)
                    }
                    
                    supabase.table('analisis_costo_beneficio').delete().eq('empresa_id', empresa_id).execute()
                    supabase.table('analisis_costo_beneficio').insert(analisis_cb).execute()
                    
                    st.success("✅ Proyección y análisis Costo-Beneficio calculados y guardados.")
                    st.session_state['mostrar_resultados_cb'] = True
                    st.rerun()
                    
                except Exception as e:
                    st.error(f"Error en el cálculo: {e}")
            
            # Mostrar resultados si existen
            if st.session_state.get('mostrar_resultados_cb', False):
                df_cb = get_datos_tabla('analisis_costo_beneficio', empresa_id)
                df_proy = get_datos_tabla('proyeccion_financiera', empresa_id)
                
                if not df_cb.empty and not df_proy.empty:
                    datos_cb = df_cb.iloc[0]
                    
                    st.divider()
                    st.subheader("📊 Resultados del Análisis Costo-Beneficio")
                    
                    # Renombrar columnas para mostrar
                    df_proy_display = df_proy.rename(columns={
                        'anio': 'Año',
                        'ingresos_proyectados': 'Ingresos Proyectados',
                        'costos_proyectados': 'Costos Proyectados',
                        'utilidad_neta_proyectada': 'Utilidad Neta Proyectada'
                    })
                    
                    st.write("**Proyección Financiera:**")
                    st.dataframe(df_proy_display, use_container_width=True, hide_index=True)
                    
                    fig = go.Figure()
                    fig.add_trace(go.Scatter(x=df_proy['anio'], y=df_proy['ingresos_proyectados'], 
                                           mode='lines+markers', name='Ingresos', line=dict(color='green')))
                    fig.add_trace(go.Scatter(x=df_proy['anio'], y=df_proy['costos_proyectados'], 
                                           mode='lines+markers', name='Costos', line=dict(color='red')))
                    fig.add_trace(go.Scatter(x=df_proy['anio'], y=df_proy['utilidad_neta_proyectada'], 
                                           mode='lines+markers', name='Utilidad Neta', line=dict(color='blue')))
                    fig.update_layout(title="Proyección Financiera", xaxis_title="Año", yaxis_title="Monto ($)")
                    st.plotly_chart(fig, use_container_width=True)
                    
                    st.write("**Indicadores Costo-Beneficio:**")
                    
                    col1, col2, col3 = st.columns(3)
                    
                    with col1:
                        st.markdown("### 💵 En DÓLARES")
                        st.metric("Relación C-B", f"{float(datos_cb['relacion_costo_beneficio_dolares']):.2f}")
                        st.write(f"**Inversión:** ${float(datos_cb['inversion_total']):,.2f}")
                        st.write(f"**VPN Total:** ${float(datos_cb['vpn_total']):,.2f}")
                        if datos_cb['relacion_costo_beneficio_dolares'] >= 1:
                            st.success(f"✅ Por cada $1 invertido, se recuperan ${float(datos_cb['relacion_costo_beneficio_dolares']):.2f}")
                        else:
                            st.error(f"❌ Por cada $1 invertido, solo se recuperan ${float(datos_cb['relacion_costo_beneficio_dolares']):.2f}")
                    
                    with col2:
                        st.markdown("### ⏱️ En TIEMPO")
                        if datos_cb['payback_periodo_anios']:
                            st.metric("Periodo de Recuperación", f"{float(datos_cb['payback_periodo_anios']):.1f} años")
                            if datos_cb['payback_periodo_anios'] <= 2:
                                st.success("✅ Recuperación rápida")
                            elif datos_cb['payback_periodo_anios'] <= 5:
                                st.info("ℹ️ Recuperación moderada")
                            else:
                                st.warning("⚠️ Recuperación lenta")
                        else:
                            st.metric("Periodo de Recuperación", "No recuperable")
                            st.error("❌ Inversión no recuperable")
                    
                    with col3:
                        st.markdown("### 📦 En UNIDADES")
                        st.metric("Beneficio/Unidad", f"${float(datos_cb['beneficio_por_unidad']):,.2f}")
                        st.metric("Costo/Unidad", f"${float(datos_cb['costo_por_unidad']):,.2f}")
                        st.metric("Relación C-B", f"{float(datos_cb['relacion_cb_unidades']):.2f}")
                        if datos_cb['relacion_cb_unidades'] >= 1:
                            st.success("✅ Rentable por unidad")
                        else:
                            st.error("❌ No rentable por unidad")
                    
                    st.divider()
                    st.subheader("📝 Interpretación Ejecutiva")
                    
                    interpretaciones = []
                    if datos_cb['relacion_costo_beneficio_dolares'] >= 1.5:
                        interpretaciones.append("**Rentabilidad Excelente:** Retorno significativo de la inversión.")
                    elif datos_cb['relacion_costo_beneficio_dolares'] >= 1:
                        interpretaciones.append("**Rentabilidad Aceptable:** El proyecto es viable.")
                    else:
                        interpretaciones.append("**Alerta:** El proyecto no es viable financieramente.")
                    
                    if datos_cb['payback_periodo_anios'] and datos_cb['payback_periodo_anios'] <= 3:
                        interpretaciones.append("**Liquidez:** Recuperación rápida del capital.")
                    
                    for interp in interpretaciones:
                        st.write(f"• {interp}")
                    
                    puntos_positivos = sum([
                        datos_cb['relacion_costo_beneficio_dolares'] >= 1,
                        datos_cb['payback_periodo_anios'] is not None and datos_cb['payback_periodo_anios'] <= 5,
                        datos_cb['relacion_cb_unidades'] >= 1
                    ])
                    
                    if puntos_positivos >= 2:
                        st.success("### ✅ RECOMENDACIÓN: APROBAR PROYECTO")
                    elif puntos_positivos == 1:
                        st.warning("### ⚠️ RECOMENDACIÓN: EVALUAR CON PRECAUCIÓN")
                    else:
                        st.error("### ❌ RECOMENDACIÓN: RECHAZAR PROYECTO")
        
        else:
            st.warning("Completa el Estado de Pérdidas y Ganancias y genera el Cuadro de Operativización para realizar la proyección.")

    
    # --- PESTAÑA DASHBOARD: BUSINESS INTELLIGENCE INTEGRADO ---
    with tab8:
        st.header("📊 Dashboard Ejecutivo - Business Intelligence")
        st.markdown("*Análisis interactivo tipo Power BI con filtros dinámicos y visualizaciones en tiempo real*")
        
        # ============================================================
        # 1. CARGA DE TODOS LOS DATOS DISPONIBLES
        # ============================================================
        
        @st.cache_data(ttl=300)
        def cargar_todos_los_datos(empresa_id):
            datos = {}
            datos['empresa'] = get_datos_empresa(empresa_id)
            datos['estrategias'] = get_datos_tabla('estrategias_generadas', empresa_id)
            datos['foda'] = get_datos_tabla('foda_cruzado', empresa_id)
            datos['pest'] = get_datos_tabla('matrices', empresa_id, tipo_matriz_filter='PEST')
            datos['operativizacion'] = get_datos_tabla('operativizacion', empresa_id)
            datos['pg'] = get_datos_tabla('perdida_ganancia', empresa_id)
            datos['proyeccion'] = get_datos_tabla('proyeccion_financiera', empresa_id)
            datos['cmi'] = get_datos_tabla('estrategias_generadas', empresa_id)
            return datos
        
        # Verificar que empresa_id existe
        if not empresa_id:
            st.error("No se ha seleccionado una empresa")
            st.stop()
            
        datos = cargar_todos_los_datos(empresa_id)
        
        # ============================================================
        # 2. PANEL DE FILTROS GLOBALES (Tipo Power BI)
        # ============================================================
        
        st.subheader("🎛️ Filtros Globales")
        
        col_f1, col_f2, col_f3, col_f4 = st.columns(4)
        
        with col_f1:
            # Filtro por Cuadrante FODA
            if not datos['estrategias'].empty and 'cuadrante' in datos['estrategias'].columns:
                cuadrantes = ['Todos'] + datos['estrategias']['cuadrante'].unique().tolist()
            else:
                cuadrantes = ['Todos']
            filtro_cuadrante = st.selectbox("📍 Cuadrante FODA", cuadrantes, key="dash_cuadrante")
        
        with col_f2:
            # Filtro por Plan Funcional
            if not datos['estrategias'].empty and 'plan_asignado' in datos['estrategias'].columns:
                planes = ['Todos'] + datos['estrategias']['plan_asignado'].unique().tolist()
            else:
                planes = ['Todos']
            filtro_plan = st.selectbox("📋 Plan Funcional", planes, key="dash_plan")
        
        with col_f3:
            # Filtro por Importancia
            if not datos['estrategias'].empty and 'importancia' in datos['estrategias'].columns:
                importancias = ['Todas'] + datos['estrategias']['importancia'].unique().tolist()
            else:
                importancias = ['Todas']
            filtro_importancia = st.selectbox("⚡ Importancia", importancias, key="dash_importancia")
        
        with col_f4:
            # Filtro por Categoría PEST
            if not datos['pest'].empty and 'categoria' in datos['pest'].columns:
                categorias_pest = ['Todas'] + datos['pest']['categoria'].unique().tolist()
            else:
                categorias_pest = ['Todas']
            filtro_pest = st.selectbox("🌍 Categoría PEST", categorias_pest, key="dash_pest")
        
        # Filtros adicionales expandibles
        with st.expander("🔍 Filtros Avanzados"):
            col_f5, col_f6 = st.columns(2)
            with col_f5:
                # Rango de impacto FODA
                if not datos['foda'].empty and 'impacto' in datos['foda'].columns:
                    impacto_max_val = int(datos['foda']['impacto'].max()) if not datos['foda']['impacto'].isna().all() else 10
                    impacto_min_val = int(datos['foda']['impacto'].min()) if not datos['foda']['impacto'].isna().all() else 0
                    impacto_min, impacto_max = st.slider(
                        "Rango de Impacto FODA", 
                        min_value=impacto_min_val, 
                        max_value=max(impacto_max_val, 10), 
                        value=(impacto_min_val, impacto_max_val),
                        key="dash_impacto"
                    )
                else:
                    impacto_min, impacto_max = 0, 10
                    st.slider("Rango de Impacto FODA", 0, 10, (0, 10), key="dash_impacto_disabled", disabled=True)
            
            with col_f6:
                # Filtro por texto
                filtro_texto = st.text_input("🔎 Buscar en estrategias", placeholder="Escribe palabra clave...", key="dash_texto")
        
        # Botón para aplicar filtros (simula refresh de Power BI)
        col_btn1, col_btn2, col_btn3 = st.columns([1, 1, 2])
        with col_btn1:
            aplicar_filtros = st.button("🔄 Aplicar Filtros", type="primary", key="dash_aplicar")
        with col_btn2:
            if st.button("❌ Limpiar Filtros", key="dash_limpiar"):
                st.rerun()
        
        st.divider()
        
        # ============================================================
        # 3. APLICACIÓN DE FILTROS Y CÁLCULO DE MÉTRICAS
        # ============================================================
        
        # Filtrar estrategias
        df_estrategias_filtrado = datos['estrategias'].copy() if not datos['estrategias'].empty else pd.DataFrame()
        
        if not df_estrategias_filtrado.empty:
            if filtro_cuadrante != 'Todos' and 'cuadrante' in df_estrategias_filtrado.columns:
                df_estrategias_filtrado = df_estrategias_filtrado[df_estrategias_filtrado['cuadrante'] == filtro_cuadrante]
            
            if filtro_plan != 'Todos' and 'plan_asignado' in df_estrategias_filtrado.columns:
                df_estrategias_filtrado = df_estrategias_filtrado[df_estrategias_filtrado['plan_asignado'] == filtro_plan]
            
            if filtro_importancia != 'Todas' and 'importancia' in df_estrategias_filtrado.columns:
                df_estrategias_filtrado = df_estrategias_filtrado[df_estrategias_filtrado['importancia'] == filtro_importancia]
            
            if filtro_texto and 'estrategia' in df_estrategias_filtrado.columns:
                mask = df_estrategias_filtrado['estrategia'].str.contains(filtro_texto, case=False, na=False)
                df_estrategias_filtrado = df_estrategias_filtrado[mask]
        
        # Filtrar PEST
        df_pest_filtrado = datos['pest'].copy() if not datos['pest'].empty else pd.DataFrame()
        
        if not df_pest_filtrado.empty:
            if filtro_pest != 'Todas' and 'categoria' in df_pest_filtrado.columns:
                df_pest_filtrado = df_pest_filtrado[df_pest_filtrado['categoria'] == filtro_pest]
        
        # ============================================================
        # 4. KPIs PRINCIPALES (Tarjetas tipo Power BI)
        # ============================================================
        
        st.subheader("📈 Indicadores Clave de Desempeño (KPIs)")
        
        col_kpi1, col_kpi2, col_kpi3, col_kpi4, col_kpi5 = st.columns(5)
        
        # KPI 1: Total Estrategias Filtradas
        with col_kpi1:
            total_est = len(df_estrategias_filtrado) if not df_estrategias_filtrado.empty else 0
            total_est_original = len(datos['estrategias']) if not datos['estrategias'].empty else 0
            delta_est = total_est_original - total_est
            st.metric(
                label="🎯 Estrategias",
                value=total_est,
                delta=f"-{delta_est} filtradas" if delta_est > 0 else "Todas",
                delta_color="off"
            )
        
        # KPI 2: Inversión Total Operativizada
        with col_kpi2:
            inversion_total = 0
            if not datos['operativizacion'].empty and 'costo' in datos['operativizacion'].columns:
                df_op = datos['operativizacion'].copy()
                if filtro_plan != 'Todos' and 'plan_asignado' in df_op.columns:
                    df_op = df_op[df_op['plan_asignado'] == filtro_plan]
                inversion_total = pd.to_numeric(df_op['costo'], errors='coerce').sum()
            st.metric(
                label="💰 Inversión Total",
                value=f"${inversion_total:,.0f}",
                delta="Presupuesto" if inversion_total > 0 else "Sin datos"
            )
        
        # KPI 3: Score PEST Promedio
        with col_kpi3:
            score_pest = 0
            if not df_pest_filtrado.empty and 'valor_ponderado' in df_pest_filtrado.columns:
                score_pest = pd.to_numeric(df_pest_filtrado['valor_ponderado'], errors='coerce').sum()
            st.metric(
                label="🌍 Score PEST",
                value=f"{score_pest:.2f}",
                delta="Favorable" if score_pest > 2.5 else "Desafiante",
                delta_color="normal" if score_pest > 2.5 else "inverse"
            )
        
        # KPI 4: Estrategias de Alta Importancia
        with col_kpi4:
            altas = 0
            if not df_estrategias_filtrado.empty and 'importancia' in df_estrategias_filtrado.columns:
                altas = len(df_estrategias_filtrado[df_estrategias_filtrado['importancia'] == 'Alta'])
            st.metric(
                label="⚡ Alta Prioridad",
                value=altas,
                delta="Requieren atención" if altas > 0 else "Sin alertas",
                delta_color="inverse" if altas > 5 else "normal"
            )
        
        # KPI 5: Rentabilidad Esperada
        with col_kpi5:
            utilidad_total = 0
            if not datos['proyeccion'].empty and 'utilidad_neta_proyectada' in datos['proyeccion'].columns:
                utilidad_total = pd.to_numeric(datos['proyeccion']['utilidad_neta_proyectada'], errors='coerce').sum()
            st.metric(
                label="📈 Utilidad Proy.",
                value=f"${utilidad_total:,.0f}",
                delta="5 años" if utilidad_total > 0 else "Sin proyección"
            )
        
        st.divider()
        
        # ============================================================
        # 5. GRÁFICOS PRINCIPALES (Layout tipo Power BI)
        # ============================================================
        
        # Fila 1: Gráficos grandes
        col_graf1, col_graf2 = st.columns([2, 1])
        
        with col_graf1:
            st.subheader("🎯 Distribución de Estrategias")
            
            if not df_estrategias_filtrado.empty and 'cuadrante' in df_estrategias_filtrado.columns and 'plan_asignado' in df_estrategias_filtrado.columns:
                try:
                    fig_burbuja = px.scatter(
                        df_estrategias_filtrado,
                        x='cuadrante',
                        y='plan_asignado',
                        color='cuadrante',
                        hover_data=['estrategia'] if 'estrategia' in df_estrategias_filtrado.columns else None,
                        size_max=60,
                        title="Mapa de Estrategias por Cuadrante y Plan"
                    )
                    fig_burbuja.update_layout(height=400)
                    st.plotly_chart(fig_burbuja, use_container_width=True)
                except Exception as e:
                    st.error(f"Error al generar gráfico: {e}")
            else:
                st.info("No hay datos suficientes de estrategias")
        
        with col_graf2:
            st.subheader("📊 Matriz de Importancia")
            
            if not df_estrategias_filtrado.empty and 'importancia' in df_estrategias_filtrado.columns:
                try:
                    imp_count = df_estrategias_filtrado['importancia'].value_counts()
                    fig_pastel = px.pie(
                        values=imp_count.values,
                        names=imp_count.index,
                        title="Distribución por Prioridad",
                        hole=0.4
                    )
                    fig_pastel.update_layout(height=400)
                    st.plotly_chart(fig_pastel, use_container_width=True)
                except Exception as e:
                    st.error(f"Error al generar gráfico: {e}")
            else:
                st.info("No hay datos de importancia")
        
        st.divider()
        
        # Fila 2: Análisis PEST y FODA
        col_graf3, col_graf4 = st.columns(2)
        
        with col_graf3:
            st.subheader("🌍 Análisis PEST Interactivo")
            
            if not df_pest_filtrado.empty and 'categoria' in df_pest_filtrado.columns and 'valor_ponderado' in df_pest_filtrado.columns:
                try:
                    pest_agrupado = df_pest_filtrado.groupby('categoria')['valor_ponderado'].sum().reset_index()
                    fig_pest = px.bar(
                        pest_agrupado,
                        x='categoria',
                        y='valor_ponderado',
                        color='categoria',
                        title="Impacto por Categoría PEST"
                    )
                    fig_pest.update_layout(height=350, showlegend=False)
                    st.plotly_chart(fig_pest, use_container_width=True)
                except Exception as e:
                    st.error(f"Error al generar gráfico PEST: {e}")
            else:
                st.info("No hay datos PEST disponibles")
        
        with col_graf4:
            st.subheader("⚔️ Análisis FODA Cruzado")
            
            if not datos['foda'].empty and 'cuadrante' in datos['foda'].columns and 'impacto' in datos['foda'].columns:
                try:
                    foda_pivot = datos['foda'].groupby('cuadrante')['impacto'].sum().reset_index()
                    
                    colores_foda = {'FO': '#2ecc71', 'FA': '#e74c3c', 'DO': '#3498db', 'DA': '#f39c12'}
                    colores_barras = [colores_foda.get(c, '#95a5a6') for c in foda_pivot['cuadrante']]
                    
                    fig_foda = go.Figure(data=go.Bar(
                        x=foda_pivot['cuadrante'],
                        y=foda_pivot['impacto'],
                        marker_color=colores_barras,
                        text=foda_pivot['impacto'],
                        textposition='auto'
                    ))
                    fig_foda.update_layout(
                        title="Puntaje por Cuadrante FODA",
                        height=350
                    )
                    st.plotly_chart(fig_foda, use_container_width=True)
                except Exception as e:
                    st.error(f"Error al generar gráfico FODA: {e}")
            else:
                st.info("No hay datos FODA disponibles")
        
        st.divider()
        
        # ============================================================
        # 6. TABLA DINÁMICA TIPO EXCEL/POWER BI
        # ============================================================
        
        st.subheader("📋 Tabla Dinámica de Estrategias")
        
        if not df_estrategias_filtrado.empty:
            # Seleccionar columnas disponibles
            columnas_disponibles = []
            for col in ['cuadrante', 'estrategia', 'plan_asignado', 'importancia', 'actividades']:
                if col in df_estrategias_filtrado.columns:
                    columnas_disponibles.append(col)
            
            if columnas_disponibles:
                st.dataframe(
                    df_estrategias_filtrado[columnas_disponibles],
                    use_container_width=True,
                    hide_index=True
                )
                
                # Botones de exportación
                col_exp1, col_exp2 = st.columns(2)
                with col_exp1:
                    try:
                        csv = df_estrategias_filtrado.to_csv(index=False).encode('utf-8')
                        st.download_button(
                            label="📥 Descargar CSV",
                            data=csv,
                            file_name=f"estrategias_{empresa_id}.csv",
                            mime="text/csv",
                            key="dash_csv"
                        )
                    except Exception as e:
                        st.error(f"Error al generar CSV: {e}")
                
                with col_exp2:
                    try:
                        buffer = io.BytesIO()
                        with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
                            df_estrategias_filtrado.to_excel(writer, sheet_name='Estrategias', index=False)
                        
                        st.download_button(
                            label="📊 Descargar Excel",
                            data=buffer.getvalue(),
                            file_name=f"estrategias_{empresa_id}.xlsx",
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                            key="dash_excel"
                        )
                    except Exception as e:
                        st.warning("Instala 'openpyxl' para exportar a Excel: pip install openpyxl")
            else:
                st.warning("No hay columnas disponibles para mostrar")
        else:
            st.warning("No hay estrategias que coincidan con los filtros")
        
        st.divider()
        
        # ============================================================
        # 7. ANÁLISIS DE PRESUPUESTO Y COSTOS
        # ============================================================
        
        if not datos['operativizacion'].empty and 'costo' in datos['operativizacion'].columns:
            st.subheader("💰 Análisis de Presupuesto")
            
            col_pres1, col_pres2 = st.columns(2)
            
            with col_pres1:
                if 'plan_asignado' in datos['operativizacion'].columns:
                    try:
                        presupuesto_plan = datos['operativizacion'].groupby('plan_asignado')['costo'].sum().reset_index()
                        fig_pres = px.bar(
                            presupuesto_plan,
                            x='plan_asignado',
                            y='costo',
                            color='plan_asignado',
                            title="Inversión por Plan Funcional"
                        )
                        fig_pres.update_layout(height=350, showlegend=False)
                        st.plotly_chart(fig_pres, use_container_width=True)
                    except Exception as e:
                        st.error(f"Error: {e}")
            
            with col_pres2:
                if 'descripcion_actividad' in datos['operativizacion'].columns:
                    try:
                        top_actividades = datos['operativizacion'].nlargest(10, 'costo')
                        if 'plan_asignado' in top_actividades.columns:
                            fig_top = px.bar(
                                top_actividades,
                                y='descripcion_actividad',
                                x='costo',
                                orientation='h',
                                color='plan_asignado',
                                title="Top 10 Actividades más Costosas"
                            )
                            fig_top.update_layout(height=350)
                            st.plotly_chart(fig_top, use_container_width=True)
                    except Exception as e:
                        st.error(f"Error: {e}")
        
        # ============================================================
        # 8. PROYECCIÓN FINANCIERA
        # ============================================================
        
        if not datos['proyeccion'].empty:
            cols_necesarios = ['anio', 'ingresos_proyectados', 'costos_proyectados', 'utilidad_neta_proyectada']
            if all(col in datos['proyeccion'].columns for col in cols_necesarios):
                st.subheader("📈 Proyección Financiera")
                
                try:
                    fig_proy = go.Figure()
                    
                    fig_proy.add_trace(go.Scatter(
                        x=datos['proyeccion']['anio'],
                        y=datos['proyeccion']['ingresos_proyectados'],
                        mode='lines+markers',
                        name='Ingresos',
                        line=dict(color='#2ecc71', width=3)
                    ))
                    
                    fig_proy.add_trace(go.Scatter(
                        x=datos['proyeccion']['anio'],
                        y=datos['proyeccion']['costos_proyectados'],
                        mode='lines+markers',
                        name='Costos',
                        line=dict(color='#e74c3c', width=3)
                    ))
                    
                    fig_proy.add_trace(go.Scatter(
                        x=datos['proyeccion']['anio'],
                        y=datos['proyeccion']['utilidad_neta_proyectada'],
                        mode='lines+markers',
                        name='Utilidad Neta',
                        line=dict(color='#3498db', width=4)
                    ))
                    
                    fig_proy.update_layout(
                        title="Proyección a 5 Años",
                        xaxis_title="Año",
                        yaxis_title="Monto ($)",
                        height=450,
                        hovermode='x unified'
                    )
                    
                    st.plotly_chart(fig_proy, use_container_width=True)
                except Exception as e:
                    st.error(f"Error en proyección: {e}")
        
        # ============================================================
        # 9. INSIGHTS CON IA
        # ============================================================
        
        st.divider()
        st.subheader("🤖 Insights y Recomendaciones")
        
        if st.button("💡 Generar Análisis con IA", key="dash_ia"):
            with st.spinner("Analizando datos..."):
                # Preparar resumen
                resumen_datos = f"""
Estrategias: {total_est}
Inversión: ${inversion_total:,.0f}
Score PEST: {score_pest:.2f}
Alta prioridad: {altas}
"""
                
                prompt_insights = f"""Como Director Ejecutivo, analiza estos datos y genera:
1. INSIGHT PRINCIPAL (hallazgo clave)
2. ALERTAS (máximo 2)
3. OPORTUNIDADES (máximo 2)
4. RECOMENDACIÓN PRIORITARIA

Datos: {resumen_datos}

Sé directo y accionable. Máximo 200 palabras."""
                
                try:
                    insights = generar_analisis(prompt_insights)
                    st.markdown("### 📋 Análisis Generado:")
                    st.info(insights)
                except Exception as e:
                    st.error(f"Error al generar análisis: {e}")    

    # --- PESTAÑA 9: RESUMEN Y CONCLUSIONES ---
    with tab9:
        st.header("Resumen, Conclusiones y Exportación")
    
        st.subheader("📄 Generar Documento Final (Formato APA)")
        
        with st.form("pdf_form"):
            col1, col2 = st.columns(2)
            with col1:
                pdf_version = st.text_input("Versión del Plan", value="1.0")
                pdf_elaborado = st.text_input("Elaborado por", value="Consultor Estratégico")
            with col2:
                pdf_revisado = st.text_input("Revisado por", value="Director de Planeación")
                pdf_aprobado = st.text_input("Aprobado por", value="Director General")
            
            st.info(
                "**Estructura del documento generado:**\n\n"
                "- **Resumen Ejecutivo** "
                "- **Plan Estratégico** "
                "- **Anexos** "
            )
            
            submitted_pdf = st.form_submit_button("🚀 Generar PDF Profesional")
        
        # El procesamiento va FUERA del with st.form()
        if submitted_pdf:
            with st.spinner("Generando documento con formato APA. Esto puede tomar un momento..."):
                # Generar PDF con formato profesional
                pdf_buffer = generar_pdf_completo_mejorado(
                    empresa_id, 
                    pdf_version, 
                    pdf_elaborado, 
                    pdf_revisado, 
                    pdf_aprobado
                )

                if pdf_buffer:
                    # Guardar PDF
                    st.session_state['pdf_bytes'] = pdf_buffer.getvalue()
                    st.session_state['pdf_nombre'] = f"Plan_Estrategico_{empresa_data.get('nombre', 'Empresa')}_V{pdf_version}.pdf"

                    st.session_state['pdf_generado'] = True
                    st.success("✅ Documento PDF generado correctamente.")
        
        # Mostrar botones de descarga si el documento fue generado
        if st.session_state.get('pdf_generado', False) and 'pdf_bytes' in st.session_state:
            col1, col2, col3 = st.columns([1, 1, 2])
            with col1:
                # Botón descargar PDF
                download_buffer = BytesIO(st.session_state['pdf_bytes'])
                st.download_button(
                    label="⬇️ Descargar PDF", 
                    data=download_buffer, 
                    file_name=st.session_state.get('pdf_nombre', 'plan_estrategico.pdf'), 
                    mime="application/pdf",
                    type="primary"
                )
            with col2:
                # Botón descargar Word
                if st.button("📄 Generar Word Editable", type="primary"):
                    with st.spinner("Generando documento Word con formato profesional..."):
                        word_buffer = generar_word_completo_mejorado(
                            empresa_id, 
                            pdf_version, 
                            pdf_elaborado, 
                            pdf_revisado, 
                            pdf_aprobado
                        )
                        if word_buffer:
                            st.session_state['word_bytes'] = word_buffer.getvalue()
                            st.session_state['word_nombre'] = f"Plan_Estrategico_{empresa_data.get('nombre', 'Empresa')}_V{pdf_version}.docx"
                            st.success("✅ Documento Word generado")
                            st.rerun()
                        else:
                            st.error("Error al generar Word")
                
                if 'word_bytes' in st.session_state:
                    st.download_button(
                        label="⬇️ Descargar Word", 
                        data=BytesIO(st.session_state['word_bytes']), 
                        file_name=st.session_state.get('word_nombre', 'plan_estrategico.docx'), 
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        type="secondary"
                    )
            with col3:
                # Mensaje actualizado que refleja ambos formatos
                if 'word_bytes' in st.session_state:
                    st.success("✅ Ambos documentos listos")
                    st.caption("📄 PDF para presentación final | 📝 Word para ediciones")
                else:
                    st.success("✅ PDF generado")
                    st.caption("Formato APA profesional. Genera el Word si necesitas editar.")
            
            # Botón para generar nuevo documento (FUERA de los columns)
            if st.button("🔄 Generar Nuevo", type="secondary"):
                if 'pdf_bytes' in st.session_state:
                    del st.session_state['pdf_bytes']
                if 'pdf_nombre' in st.session_state:
                    del st.session_state['pdf_nombre']
                if 'word_bytes' in st.session_state:
                    del st.session_state['word_bytes']
                if 'word_nombre' in st.session_state:
                    del st.session_state['word_nombre']
                st.session_state['pdf_generado'] = False
                st.rerun()
                
def pantalla_acceso():
    st.sidebar.title("Estratega Pro UG-UCE")
    opcion = st.sidebar.radio("Acceso al Sistema", ["Entrar", "Crear Cuenta"])
    col1, col2 = st.columns([2,1])
    with col1:
        if opcion == "Entrar":
            st.subheader("🔐 Iniciar Sesión")
            with st.form("login_form"):
                email = st.text_input("Correo Electrónico")
                password = st.text_input("Contraseña", type="password")
                if st.form_submit_button("Acceder"):
                    if supabase:
                        try:
                            res = supabase.auth.sign_in_with_password({"email": email, "password": password})
                            st.session_state.user = res.user
                            st.session_state.logged_in = True
                            st.rerun()
                        except Exception as e:
                            st.error(f"Error al iniciar sesión: {e}")
        else:
            st.subheader("📝 Registro de Nuevo Consultor")
            with st.form("signup_form"):
                nombre = st.text_input("Nombre Completo")
                email = st.text_input("Correo Electrónico")
                password = st.text_input("Contraseña", type="password")
                if st.form_submit_button("Finalizar Registro"):
                    if nombre and email and password and supabase:
                        try:
                            res = supabase.auth.sign_up({"email": email, "password": password, "options": {"data": {"full_name": nombre}}})
                            if res.user:
                                st.success("¡Registro exitoso! Ya puedes iniciar sesión.")
                            else:
                                st.error("No se pudo completar el registro.")
                        except Exception as e:
                            st.error(f"Error en el registro: {e}")
                    else:
                        st.warning("Por favor, llena todos los campos.")

def main():
    if "logged_in" not in st.session_state:
        st.session_state.logged_in = False
    if not st.session_state.get("logged_in"):
        pantalla_acceso()
    else:
        aplicacion_principal()

if __name__ == "__main__":
    supabase = init_supabase()
    if supabase:
        main()
    else:
        st.error("La aplicación no puede iniciarse. Revisa la conexión con la base de datos (Supabase).")





















